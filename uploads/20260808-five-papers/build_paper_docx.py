#!/usr/bin/env python3
"""Render the five rewritten manuscripts from constrained markdown to journal-format Word.

The markdown follows the fixed template the author agents were given (title,
CJK/EN abstracts, numbered sections, three-line tables, GB/T 7714 references),
so this parses that template rather than general markdown. Citation brackets
like [3] or [2-4] become superscript runs everywhere except the reference list.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x1A, 0x1A, 0x1A)
BODY_PT = 10.5           # 五号
REF_PT = 9               # 小五
CITE = re.compile(r"\[(\d+(?:\s*[,，\-–]\s*\d+)*)\]")
BOLD = re.compile(r"\*\*(.+?)\*\*")


def set_fonts(run, *, size=BODY_PT, bold=False, italic=False, east="宋体", west="Times New Roman"):
    run.font.name = west
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = INK
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)


def add_runs(p, text, *, size=BODY_PT, east="宋体", superscript_cites=True, base_bold=False):
    """Split text into bold / citation / plain runs."""
    pos = 0
    tokens = []
    for m in BOLD.finditer(text):
        tokens.append((pos, m.start(), False))
        tokens.append((m.start(1), m.end(1), True))
        pos = m.end()
    tokens.append((pos, len(text), False))
    for start, end, bold in tokens:
        chunk = text[start:end]
        if not chunk:
            continue
        cpos = 0
        for cm in CITE.finditer(chunk):
            if cm.start() > cpos:
                set_fonts(p.add_run(chunk[cpos:cm.start()]), size=size, bold=bold or base_bold, east=east)
            cite = p.add_run(cm.group(0))
            set_fonts(cite, size=size, bold=bold or base_bold, east=east)
            if superscript_cites:
                cite.font.superscript = True
            cpos = cm.end()
        if cpos < len(chunk):
            set_fonts(p.add_run(chunk[cpos:]), size=size, bold=bold or base_bold, east=east)


def three_line_table(doc, header, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 12), ("bottom", 12)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(cell.paragraphs[0], cell_text, size=9, base_bold=True)
        cell_borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1A1A1A")
        cell_borders.append(bottom)
        cell._tc.get_or_add_tcPr().append(cell_borders)
    for i, row in enumerate(rows, start=1):
        for j, cell_text in enumerate(row):
            if j >= len(header):
                break
            cell = table.cell(i, j)
            add_runs(cell.paragraphs[0], cell_text, size=9)
    return table


def paragraph(doc, *, align=None, indent=False, space_after=6, line=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Cm(0.74)
    return p


def parse_table_block(lines, i):
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    rows = []
    i += 2  # skip separator
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def build(md_path, out_path):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.8)
    section.top_margin = section.bottom_margin = Cm(2.5)

    in_refs = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):  # Chinese title
            p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line=1.3)
            add_runs(p, line[2:].strip(), size=15, east="黑体", base_bold=True, superscript_cites=False)
        elif line.startswith("**英文题目"):
            text = re.sub(r"^\*\*英文题目[:：]?\*\*\s*", "", line)
            p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            add_runs(p, text, size=12, base_bold=True, superscript_cites=False)
        elif line.startswith("## "):
            title = line[3:].strip()
            in_refs = title.startswith("参考文献")
            p = paragraph(doc, space_after=6, line=1.3)
            p.paragraph_format.space_before = Pt(12)
            add_runs(p, title, size=12 if title in ("摘要", "Abstract") or in_refs else 13,
                     east="黑体", base_bold=True, superscript_cites=False)
        elif line.startswith("### "):
            p = paragraph(doc, space_after=4, line=1.3)
            p.paragraph_format.space_before = Pt(8)
            add_runs(p, line[4:].strip(), size=11, east="黑体", base_bold=True, superscript_cites=False)
        elif line.strip().startswith("|"):
            if i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(":", "").strip()) <= {"-", " "}:
                header, rows, i = parse_table_block(lines, i)
                three_line_table(doc, header, rows)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                continue
            p = paragraph(doc, indent=True)
            add_runs(p, line.strip())
        elif in_refs and re.match(r"^\[\d+\]", line.strip()):
            p = paragraph(doc, space_after=3, line=1.3)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.9)
            pf.first_line_indent = Cm(-0.9)
            add_runs(p, line.strip(), size=REF_PT, superscript_cites=False)
        else:
            text = line.strip()
            is_label = text.startswith("**关键词") or text.startswith("**Keywords") or \
                text.startswith("**目的") or text.startswith("**Objective")
            table_title = re.match(r"^\*\*表\s*\d+\*\*", text)
            if table_title:
                p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
                add_runs(p, text, size=9.5, superscript_cites=False)
            else:
                p = paragraph(doc, indent=not is_label)
                add_runs(p, text, superscript_cites=not in_refs)
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src_dir = Path(__file__).parent / "papers"
    out_dir = Path(__file__).parent / "docx"
    out_dir.mkdir(exist_ok=True)
    targets = sys.argv[1:] or sorted(p.stem for p in src_dir.glob("*.md"))
    for stem in targets:
        md = src_dir / f"{stem}.md"
        first = md.read_text(encoding="utf-8").splitlines()[0].lstrip("# ").strip()
        safe = re.sub(r"[\\/:*?\"<>|]", "_", first)[:60]
        out = build(md, out_dir / f"{safe}.docx")
        print(f"{md.name} -> {out.name}")
