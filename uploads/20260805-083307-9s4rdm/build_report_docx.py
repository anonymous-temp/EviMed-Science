#!/usr/bin/env python3
"""Render the TDM research-design report to Word.

Content is held as a structure so the same source can be re-rendered after an
edit without retyping; every number in it traces to tdm-worked-example.py or to
a cited work.
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ACCENT = RGBColor(0xC1, 0x5F, 0x3C)
INK = RGBColor(0x1F, 0x1D, 0x1B)
MUTED = RGBColor(0x5A, 0x54, 0x4E)


def setup(doc):
    for name, size, bold, colour in (
        ("Normal", 10.5, False, INK),
        ("Heading 1", 16, True, ACCENT),
        ("Heading 2", 13, True, INK),
        ("Heading 3", 11.5, True, MUTED),
    ):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = colour
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if name == "Normal" else "黑体")
        pf = st.paragraph_format
        pf.space_before = Pt(14 if name.startswith("Heading") else 0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.4 if name == "Normal" else 1.15
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.6)
    section.top_margin = section.bottom_margin = Cm(2.4)


def para(doc, text, *, style=None, italic=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    for chunk, bold in _bold_split(text):
        run = p.add_run(chunk)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def _bold_split(text):
    """`**bold**` inside plain prose, without pulling in a markdown dependency."""
    out, rest = [], text
    while "**" in rest:
        head, _, tail = rest.partition("**")
        strong, _, rest = tail.partition("**")
        if head:
            out.append((head, False))
        out.append((strong, True))
    if rest:
        out.append((rest, False))
    return out or [(text, False)]


def bullets(doc, items, style="List Bullet"):
    for item in items:
        para(doc, item, style=style, space_after=2)


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(t.rows[0].cells, header):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = ""
            for chunk, bold in _bold_split(str(text)):
                run = cell.paragraphs[0].add_run(chunk)
                run.bold = bold
                run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Cm(w)
    doc.add_paragraph()
    return t


def caption(doc, text):
    para(doc, text, italic=True, size=9, space_after=10)


def build():
    import report_content as C
    doc = Document()
    setup(doc)

    para(doc, C.TITLE, style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, C.SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    para(doc, C.DATELINE, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True, space_after=16)

    para(doc, "摘要", style="Heading 2")
    for block in C.ABSTRACT.split("\n\n"):
        para(doc, block)

    para(doc, "一、数据来源与字段结构", style="Heading 2")
    para(doc, C.DATA_OVERVIEW_INTRO)
    table(doc, *C.TABLE_INVENTORY, widths=[2.6, 2.2, 2.4, 8.2])
    para(doc, C.ANALYTES_INTRO)
    table(doc, *C.TABLE_ANALYTES, widths=[4.2, 1.6, 4.2, 5.4])

    para(doc, "二、编码规范与数据治理发现", style="Heading 2")
    para(doc, C.CODING_INTRO)
    for heading, body in C.CODING_FINDINGS:
        para(doc, heading, style="Heading 3")
        para(doc, body)

    para(doc, "三、派生量与一致性校验", style="Heading 2")
    para(doc, C.DERIVED_INTRO)
    para(doc, "3.1 总量恒等式：全部后续推断的合法性前提", style="Heading 3")
    table(doc, *C.TABLE_IDENTITY, widths=[2.0, 2.8, 2.0, 2.0, 2.2, 2.2, 2.2])
    para(doc, C.IDENTITY_NOTE)
    para(doc, "3.2 代谢比：数据中已有的 CYP2D6 表型代理", style="Heading 3")
    table(doc, *C.TABLE_MR, widths=[2.0, 2.8, 2.6, 2.8, 2.0, 3.2])
    para(doc, C.MR_NOTE)
    para(doc, "3.3 剂量归一化浓度：由医嘱重建的暴露", style="Heading 3")
    table(doc, *C.TABLE_CD, widths=[2.0, 2.8, 5.0, 2.4, 3.2])
    para(doc, C.CD_NOTE)

    para(doc, "四、分析思路：由一个测量点向外扩展", style="Heading 2")
    para(doc, C.EXPANSION_INTRO)
    bullets(doc, C.EXPANSION_STEPS)
    para(doc, C.VITALS_NOTE)

    para(doc, "五、五个研究方向", style="Heading 2")
    for group in C.GROUPS:
        para(doc, f"{group['id']}　{group['title']}", style="Heading 3")
        para(doc, "文献小综述", style="Heading 3")
        for block in group["review"].split("\n\n"):
            para(doc, block)
        para(doc, "设计思路", style="Heading 3")
        for block in group["rationale"].split("\n\n"):
            para(doc, block)
        para(doc, "研究方案", style="Heading 3")
        for label, body in group["design"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}——")
            run.bold = True
            for chunk, bold in _bold_split(body):
                r = p.add_run(chunk)
                r.bold = bold
            p.paragraph_format.space_after = Pt(4)

    para(doc, "六、跨领域方法借鉴与转移边界", style="Heading 2")
    para(doc, C.CROSSFIELD_INTRO)
    table(doc, *C.TABLE_CROSSFIELD, widths=[1.4, 3.4, 2.8, 4.4, 4.4])

    para(doc, "七、字段使用总表", style="Heading 2")
    para(doc, C.FIELDUSE_INTRO)
    table(doc, *C.TABLE_FIELDUSE, widths=[4.0, 2.0, 4.6, 5.8])

    para(doc, "八、全量数据需求与治理要求", style="Heading 2")
    para(doc, C.GOVERNANCE_INTRO)
    bullets(doc, C.GOVERNANCE_ITEMS)

    para(doc, "九、局限", style="Heading 2")
    bullets(doc, C.LIMITS)

    para(doc, "十、实施次序", style="Heading 2")
    para(doc, C.CLOSING)

    doc.save("住院TDM数据研究方向与实施方案.docx")
    print("saved 住院TDM数据研究方向与实施方案.docx")


if __name__ == "__main__":
    build()
