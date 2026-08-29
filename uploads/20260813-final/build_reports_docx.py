#!/usr/bin/env python3
"""Render the platform's evidence reports to Word for reading.

These are the pipeline's own deliverables, not the hand-rewritten manuscripts:
markdown with `[n](url)` citations, hidden `<!-- claim:CLM-nnn -->` markers and
GFM tables. The markers are provenance for the gate, not for a reader, so they
come out; the citation numbers stay and become superscripts.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).parent
INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0xC1, 0x5F, 0x3C)
MUTED = RGBColor(0x5A, 0x54, 0x4E)
BODY_PT = 10.5

CLAIM_MARKER = re.compile(r"<!--\s*claim:[A-Z]+-\d+\s*-->")
HTML_COMMENT = re.compile(r"<!--.*?-->", re.S)
# `[3](https://…)` and bare `[3]` both mean the same thing to a reader.
LINKED_CITE = re.compile(r"\[(\d+)\]\((?:https?://)[^)]*\)")
CITE = re.compile(r"\[(\d+(?:\s*[,，]\s*\d+)*)\]")
BOLD = re.compile(r"\*\*(.+?)\*\*")


def clean(text):
    text = CLAIM_MARKER.sub("", text)
    text = HTML_COMMENT.sub("", text)
    text = LINKED_CITE.sub(r"[\1]", text)
    return re.sub(r"[ \t]{2,}", " ", text)


def font(run, *, size=BODY_PT, bold=False, east="宋体", colour=INK):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = colour
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)


def add_runs(p, text, *, size=BODY_PT, east="宋体", base_bold=False, superscript=True, colour=INK):
    pos, tokens = 0, []
    for m in BOLD.finditer(text):
        tokens.append((pos, m.start(), False))
        tokens.append((m.start(1), m.end(1), True))
        pos = m.end()
    tokens.append((pos, len(text), False))
    for start, end, bold in tokens:
        chunk = text[start:end]
        if not chunk:
            continue
        cpos = 0
        for cm in CITE.finditer(chunk) if superscript else []:
            if cm.start() > cpos:
                font(p.add_run(chunk[cpos:cm.start()]), size=size, bold=bold or base_bold, east=east, colour=colour)
            cite = p.add_run(cm.group(0))
            font(cite, size=size, bold=bold or base_bold, east=east, colour=colour)
            cite.font.superscript = True
            cpos = cm.end()
        if cpos < len(chunk):
            font(p.add_run(chunk[cpos:]), size=size, bold=bold or base_bold, east=east, colour=colour)


def para(doc, *, align=None, indent=False, after=6, line=1.45, before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Cm(0.74)
    return p


def three_line_table(doc, header, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 12), ("bottom", 12)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    table._tbl.tblPr.append(borders)
    for j, text in enumerate(header):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(cell.paragraphs[0], text, size=8.5, base_bold=True)
        tc = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1A1A1A")
        tc.append(bottom)
        cell._tc.get_or_add_tcPr().append(tc)
    for i, row in enumerate(rows, start=1):
        for j, text in enumerate(row[:len(header)]):
            add_runs(table.cell(i, j).paragraphs[0], text, size=8.5)


def render(doc, md_text):
    lines = clean(md_text).splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            sizes = {1: 15, 2: 12.5, 3: 11.5}
            p = para(doc, after=5, line=1.25, before=12 if level > 1 else 0,
                     align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else None)
            add_runs(p, title, size=sizes.get(level, 11), east="黑体", base_bold=True,
                     superscript=False, colour=ACCENT if level == 1 else INK)
        elif line.strip().startswith("|") and i + 1 < len(lines) and \
                set(lines[i + 1].replace("|", "").replace(":", "").strip()) <= {"-", " "}:
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows, i = [], i + 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            three_line_table(doc, header, rows)
            para(doc, after=4)
            continue
        elif re.match(r"^\s*[-*]\s+", line):
            p = para(doc, after=3, line=1.4)
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, "· " + re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = para(doc, after=3, line=1.4)
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, line.strip())
        elif line.startswith(">"):
            p = para(doc, after=4)
            p.paragraph_format.left_indent = Cm(0.8)
            add_runs(p, line.lstrip("> ").strip(), size=9.5, colour=MUTED)
        else:
            add_runs(para(doc, indent=not line.strip().startswith("**")), line.strip())
        i += 1


def setup(doc):
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.6)
    section.top_margin = section.bottom_margin = Cm(2.4)


def main():
    src = BASE / "reports"
    out = BASE / "docx"
    out.mkdir(exist_ok=True)
    dirs = sorted(d for d in src.iterdir() if d.is_dir() and d.name.startswith("RQ-"))
    combined = Document()
    setup(combined)
    written = 0
    for d in dirs:
        report = d / "clinical-evidence-report.md"
        if not report.exists():
            report = d / "safety-report.md"
        if not report.exists():
            continue
        rq = d.name.split("_")[0]
        text = report.read_text(encoding="utf-8")

        single = Document()
        setup(single)
        render(single, text)
        title = clean(text).splitlines()[0].lstrip("# ").strip()
        safe = re.sub(r'[\\/:*?"<>|]', "_", f"{rq} {title}")[:80]
        single.save(out / f"{safe}.docx")

        if written:
            combined.add_section(WD_SECTION.NEW_PAGE)
        render(combined, text)
        written += 1
        print(f"{rq}  {title[:52]}")
    combined.save(out / "速效救心丸语义问题深度分析报告合集.docx")
    print(f"\n{written} 篇 -> {out}")


if __name__ == "__main__":
    main()
