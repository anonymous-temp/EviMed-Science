from pathlib import Path
from io import BytesIO
from uuid import uuid4
import json
import zipfile

from docx import Document

from new_meta.core.artifact_package import (
    create_artifact_package,
    _build_citation_audit_review,
    _build_clinical_interpretation_audit_review,
    _build_cross_reference_audit_review,
    _build_figure_audit_review,
    _build_figure_legend_audit_review,
    _build_manuscript_polish_audit_review,
    _build_prisma_audit_review,
    _build_submission_readiness_review,
    _build_table_footnote_audit_review,
    _has_publication_section_shape,
    _manuscript_manifest_summary,
    _manuscript_content_summary,
    _review_language_from_text,
)
from new_meta.core.artifact_package_entries import iter_existing_package_files
from new_meta.core.artifact_package_publication_similarity import build_publication_similarity_review
from new_meta.core.benchmark_review import build_benchmark_review_payload
from new_meta.core.benchmark_source_decisions import save_benchmark_source_decision
from new_meta.core.manuscript_facts import validate_and_repair_manuscript
from new_meta.core.project import Project
from new_meta.agents.writing_agent import WritingAgent
from new_meta.schemas.study import ExtractedStudy, OutcomeData, StudyCharacteristics


MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05"
    b"\xfe\x02\xfeA\xe2\x95\xb3\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_artifact_package_orchestrator_stays_below_monolith_budget() -> None:
    package_path = Path("new_meta/core/artifact_package.py")
    html_module_path = Path("new_meta/core/artifact_package_html.py")
    manifest_module_path = Path("new_meta/core/artifact_package_manifest.py")

    package_text = package_path.read_text(encoding="utf-8")
    manifest_text = manifest_module_path.read_text(encoding="utf-8") if manifest_module_path.exists() else ""

    assert len(package_text.splitlines()) <= 3500
    assert "def _render_abstract_audit_html" not in package_path.read_text(encoding="utf-8")
    assert "def _render_abstract_audit_html" in html_module_path.read_text(encoding="utf-8")
    assert "def _manuscript_manifest_summary" not in package_text
    assert "def manuscript_manifest_summary" in manifest_text
    assert "def review_manifest_summary" in manifest_text


def test_artifact_package_keeps_citation_audit_in_dedicated_module() -> None:
    package_text = Path("new_meta/core/artifact_package.py").read_text(encoding="utf-8")
    citation_module = Path("new_meta/core/artifact_package_citation_audit.py")
    citation_text = citation_module.read_text(encoding="utf-8") if citation_module.exists() else ""

    assert "def _build_citation_audit_review" not in package_text
    assert "def _uncited_discussion_context_claims" not in package_text
    assert "def build_citation_audit_review" in citation_text
    assert "def uncited_discussion_context_claims" in citation_text


def test_artifact_package_uses_validation_main_text_word_count(tmp_path: Path) -> None:
    project = Project("manifest word count", output_dir=tmp_path / uuid4().hex)
    table_payload = " ".join(["tableword"] * 400)
    manuscript = "\n\n".join([
        "# Meta-analysis",
        "## Abstract",
        "Brief abstract.",
        "## Introduction",
        "Brief introduction.",
        "## Methods",
        "Brief methods.",
        "## Results",
        "Brief results.",
        "### Table 1. Long extracted evidence table",
        "| Study | Source quote |",
        "|---|---|",
        f"| Trial A | {table_payload} |",
        "## Discussion",
        "Brief discussion.",
        "## References",
        "1. Trial reference.",
    ])
    facts = {
        "report_type": "meta",
        "primary_effect": {"n_studies": 2},
        "evidence_readiness": {"status": "ready", "blockers": [], "warnings": []},
    }
    project.save_text("draft.md", manuscript, subdir="manuscript")
    project.save_json("manuscript_facts.json", facts, subdir="manuscript")

    _, validation = validate_and_repair_manuscript(manuscript, facts)
    manifest = _manuscript_manifest_summary(project)

    assert manifest["main_word_count"] == validation["facts_summary"]["main_word_count"]
    assert manifest["main_word_count"] < 100


def test_citation_audit_uses_validation_main_text_word_count(tmp_path: Path) -> None:
    project = Project("citation word count", output_dir=tmp_path / uuid4().hex)
    table_payload = " ".join(["tableword"] * 400)
    manuscript = "\n\n".join([
        "# Meta-analysis",
        "## Abstract",
        "Brief abstract [1].",
        "## Introduction",
        "Brief introduction [1].",
        "## Methods",
        "Brief methods [1].",
        "## Results",
        "Brief results [1].",
        "### Table 1. Long extracted evidence table",
        "| Study | Source quote |",
        "|---|---|",
        f"| Trial A | {table_payload} |",
        "## Discussion",
        "Brief discussion [1].",
        "## References",
        "[1] Trial reference.",
    ])
    facts = {
        "report_type": "meta",
        "primary_effect": {"n_studies": 2},
        "evidence_readiness": {"status": "ready", "blockers": [], "warnings": []},
    }
    project.save_text("draft.md", manuscript, subdir="manuscript")
    project.save_json("manuscript_facts.json", facts, subdir="manuscript")

    _, validation = validate_and_repair_manuscript(manuscript, facts)
    audit = _build_citation_audit_review(project)

    assert audit["summary"]["main_text_word_count"] == validation["facts_summary"]["main_word_count"]
    assert audit["summary"]["main_text_word_count"] < 100


def test_artifact_package_uses_two_study_adaptive_length_target(tmp_path: Path) -> None:
    project = Project("adaptive length target", output_dir=tmp_path / uuid4().hex)
    manuscript = "\n\n".join([
        "# Meta-analysis",
        "## Abstract",
        " ".join(["abstract"] * 250),
        "## Introduction",
        " ".join(["introduction"] * 1200),
        "## Methods",
        " ".join(["methods"] * 1200),
        "## Results",
        " ".join(["results"] * 900),
        "## Discussion",
        " ".join(["discussion"] * 1450),
        "## References",
        "1. Trial reference.",
    ])
    facts = {
        "report_type": "meta",
        "primary_effect": {"n_studies": 2},
        "evidence_readiness": {"status": "ready", "blockers": [], "warnings": []},
    }
    project.save_text("draft.md", manuscript, subdir="manuscript")
    project.save_json("manuscript_facts.json", facts, subdir="manuscript")

    manifest = _manuscript_manifest_summary(project)

    assert manifest["minimum_main_words"] == 4500
    assert manifest["main_word_count"] >= manifest["minimum_main_words"]


def test_artifact_package_migrates_legacy_two_study_length_target(tmp_path: Path) -> None:
    project = Project("legacy adaptive length target", output_dir=tmp_path / uuid4().hex)
    manuscript = "\n\n".join([
        "# Meta-analysis",
        "## Abstract",
        " ".join(["abstract"] * 250),
        "## Introduction",
        " ".join(["introduction"] * 1200),
        "## Methods",
        " ".join(["methods"] * 1200),
        "## Results",
        " ".join(["results"] * 900),
        "## Discussion",
        " ".join(["discussion"] * 1450),
        "## References",
        "1. Trial reference.",
    ])
    facts = {
        "report_type": "meta",
        "primary_effect": {"n_studies": 2},
        "evidence_readiness": {"status": "ready", "blockers": [], "warnings": []},
        "writing_constraints": {"publication_min_main_words": 6000},
    }
    project.save_text("draft.md", manuscript, subdir="manuscript")
    project.save_json("manuscript_facts.json", facts, subdir="manuscript")

    manifest = _manuscript_manifest_summary(project)

    assert manifest["minimum_main_words"] == 4500
    assert manifest["main_word_count"] >= manifest["minimum_main_words"]


def _clinical_discussion_fixture(effect: str = "OR 0.66", endpoint: str = "mortality") -> list[str]:
    return [
        "## Discussion",
        f"The pooled {effect} should be interpreted as the direction and magnitude of effect for {endpoint} [1].",
        "Clinical translation depends on baseline risk, absolute risk difference, and number needed to treat.",
        "The endpoint should be interpreted by its clinical components and follow-up time.",
        "Benefit-harm balance requires safety, adverse events, tolerability, and treatment discontinuation to be considered [1].",
        "Applicability depends on patient age, comorbidity, disease severity, renal function, and background therapy.",
        "Implementation requires monitoring, follow-up, patient preference, cost, and access to shape decisions.",
        "Certainty of evidence, heterogeneity, publication bias, and other limitations should temper inference [1].",
        "",
        "## Conclusion",
        f"The result supports cautious clinical interpretation of {effect} for {endpoint}, individualized to baseline risk, safety, certainty, and patient preferences [1].",
        "",
    ]


def _save_passing_submission_quality_gate(project: Project) -> None:
    project.save_json(
        "submission_quality_gate.json",
        {
            "status": "pass",
            "failed_count": 0,
            "warning_count": 0,
            "checks": [
                {"name": "claim_source_resolution", "status": "pass", "message": "ok"},
                {"name": "citation_contract", "status": "pass", "message": "ok"},
                {"name": "claim_map_authoring", "status": "pass", "message": "ok"},
            ],
        },
        subdir="manuscript",
    )


def test_create_artifact_package_collects_submission_files(tmp_path: Path) -> None:
    project = Project("artifact package", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# Manuscript", subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "grade": {
                "outcomes": [
                    {
                        "outcome_name": "28-day mortality",
                        "n_studies": 1,
                        "certainty": "Low",
                        "effect_summary": "OR 0.45 (95% CI 0.20 to 0.99)",
                        "domains": [
                            {
                                "domain": "indirectness",
                                "rating": "serious",
                                "rationale": "Comparator source metadata were incomplete.",
                                "details": {
                                    "method": "rule_based_pico_directness_v1",
                                    "n_contributing": 1,
                                    "target_outcome": "28-day mortality",
                                    "protocol_primary_outcome": "28-day mortality",
                                    "source_verified_direct_rows": False,
                                    "surrogate_outcome": False,
                                    "dimensions": {
                                        "population": {"mismatch": 0, "unverified": 0, "total": 1},
                                        "intervention": {"mismatch": 0, "unverified": 0, "total": 1},
                                        "comparator": {"mismatch": 0, "unverified": 1, "total": 1},
                                        "outcome": {"mismatch": 0, "unverified": 0, "total": 1},
                                        "design": {"non_randomized": 0, "unverified": 0, "total": 1},
                                    },
                                },
                            },
                        ],
                    }
                ]
            },
            "evidence_readiness": {
                "status": "blocked",
                "blocker_codes": [
                    "primary_counts_not_source_verified",
                    "primary_timepoint_not_source_verified",
                ],
                "blockers": [
                    {
                        "code": "primary_counts_not_source_verified",
                        "row_id": "S1:0",
                        "missing_values": ["total_intervention=76"],
                    },
                    {
                        "code": "primary_timepoint_not_source_verified",
                        "row_id": "S1:0",
                    },
                ],
                "warnings": [],
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "28-day mortality",
                        "events_intervention": 11,
                        "total_intervention": 76,
                        "events_control": 20,
                        "total_control": 73,
                        "source_quote": "11 deaths and 20 deaths.",
                        "source_location": "Abstract",
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True}, subdir="manuscript")
    project.save_json("manuscript_quality_gate.json", {"passed": True}, subdir="manuscript")
    project.save_json(
        "submission_quality_gate.json",
        {
            "status": "pass",
            "failed_count": 0,
            "warning_count": 0,
            "checks": [
                {"name": "claim_source_resolution", "status": "pass", "message": "ok"},
                {"name": "citation_contract", "status": "pass", "message": "ok"},
                {"name": "claim_map_authoring", "status": "pass", "message": "ok"},
            ],
        },
        subdir="manuscript",
    )
    project.save_json("quality_gate.json", {"status": "pass"}, subdir="manuscript")
    project.save_json(
        "claim_map.json",
        [
            {
                "id": "intro_objective",
                "section": "Introduction",
                "claim": "The review assesses source-verified mortality evidence.",
                "claim_type": "objective",
                "can_write_main_text": True,
                "source_refs": [{"source_type": "structured_fact", "source_id": "protocol"}],
            }
        ],
        subdir="manuscript",
    )
    project.save_json(
        "claim_map_audit.json",
        {"status": "ok", "writable_claims": 1},
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_plan.json",
        {"sections": [{"section": "Introduction", "claim_ids": ["intro_objective"]}]},
        subdir="manuscript",
    )
    project.save_json(
        "claim_source_resolution_audit.json",
        {"summary": {"unresolved_count": 0}, "unresolved_claims": []},
        subdir="manuscript",
    )
    project.save_json(
        "claim_source_alignment_audit.json",
        {"status": "ok", "alignment_input_hash": "a" * 64, "items": []},
        subdir="manuscript",
    )
    project.save_json(
        "citation_contract.json",
        {
            "schema_version": 1,
            "status": "ok",
            "items": [
                {
                    "claim_id": "intro_objective",
                    "citation": "[1]",
                    "reference_numbers": [1],
                    "source_spans": [{"source_id": "protocol", "quote": "mortality evidence"}],
                }
            ],
        },
        subdir="manuscript",
    )
    project.save_json(
        "claim_map_citation_plan.json",
        {"status": "ok", "claim_ids": ["intro_objective"]},
        subdir="manuscript",
    )
    project.save_json(
        "final_claim_map_citation_plan.json",
        {"status": "ok", "claim_ids": ["intro_objective"]},
        subdir="manuscript",
    )
    project.save_json(
        "claim_map_authoring_audit.json",
        {"status": "ok", "accepted_sections": 1, "rejected_sections": []},
        subdir="manuscript",
    )
    project.save_json(
        "citation_audit_review.json",
        {"status": "ok", "unsupported_citations": []},
        subdir="manuscript",
    )
    project.save_json(
        "citation_grounding_audit.json",
        {"status": "ok", "unsupported_claims": []},
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_semantic_edit_audit.json",
        {"status": "ok", "edits": []},
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_style_audit.json",
        {"status": "ok", "issues": []},
        subdir="manuscript",
    )
    project.save_json(
        "low_k_methodology_review_audit.json",
        {"status": "ok", "k": 2},
        subdir="manuscript",
    )
    project.save_text("references.bib", "@article{x}")
    project.save_json("pipeline_warnings.json", [{"stage": "figures"}])
    project.save_json(
        "protocol_overrides.json",
        {
            "schema_version": 1,
            "current_revision": 1,
            "overrides": [
                {
                    "revision": 1,
                    "updated_by": "tester",
                    "fields": {"effect_measure": {"old": "RR", "new": "OR"}},
                }
            ],
        },
    )
    project.save_json(
        "llm_usage_manifest.json",
        {
            "schema_version": 1,
            "summary": {
                "total_calls": 3,
                "total_tokens": 9000,
                "prompt_tokens": 3000,
                "completion_tokens": 6000,
                "estimated_cost_usd": 0.010728,
            },
            "events": [],
        },
    )
    project.save_json(
        "all_extractions.json",
        [
            ExtractedStudy(
                characteristics=StudyCharacteristics(
                    study_id="S1",
                    pmid="S1",
                    doi="10.1000/artifact",
                    title="Artifact Trial",
                    source_type="user_upload",
                    pdf_path="/tmp/artifact.pdf",
                ),
                outcomes=[
                    OutcomeData(
                        outcome_name="28-day mortality",
                        outcome_type="dichotomous",
                        events_intervention=11,
                        total_intervention=76,
                        events_control=20,
                        total_control=73,
                        source_quote="11 deaths and 20 deaths.",
                        source_location="Abstract",
                        source_quote_verified=True,
                        extraction_confidence="medium",
                    )
                ],
            )
        ],
        subdir="extraction",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "summary": {},
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "study_label": "Smith 2024",
                    "title": "Artifact Trial",
                    "outcome_name": "28-day mortality",
                    "outcome_type": "dichotomous",
                    "source_quote": "11 deaths and 20 deaths.",
                    "source_location": "Abstract",
                    "source_quote_verified": True,
                    "extraction_confidence": "medium",
                    "requires_review": True,
                    "conflicts": [
                        {
                            "field": "total_intervention",
                            "message": "Needs source-backed total.",
                            "sources": ["schema_count_validation"],
                        }
                    ],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {
                "full_text": (
                    "[PAGE 4]\nResults text before the table. "
                    "11 deaths and 20 deaths. "
                    "Results text after the table."
                ),
                "page_map": [{"page_number": 4, "start_char": 0, "end_char": 92}],
            }
        },
        subdir="papers",
    )
    project.save_json(
        "extraction_overrides.json",
        {"schema_version": 1, "current_revision": 2, "overrides": []},
        subdir="extraction",
    )
    project.save_json("effect_selection_audit.json", [], subdir="analysis")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_benchmark",
            "summary_card": {
                "benchmark_id": "mini_benchmark",
                "status": "blocked",
                "passed": False,
                "published_anchor": {
                    "n_trials": 2,
                    "n_participants": 120,
                    "effect_measure": "OR",
                    "effect": 0.66,
                    "ci_lower": 0.53,
                    "ci_upper": 0.82,
                },
                "observed_primary": {
                    "effect_measure": "RR",
                    "n_studies": 1,
                    "effect": 0.74,
                    "ci_lower": 0.60,
                    "ci_upper": 0.91,
                    "total_participants": 80,
                    "participant_difference": -40,
                },
                "failing_gates": [
                    {
                        "gate": "primary_full_text_recall",
                        "label": "Primary full-text recall",
                        "passed": False,
                        "matched": 1,
                        "total": 2,
                        "failure_reasons": ["primary_full_text_recall_below_threshold"],
                    }
                ],
                "missing_primary_full_texts": [
                    {
                        "trial_id": "trial_b",
                        "trial_name": "Trial B",
                        "publication_pmids": ["12345"],
                        "publication_dois": ["10.1000/trial-b"],
                    }
                ],
                "next_actions": [
                    {
                        "type": "upload_full_texts",
                        "message": "Upload source PDF/HTML for the missing expected publications before accepting a publication-style manuscript.",
                    }
                ],
            },
            "primary_analysis": {
                "missing": [
                    {
                        "trial_id": "trial_b",
                        "trial_name": "Trial B",
                        "expected_events_intervention": 11,
                        "expected_total_intervention": 20,
                        "expected_events_control": 20,
                        "expected_total_control": 40,
                    }
                ]
            },
        },
        subdir="benchmark",
    )
    project.save_json(
        "benchmark_summary_card.json",
        {"benchmark_id": "mini_benchmark", "status": "blocked", "passed": False},
        subdir="benchmark",
    )
    project.save_json(
        "benchmark_source_manifest.json",
        {
            "schema_version": 1,
            "sources": [
                {
                    "task_id": "full_text:trial_b",
                    "trial_id": "trial_b",
                    "trial_name": "Trial B",
                    "source_kind": "supplement",
                    "filename": "trial-b-supplement.pdf",
                    "local_path": "/tmp/trial-b-supplement.pdf",
                    "status": "uploaded_needs_review",
                    "sha256": "abc",
                    "size_bytes": 123,
                    "parse_status": "ok",
                    "parsed_path": "benchmark/source_parsed/abc.json",
                    "text_chars": 240,
                    "page_count": 2,
                    "table_count": 1,
                    "text_preview": "Trial B appendix mortality table.",
                }
            ],
        },
        subdir="benchmark",
    )
    project.save_json(
        "abc.json",
        {
            "full_text": "Trial B appendix mortality table: 11 deaths among 20 treatment participants and 20 deaths among 40 control participants.",
            "tables": ["| deaths | total |"],
            "page_map": [{"page_number": 3, "start_char": 0, "end_char": 118}],
        },
        subdir="benchmark/source_parsed",
    )
    benchmark_before_package = build_benchmark_review_payload(project)
    benchmark_task = benchmark_before_package["source_acquisition_tasks"][0]
    benchmark_source = benchmark_task["uploaded_sources"][0]
    benchmark_candidate = benchmark_source["quote_candidates"][0]
    save_benchmark_source_decision(
        project,
        task_id=benchmark_task["task_id"],
        trial_id=benchmark_task["trial_id"],
        source=benchmark_source,
        candidate=benchmark_candidate,
        decision="accepted",
        reason="Reviewer confirmed source table.",
        updated_by="tester",
        expected_revision=0,
    )
    project.save_json(
        "benchmark_source_applications.json",
        {
            "schema_version": 1,
            "current_revision": 1,
            "applications": [
                {
                    "candidate_id": benchmark_candidate["candidate_id"],
                    "task_id": benchmark_task["task_id"],
                    "trial_id": benchmark_task["trial_id"],
                    "action": "updated_existing_outcome",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "applied_fields": ["events_intervention"],
                    "values_applied": {"events_intervention": 11},
                    "previous_values": {"events_intervention": None},
                    "updated_by": "tester",
                    "revision": 1,
                }
            ],
        },
        subdir="benchmark",
    )
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(b"fake-png")

    package_path = create_artifact_package(project)

    assert package_path == project.base_dir / "package" / "metaagent_export.zip"
    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        review = json.loads(zf.read("review/evidence_readiness_review.json"))
        review_html = zf.read("review/extraction_review.html").decode("utf-8")
        benchmark_review = json.loads(zf.read("review/benchmark_review.json"))
        benchmark_html = zf.read("review/benchmark_review.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))
        llm_usage = json.loads(zf.read("llm_usage_manifest.json"))

    assert "manuscript/draft.md" in names
    assert "manuscript/manuscript_facts.json" in names
    assert "manuscript/manuscript_validation.json" in names
    assert "manuscript/manuscript_quality_gate.json" in names
    assert "manuscript/submission_quality_gate.json" in names
    assert "manuscript/quality_gate.json" in names
    assert "manuscript/claim_map.json" in names
    assert "manuscript/claim_map_audit.json" in names
    assert "manuscript/manuscript_plan.json" in names
    assert "manuscript/claim_source_resolution_audit.json" in names
    assert "manuscript/claim_source_alignment_audit.json" in names
    assert "manuscript/citation_contract.json" in names
    assert "manuscript/claim_map_citation_plan.json" in names
    assert "manuscript/final_claim_map_citation_plan.json" in names
    assert "manuscript/claim_map_authoring_audit.json" in names
    assert "manuscript/citation_audit_review.json" in names
    assert "manuscript/citation_grounding_audit.json" in names
    assert "manuscript/manuscript_semantic_edit_audit.json" in names
    assert "manuscript/manuscript_style_audit.json" in names
    assert "manuscript/low_k_methodology_review_audit.json" in names
    assert "references.bib" in names
    assert "pipeline_warnings.json" in names
    assert "protocol_overrides.json" in names
    assert "llm_usage_manifest.json" in names
    assert "extraction/extraction_audit.json" in names
    assert "analysis/effect_selection_audit.json" in names
    assert "benchmark/benchmark_report.json" in names
    assert "benchmark/benchmark_summary_card.json" in names
    assert "benchmark/benchmark_source_manifest.json" in names
    assert "benchmark/benchmark_source_decisions.json" in names
    assert "benchmark/benchmark_source_applications.json" in names
    assert "benchmark/source_parsed/abc.json" in names
    assert "figures/forest_plot.png" in names
    assert "review/evidence_readiness_review.json" in names
    assert "review/extraction_review.html" in names
    assert "review/benchmark_review.json" in names
    assert "review/benchmark_review.html" in names
    assert "package_manifest.json" in names
    assert set(manifest["entries"]) == names
    assert manifest["entry_count"] == len(names)
    assert manifest["review"]["included"] is True
    assert manifest["review"]["html_review"] is True
    assert manifest["review"]["blocker_count"] == 2
    assert manifest["review"]["extraction_source_cards"] == 1
    assert manifest["review"]["extraction_review_cards"] == 1
    assert manifest["review"]["source_context_available_cards"] == 1
    assert manifest["review"]["source_context_missing_cards"] == 0
    assert manifest["review"]["source_context_coverage"] == 1.0
    assert manifest["review"]["selected_primary_source_cards"] == 1
    assert manifest["review"]["selected_primary_source_context_available_cards"] == 1
    assert manifest["review"]["selected_primary_source_context_missing_cards"] == 0
    assert manifest["review"]["selected_primary_source_context_coverage"] == 1.0
    assert manifest["review"]["grade_review_domains"] == 1
    assert manifest["review"]["benchmark_included"] is True
    assert manifest["review"]["benchmark_status"] == "blocked"
    assert manifest["review"]["benchmark_failing_gates"] == 1
    assert manifest["review"]["benchmark_source_acquisition_tasks"] == 1
    assert manifest["review"]["benchmark_attached_source_tasks"] == 1
    assert manifest["review"]["benchmark_accepted_source_candidates"] == 1
    assert manifest["review"]["benchmark_source_applications"] == 1
    assert manifest["review"]["benchmark_html_review"] is True
    assert manifest["llm_usage"]["included"] is True
    assert manifest["llm_usage"]["total_calls"] == 3
    assert manifest["llm_usage"]["total_tokens"] == 9000
    assert manifest["llm_usage"]["estimated_cost_usd"] == 0.010728
    assert llm_usage["summary"]["prompt_tokens"] == 3000
    assert review["primary_count_verification_rows"][0]["missing_values"] == ["total_intervention=76"]
    assert review["summary"]["source_context_available_cards"] == 1
    assert review["summary"]["source_context_missing_cards"] == 0
    assert review["summary"]["source_context_coverage"] == 1.0
    assert review["summary"]["selected_primary_source_context_available_cards"] == 1
    assert review["summary"]["selected_primary_source_context_missing_cards"] == 0
    assert review["summary"]["selected_primary_source_context_coverage"] == 1.0
    assert review["missing_source_context_cards"] == []
    assert review["missing_selected_primary_source_context_cards"] == []
    assert review["timepoint_adjudication_rows"][0]["requires_user_adjudication"] is True
    assert review["grade_review"]["outcomes"][0]["outcome_name"] == "28-day mortality"
    grade_domain = review["grade_review"]["outcomes"][0]["domains"][0]
    assert grade_domain["domain"] == "indirectness"
    assert grade_domain["details"]["dimensions"]["comparator"]["unverified"] == 1
    source_card = review["extraction_source_cards"][0]
    assert source_card["source"]["quote_verified"] is True
    assert source_card["source_anchor"]["kind"] == "pdf_text_quote"
    assert source_card["source_anchor"]["highlight_text"] == "11 deaths and 20 deaths."
    assert source_card["source_anchor"]["can_open_pdf"] is True
    assert source_card["trust"]["status"] == "needs_review"
    assert source_card["trust"]["confidence"] == "medium"
    assert source_card["override"]["current_revision"] == 2
    assert source_card["review_action"]["current_revision"] == 0
    assert source_card["review_action"]["save_message_type"] == "extraction_review_decision"
    assert source_card["review_action"]["suggested_decision"]["row_id"] == "S1:0"
    assert source_card["review_action"]["suggested_decision"]["decision"] == "accepted"
    values = {item["field"]: item for item in source_card["values"]}
    assert values["events_intervention"]["value"] == 11
    assert values["total_intervention"]["conflicts"][0]["message"] == "Needs source-backed total."
    assert "MetaAgent Extraction Review" in review_html
    assert "GRADE Quality Review" in review_html
    assert "28-day mortality" in review_html
    assert "indirectness" in review_html
    assert "comparator" in review_html
    assert "unverified" in review_html
    assert "Artifact Trial" in review_html
    assert "28-day mortality" in review_html
    assert "11 deaths and 20 deaths." in review_html
    assert "Trust status: needs_review" in review_html
    assert "Source anchor: pdf_text_quote" in review_html
    assert "Highlight: 11 deaths and 20 deaths." in review_html
    assert "Source Context" in review_html
    assert "Results text before the table." in review_html
    assert "<mark>11 deaths and 20 deaths.</mark>" in review_html
    assert "Results text after the table." in review_html
    assert "Open PDF: /tmp/artifact.pdf" in review_html
    assert "Needs source-backed total." in review_html
    assert "Confirm this row" in review_html
    assert "Review Decision Payload Seed" in review_html
    assert "extraction_review_decision" in review_html
    assert benchmark_review["status"] == "blocked"
    assert benchmark_review["summary"]["failing_gates"] == 1
    assert benchmark_review["summary"]["source_acquisition_tasks"] == 1
    assert benchmark_review["source_acquisition_tasks"][0]["task_type"] == "full_text_upload"
    assert benchmark_review["source_acquisition_tasks"][0]["status"] == "source_candidate_accepted_needs_override"
    assert benchmark_review["source_acquisition_tasks"][0]["uploaded_sources"][0]["filename"] == "trial-b-supplement.pdf"
    assert benchmark_review["source_acquisition_tasks"][0]["uploaded_sources"][0]["parsed_path"] == "benchmark/source_parsed/abc.json"
    assert benchmark_review["source_acquisition_tasks"][0]["uploaded_sources"][0]["text_chars"] == 240
    assert benchmark_review["source_acquisition_tasks"][0]["uploaded_sources"][0]["quote_candidates"][0]["source_page"] == 3
    assert (
        benchmark_review["source_acquisition_tasks"][0]["uploaded_sources"][0]["quote_candidates"][0]["review_decision"]["decision"]
        == "accepted"
    )
    assert "MetaAgent Benchmark Review" in benchmark_html
    assert "mini_benchmark" in benchmark_html
    assert "Trial B" in benchmark_html
    assert "Source Acquisition Tasks" in benchmark_html
    assert "full_text_upload" in benchmark_html
    assert "240 chars" in benchmark_html
    assert "1 quote candidate" in benchmark_html
    assert "matched 11, 20, 40" in benchmark_html
    assert "accepted by tester" in benchmark_html


def test_artifact_package_file_discovery_skips_empty_files_and_includes_dynamic_assets(tmp_path: Path) -> None:
    project = Project("artifact package discovery", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# Manuscript", subdir="manuscript")
    project.save_text("references.bib", "@article{x}")
    (project.base_dir / "pipeline_warnings.json").write_text("", encoding="utf-8")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)
    project.save_json("abc.json", {"full_text": "Parsed supplement"}, subdir="benchmark/source_parsed")

    arcnames = {arcname for _, arcname in iter_existing_package_files(project)}

    assert "references.bib" in arcnames
    assert "manuscript/draft.md" in arcnames
    assert "figures/forest_plot.png" in arcnames
    assert "benchmark/source_parsed/abc.json" in arcnames
    assert "pipeline_warnings.json" not in arcnames


def test_artifact_package_exports_submission_docx(tmp_path: Path) -> None:
    project = Project("docx package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Systemic corticosteroids and mortality",
            "",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "",
            "## Methods",
            "PubMed-style Boolean strategy was used.",
            "",
            "| Trial | Events | Source |",
            "|---|---:|---|",
            "| RECOVERY | 95/324 vs 283/683 | Figure 2 |",
            "",
            "## Figures",
            "![Figure 1. Forest plot](../figures/forest_plot.png)",
            "",
            "## References",
            "[1] WHO REACT Working Group. Association Between Administration of Systemic Corticosteroids and Mortality Among Critically Ill Patients With COVID-19. *JAMA*. 2020.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "evidence_readiness": {"status": "ready"}},
        subdir="manuscript",
    )
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    assert (project.base_dir / "manuscript" / "draft.docx").exists()
    assert (project.base_dir / "manuscript" / "draft.pdf").exists()
    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        docx_bytes = zf.read("manuscript/draft.docx")
        pdf_bytes = zf.read("manuscript/draft.pdf")
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "manuscript/draft.docx" in names
    assert "manuscript/draft.pdf" in names
    assert "manuscript/draft.docx" in manifest["entries"]
    assert "manuscript/draft.pdf" in manifest["entries"]
    assert manifest["manuscript"]["included"] is True
    assert manifest["manuscript"]["markdown"] is True
    assert manifest["manuscript"]["docx"] is True
    assert manifest["manuscript"]["pdf"] is True
    assert manifest["manuscript"]["word_count"] >= 30
    assert pdf_bytes.startswith(b"%PDF")
    assert len(pdf_bytes) > 1000
    document = Document(BytesIO(docx_bytes))
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    assert "Systemic corticosteroids and mortality" in paragraph_text
    assert "Abstract" in paragraph_text
    assert "Primary outcome OR 0.66" in paragraph_text
    assert "Figure 1. Forest plot" in paragraph_text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "RECOVERY"
    with zipfile.ZipFile(BytesIO(docx_bytes)) as docx_zip:
        assert any(name.startswith("word/media/") for name in docx_zip.namelist())


def test_artifact_package_surfaces_limited_text_source_coverage(tmp_path: Path) -> None:
    project = Project("source coverage package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "# Manuscript\n\n## Abstract\n\nShort abstract.\n",
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "evidence_gap", "evidence_readiness": {"status": "ready"}},
        subdir="manuscript",
    )
    project.save_json(
        "text_source_warnings.json",
        [
            {
                "pmid": "31523904",
                "doi": "10.1002/ejhf.1596",
                "title": "EMPEROR-Preserved trial design",
                "trial_registration": "NCT03057951",
                "text_availability": "abstract_only",
                "warning": "Only structured abstract text was retrievable automatically; extraction requires manual verification.",
            },
            {
                "pmid": "",
                "doi": "",
                "title": "Registry-only Trial",
                "trial_registration": "NCT00000001",
                "text_availability": "metadata_only",
                "warning": "Only registry metadata is available; outcome extraction requires user-uploaded full text or verified source data.",
            },
        ],
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        audit = json.loads(zf.read("review/text_source_coverage_audit.json"))
        audit_html = zf.read("review/text_source_coverage_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "text_source_warnings.json" in names
    assert "review/text_source_coverage_audit.json" in names
    assert "review/text_source_coverage_audit.html" in names
    assert audit["status"] == "ready"
    assert audit["summary"]["limited_source_records"] == 2
    assert audit["summary"]["abstract_only_records"] == 1
    assert audit["summary"]["metadata_only_records"] == 1
    assert audit["summary"]["action_required_limited_records"] == 0
    assert audit["summary"]["screening_only_limited_records"] == 2
    records_by_title = {record["title"]: record for record in audit["records"]}
    assert records_by_title["EMPEROR-Preserved trial design"]["source_level"] == "abstract_only"
    assert records_by_title["Registry-only Trial"]["source_level"] == "metadata_only"
    assert records_by_title["EMPEROR-Preserved trial design"]["impact_scope"] == "screening_only"
    assert records_by_title["Registry-only Trial"]["impact_scope"] == "screening_only"
    assert records_by_title["EMPEROR-Preserved trial design"]["requires_review"] is False
    assert audit["issues"] == []
    assert "EMPEROR-Preserved trial design" in audit_html
    assert "NCT00000001" in audit_html
    source_gate = next(gate for gate in readiness["gates"] if gate["id"] == "source_coverage")
    assert source_gate["status"] == "pass"
    assert "limited source record(s)=2" in source_gate["detail"]
    assert "action-required limited source record(s)=0" in source_gate["detail"]
    assert manifest["review"]["text_source_coverage_included"] is True
    assert manifest["review"]["text_source_coverage_limited_records"] == 2
    assert manifest["review"]["text_source_coverage_action_required_records"] == 0
    assert manifest["review"]["text_source_coverage_html_review"] is True


def test_source_coverage_warns_when_limited_source_feeds_primary_analysis(tmp_path: Path) -> None:
    project = Project("source coverage primary", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "# Manuscript\n\n## Abstract\n\nShort abstract.\n",
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "evidence_gap", "evidence_readiness": {"status": "ready"}},
        subdir="manuscript",
    )
    project.save_json(
        "text_source_warnings.json",
        [
            {
                "pmid": "31523904",
                "doi": "10.1002/ejhf.1596",
                "title": "EMPEROR-Preserved trial design",
                "text_availability": "abstract_only",
                "warning": "Only structured abstract text was retrievable automatically.",
            }
        ],
    )
    project.save_json(
        "all_extractions.json",
        [
            {
                "characteristics": {
                    "study_id": "31523904",
                    "pmid": "31523904",
                    "doi": "10.1002/ejhf.1596",
                    "title": "EMPEROR-Preserved trial design",
                },
                "outcomes": [
                    {
                        "outcome_name": "Composite endpoint",
                        "outcome_type": "dichotomous",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                    }
                ],
            }
        ],
        subdir="extraction",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "Composite endpoint",
                "studies": [{"study_id": "31523904", "study_label": "Anker 2019"}],
            }
        },
        subdir="analysis",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/text_source_coverage_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    record = audit["records"][0]
    assert record["impact_scope"] == "primary_analysis"
    assert record["requires_review"] is True
    assert "primary_meta_analysis" in record["downstream_uses"]
    assert "extracted_outcome" in record["downstream_uses"]
    assert audit["summary"]["action_required_limited_records"] == 1
    assert audit["summary"]["screening_only_limited_records"] == 0
    assert audit["issues"][0]["code"] == "limited_text_source_used_downstream"
    source_gate = next(gate for gate in readiness["gates"] if gate["id"] == "source_coverage")
    assert source_gate["status"] == "warn"
    assert "action-required limited source record(s)=1" in source_gate["detail"]


def test_text_source_coverage_review_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese source coverage package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n评价治疗是否改善主要结局。",
            "## 方法\n全文来源不足的研究必须人工复核。",
            "## 结果\n主要分析纳入1项研究。",
            "## 讨论\n有限文本来源会影响证据可信度。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "output_language": "zh", "evidence_readiness": {"status": "ready"}},
        subdir="manuscript",
    )
    project.save_json(
        "text_source_warnings.json",
        [
            {
                "pmid": "31523904",
                "doi": "10.1002/ejhf.1596",
                "title": "EMPEROR-Preserved trial design",
                "text_availability": "abstract_only",
                "warning": "Only structured abstract text was retrievable automatically.",
            }
        ],
    )
    project.save_json(
        "all_extractions.json",
        [
            {
                "characteristics": {
                    "study_id": "31523904",
                    "pmid": "31523904",
                    "doi": "10.1002/ejhf.1596",
                    "title": "EMPEROR-Preserved trial design",
                },
                "outcomes": [
                    {
                        "outcome_name": "Composite endpoint",
                        "outcome_type": "dichotomous",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                    }
                ],
            }
        ],
        subdir="extraction",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "Composite endpoint",
                "studies": [{"study_id": "31523904", "study_label": "Anker 2019"}],
            }
        },
        subdir="analysis",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/text_source_coverage_audit.json"))
        audit_html = zf.read("review/text_source_coverage_audit.html").decode("utf-8")

    assert audit["language"] == "zh"
    assert '<html lang="zh">' in audit_html
    assert "MetaAgent 文本来源覆盖复核" in audit_html
    assert "复核仅有摘要、元数据、注册记录或未知文本来源的记录" in audit_html
    assert "记录总数" in audit_html
    assert "全文" in audit_html
    assert "仅摘要" in audit_html
    assert "需人工复核" in audit_html
    assert "来源记录" in audit_html
    assert "复核问题" in audit_html
    assert "来源级别" in audit_html
    assert "影响范围" in audit_html
    assert "主要分析" in audit_html
    assert "有限文本来源进入下游分析" in audit_html
    assert "是" in audit_html
    assert "MetaAgent Text Source Coverage" not in audit_html
    assert "Review records that rely on abstract-only" not in audit_html


def test_artifact_package_includes_manuscript_polish_audit_summary(tmp_path: Path) -> None:
    project = Project("artifact polish audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66. "
            "The pooled RR was 0.74. "
            "The pooled MD was -1.0. "
            "The pooled HR was 0.81.\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "rewrite_scope": "targeted",
            "accepted_chunks": 2,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "attempted_chunks": 3,
            "skipped_chunks": 2,
            "total_rewrite_chunks": 5,
            "targeted_chunks": 3,
            "non_target_chunks": 2,
            "rewrite_retries": 2,
            "retry_recovered_chunks": 1,
            "polish_budget_exhausted": True,
            "skipped_chunk_details": [
                {
                    "heading": "Discussion",
                    "chunk_index": 3,
                    "chunk_count": 5,
                    "reason": "polish_budget_exhausted",
                    "original_text": "It is important to note that the skipped RR paragraph stayed unchanged [1].",
                    "kept_text": "The skipped RR paragraph stayed unchanged [1].",
                    "deterministic_cleanup_applied": True,
                    "review_action": "rerun_with_higher_polish_budget",
                }
            ],
            "accepted_sections": 1,
            "rejected_sections": 1,
            "accepted_edit_count": 1,
            "accepted_edits": [
                {
                    "heading": "Discussion",
                    "chunk_index": 0,
                    "chunk_count": 1,
                    "original_text": "It is important to note that the pooled OR was 0.66 [1].",
                    "candidate_text": "The pooled OR was 0.66 [1].",
                    "diff": "--- original\n+++ polished\n@@ -1 +1 @@\n-It is important to note that the pooled OR was 0.66 [1].\n+The pooled OR was 0.66 [1].",
                    "review_action": "accepted_fact_preserving_polish",
                }
            ],
            "style_policy": {
                "name": "MetaAgent conservative scholarly polish",
                "detector_evasion": False,
                "detector_optimization": "disabled",
                "external_proofreader_role": "review_only",
            },
            "proofreading": {
                "enabled": True,
                "status": "ok",
                "provider": "languagetool",
                "language_code": "en-US",
                "issue_count": 1,
                "issues": [{"rule_id": "STYLE_PASSIVE", "message": "Consider direct phrasing."}],
            },
            "before": {
                "ai_style_signal": {
                    "score": 3,
                    "issues": [{"code": "template_phrase_hits"}, {"code": "repeated_sentence_starts"}],
                }
            },
            "after": {
                "ai_style_signal": {
                    "score": 1,
                    "issues": [{"code": "low_sentence_length_variation"}],
                }
            },
            "issues": [
                {
                    "code": "numeric_tokens_changed",
                    "heading": "Discussion",
                    "message": "Polish rewrite changed numeric values.",
                    "original_text": "It is important to note that the pooled OR was 0.66 [1].",
                    "candidate_text": "The pooled OR was 0.66 [1].",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "manuscript/manuscript_polish_audit.json" in names
    assert "review/manuscript_polish_audit.json" in names
    assert "review/manuscript_polish_audit.html" in names
    assert review["summary"]["before_ai_style_score"] == 3
    assert review["summary"]["after_ai_style_score"] == 1
    assert review["summary"]["fact_guard_issues"] == 1
    assert review["summary"]["rewrite_scope"] == "targeted"
    assert review["summary"]["attempted_chunks"] == 3
    assert review["summary"]["skipped_chunks"] == 2
    assert review["summary"]["total_rewrite_chunks"] == 5
    assert review["summary"]["targeted_chunks"] == 3
    assert review["summary"]["non_target_chunks"] == 2
    assert review["summary"]["rewrite_retries"] == 2
    assert review["summary"]["retry_recovered_chunks"] == 1
    assert review["summary"]["polish_budget_exhausted"] is True
    assert review["summary"]["skipped_chunk_detail_count"] == 1
    assert review["summary"]["resolved_ai_style_issues"] == 2
    assert review["summary"]["proofreading_issues"] == 1
    assert review["summary"]["accepted_edit_count"] == 1
    assert [item["code"] for item in review["resolved_style_issues"]] == [
        "template_phrase_hits",
        "repeated_sentence_starts",
    ]
    assert review["resolved_style_issues"][0]["status"] == "resolved_after_polish"
    assert review["skipped_chunk_details"][0]["heading"] == "Discussion"
    assert review["skipped_chunk_details"][0]["chunk_index"] == 3
    assert "RR paragraph" in review["skipped_chunk_details"][0]["kept_text"]
    assert review["accepted_edits"][0]["heading"] == "Discussion"
    assert "It is important to note" in review["accepted_edits"][0]["original_text"]
    assert review["style_policy"]["detector_evasion"] is False
    assert review["proofreading"]["provider"] == "languagetool"
    queue = review["review_queue"]
    assert queue["status"] == "human_review_required"
    assert queue["accepted_auto_edits"] == 1
    assert queue["rewrite_retries"] == 2
    assert queue["retry_recovered_chunks"] == 1
    assert queue["rejected_candidates"] == 1
    assert queue["remaining_style_issues"] == 1
    assert queue["proofreading_issues"] == 1
    assert queue["manual_review_items"] == 3
    assert queue["can_auto_apply_rejected_edits"] is False
    assert any("human confirms" in action for action in queue["next_actions"])
    rejected = review["rejected_edits"][0]
    assert rejected["candidate_id"].startswith("rejected:0:numeric_tokens_changed")
    assert rejected["code"] == "numeric_tokens_changed"
    assert rejected["heading"] == "Discussion"
    assert rejected["can_auto_apply"] is False
    assert rejected["manual_accept_allowed"] is True
    assert "human confirms" in rejected["manual_accept_condition"]
    assert "It is important to note" not in rejected["candidate_text"]
    assert "-It is important to note" in rejected["diff"]
    remaining = review["remaining_style_issues"][0]
    assert remaining["code"] == "low_sentence_length_variation"
    assert "message" in remaining
    assert "suggested_action" in remaining
    assert "evidence_excerpt" in remaining
    assert "pooled OR" in remaining["evidence_excerpt"]
    assert "manual" in remaining["suggested_action"].lower()
    assert "MetaAgent Manuscript Polish Audit" in html
    assert "LanguageTool" in html
    assert "Accepted Polish Edits" in html
    assert "Review Queue" in html
    assert "Skipped Polish Chunks" in html
    assert "Rewrite retries" in html
    assert "Recovered chunks" in html
    assert "human_review_required" in html
    assert "RR paragraph" in html
    assert "accepted_fact_preserving_polish" in html
    assert "It is important to note" in html
    assert "detector evasion disabled" in html
    assert "numeric_tokens_changed" in html
    assert "Rejected Polish Candidates" in html
    assert "human confirms" in html
    assert "low_sentence_length_variation" in html
    assert "Resolved Style Signals" in html
    assert "resolved_after_polish" in html
    assert "Sentence lengths are too uniform" in html
    assert "pooled OR" in html
    assert "manual_review_required" in html
    assert manifest["review"]["manuscript_polish_included"] is True
    assert manifest["review"]["manuscript_polish_html_review"] is True
    assert manifest["review"]["manuscript_polish_enabled"] is True
    assert manifest["review"]["manuscript_polish_before_ai_style_score"] == 3
    assert manifest["review"]["manuscript_polish_after_ai_style_score"] == 1
    assert manifest["review"]["manuscript_polish_ai_style_delta"] == -2
    assert manifest["review"]["manuscript_polish_remaining_ai_style_issues"] == 1
    assert manifest["review"]["manuscript_polish_rewrite_scope"] == "targeted"
    assert manifest["review"]["manuscript_polish_accepted_edit_count"] == 1
    assert manifest["review"]["manuscript_polish_rejected_chunks"] == 1
    assert manifest["review"]["manuscript_polish_attempted_chunks"] == 3
    assert manifest["review"]["manuscript_polish_skipped_chunks"] == 2
    assert manifest["review"]["manuscript_polish_skipped_chunk_detail_count"] == 1
    assert manifest["review"]["manuscript_polish_total_rewrite_chunks"] == 5
    assert manifest["review"]["manuscript_polish_targeted_chunks"] == 3
    assert manifest["review"]["manuscript_polish_non_target_chunks"] == 2
    assert manifest["review"]["manuscript_polish_rewrite_retries"] == 2
    assert manifest["review"]["manuscript_polish_retry_recovered_chunks"] == 1
    assert manifest["review"]["manuscript_polish_budget_exhausted"] is True
    assert manifest["review"]["manuscript_polish_fact_guard_issues"] == 1
    assert manifest["review"]["manuscript_polish_resolved_ai_style_issues"] == 2
    assert manifest["review"]["manuscript_polish_review_queue_status"] == "human_review_required"
    assert manifest["review"]["manuscript_polish_manual_review_items"] == 3
    assert manifest["review"]["manuscript_polish_rejected_candidates"] == 1
    assert manifest["review"]["manuscript_polish_proofreading_issues"] == 1
    assert manifest["review"]["manuscript_polish_detector_evasion"] is False
    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")
    assert "resolved_style_issues=2" in polish_gate["detail"]
    assert "rewrite_retries=2" in polish_gate["detail"]
    assert "retry_recovered_chunks=1" in polish_gate["detail"]


def test_artifact_package_flags_failed_manuscript_proofreader(tmp_path: Path) -> None:
    project = Project("artifact failed proofreader", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 0,
            "unchanged_chunks": 0,
            "proofreading": {
                "enabled": True,
                "status": "failed",
                "provider": "languagetool",
                "language_code": "en-US",
                "issue_count": 0,
                "issues": [],
                "error": "timeout",
            },
            "before": {"ai_style_signal": {"score": 1, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"ai_style_signal": {"score": 0, "issues": []}},
            "issues": [],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    queue = review["review_queue"]
    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")

    assert review["passed"] is False
    assert review["summary"]["proofreading_failed"] is True
    assert review["summary"]["proofreading_error"] == "timeout"
    assert review["proofreading"]["status"] == "failed"
    assert review["proofreading"]["error"] == "timeout"
    assert queue["status"] == "human_review_required"
    assert queue["proofreading_failed"] is True
    assert queue["proofreading_error"] == "timeout"
    assert queue["manual_review_items"] == 1
    assert any("rerun" in action.lower() for action in queue["next_actions"])
    assert "Proofreader failed" in html
    assert "timeout" in html
    assert polish_gate["status"] == "warn"
    assert "proofreading_failed=True" in polish_gate["detail"]
    assert any("rerun" in action.lower() for action in polish_gate["next_actions"])
    assert manifest["review"]["manuscript_polish_proofreading_failed"] is True
    assert manifest["review"]["manuscript_polish_proofreading_error"] == "timeout"
    assert manifest["review"]["manuscript_polish_manual_review_items"] == 1


def test_manuscript_polish_review_passes_safe_final_text_with_rejected_candidates(tmp_path: Path) -> None:
    project = Project("artifact safe rejected polish", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 1,
            "before": {"ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"ai_style_signal": {"score": 0, "issues": []}},
            "proofreading": {"enabled": False, "status": "disabled", "provider": "none", "issue_count": 0},
            "issues": [
                {
                    "code": "numeric_tokens_changed",
                    "heading": "Discussion",
                    "message": "Polish rewrite changed numeric values.",
                    "original_text": "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].",
                    "candidate_text": "The pooled OR was 0.68 (95% CI 0.53 to 0.82) [1].",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    review = _build_manuscript_polish_audit_review(project)

    assert review["passed"] is True
    assert review["summary"]["fact_guard_issues"] == 1
    assert review["summary"]["remaining_ai_style_issues"] == 0
    assert review["summary"]["proofreading_issues"] == 0
    assert review["review_queue"]["status"] == "polish_guard_discarded_candidates_no_review_required"
    assert review["review_queue"]["rejected_candidates"] == 1
    assert review["review_queue"]["manual_review_items"] == 0
    assert review["rejected_edits"][0]["can_auto_apply"] is False


def test_submission_readiness_passes_safe_final_text_with_discarded_polish_candidates(tmp_path: Path) -> None:
    project = Project("artifact safe discarded polish readiness", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 1,
            "before": {"ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"ai_style_signal": {"score": 0, "issues": []}},
            "proofreading": {"enabled": True, "status": "ok", "provider": "languagetool", "issue_count": 0, "issues": []},
            "issues": [
                {
                    "code": "numeric_tokens_changed",
                    "heading": "Discussion",
                    "message": "Polish rewrite changed numeric values.",
                    "original_text": "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].",
                    "candidate_text": "The pooled OR was 0.68 (95% CI 0.53 to 0.82) [1].",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")

    assert review["review_queue"]["status"] == "polish_guard_discarded_candidates_no_review_required"
    assert review["review_queue"]["manual_review_items"] == 0
    assert polish_gate["status"] == "pass"
    assert not polish_gate.get("next_actions")


def test_submission_readiness_does_not_warn_when_polish_is_disabled_by_policy(tmp_path: Path) -> None:
    project = Project("artifact disabled polish readiness", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1].\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": False,
            "language": "en",
            "rewrite_scope": "targeted",
            "accepted_chunks": 0,
            "rejected_chunks": 0,
            "after": {"ai_style_signal": {"score": 1, "issues": [{"code": "template_phrase_hits"}]}},
            "proofreading": {"enabled": False, "status": "disabled", "provider": "none", "issue_count": 0},
            "review_queue": {
                "status": "human_review_required",
                "manual_review_items": 1,
                "next_actions": ["Review remaining style signals manually."],
            },
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))
        polish_review = json.loads(zf.read("review/manuscript_polish_audit.json"))

    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")

    assert polish_gate["status"] == "pass"
    assert "enabled=False" in polish_gate["detail"]
    assert not polish_gate.get("next_actions")
    assert polish_review["review_queue"]["status"] == "no_polish_review_needed"
    assert polish_review["review_queue"]["manual_review_items"] == 0
    assert polish_review["review_queue"]["next_actions"] == []
    assert manifest["review"]["manuscript_polish_review_queue_status"] == "no_polish_review_needed"
    assert manifest["review"]["manuscript_polish_manual_review_items"] == 0
    assert manifest["review"]["manuscript_polish_next_actions"] == []


def test_manuscript_polish_review_rescores_current_draft_when_after_signal_is_stale(tmp_path: Path) -> None:
    project = Project("artifact stale polish signal", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Discussion\n\n"
            "The pooled OR was 0.66 (95% CI 0.53 to 0.82) [1]. "
            "The interpretation remains source linked.\n\n"
            "## References\n\n"
            "[1] Example reference.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 0,
            "before": {"ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {
                "ai_style_signal": {
                    "score": 1,
                    "issues": [{"code": "low_sentence_length_variation"}],
                }
            },
            "proofreading": {"enabled": False, "status": "disabled", "provider": "none", "issue_count": 0},
            "issues": [],
        },
        subdir="manuscript",
    )

    review = _build_manuscript_polish_audit_review(project)

    assert review["passed"] is True
    assert review["summary"]["polish_audit_stale"] is True
    assert review["summary"]["stored_after_ai_style_score"] == 1
    assert review["summary"]["after_ai_style_score"] == 0
    assert review["summary"]["remaining_ai_style_issues"] == 0
    assert review["current_draft_ai_style_signal"]["score"] == 0
    assert review["review_queue"]["remaining_style_issues"] == 0
    assert review["review_queue"]["status"] == "no_polish_review_needed"

    package_path = create_artifact_package(project)
    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        packaged_review = json.loads(zf.read("review/manuscript_polish_audit.json"))

    assert packaged_review["summary"]["after_ai_style_score"] == 0
    assert packaged_review["summary"]["polish_audit_stale"] is True
    assert manifest["review"]["manuscript_polish_after_ai_style_score"] == 0
    assert manifest["review"]["manuscript_polish_stored_after_ai_style_score"] == 1
    assert manifest["review"]["manuscript_polish_current_draft_ai_style_score"] == 0
    assert manifest["review"]["manuscript_polish_remaining_ai_style_issues"] == 0
    assert manifest["review"]["manuscript_polish_audit_stale"] is True


def test_artifact_package_localizes_mixed_chinese_polish_style_issues(tmp_path: Path) -> None:
    project = Project("artifact mixed zh polish audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "# 中文稿\n\n## 摘要\n\n本研究显示疗效降低。本研究显示疗效稳定。本研究显示疗效可信。本研究显示疗效相关。\n",
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "mixed",
            "after": {
                "ai_style_signal": {
                    "score": 1,
                    "issues": [{"code": "repeated_sentence_starts", "count": 1, "openings": {"hfmref": 2}}],
                }
            },
            "issues": [
                {
                    "code": "directional_terms_changed",
                    "heading": "讨论",
                    "message": "Polish rewrite changed directional conclusion terms.",
                    "original_text": "结果未显示安全性结局增加。",
                    "candidate_text": "结果显示安全性结局增加。",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    issue = review["remaining_style_issues"][0]
    assert review["language"] == "mixed"
    assert review["review_language"] == "zh"
    assert "多个句子使用相同开头" in issue["message"]
    assert "人工改写" in issue["suggested_action"]
    assert "本研究显示疗效" in issue["evidence_excerpt"]
    rejected = review["rejected_edits"][0]
    assert "结论方向" in rejected["message"]
    assert "结论方向" in rejected["blocking_reason"]
    assert "剩余风格信号" in html
    assert "多个句子使用相同开头" in html
    assert "结论方向" in html
    assert "审查语言" in html
    assert ">中文<" in html or "审查语言: 中文" in html
    assert "语言: mixed" not in html
    assert "方向性术语变更" in html
    assert "重复句首" in html
    assert "directional_terms_changed" not in html
    assert "repeated_sentence_starts" not in html
    assert "证据摘录" in html
    assert "处理建议" in html
    assert "evidence_excerpt" not in html
    assert "suggested_action" not in html


def test_artifact_package_does_not_count_polish_budget_as_fact_guard_issue(tmp_path: Path) -> None:
    project = Project("artifact polish budget only", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "# Title\n\n## Discussion\n\nThe manuscript text was kept unchanged [1].\n\n## References\n\n[1] Example reference.\n",
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 0,
            "unchanged_chunks": 0,
            "attempted_chunks": 1,
            "skipped_chunks": 3,
            "total_rewrite_chunks": 4,
            "polish_budget_exhausted": True,
            "accepted_sections": 1,
            "rejected_sections": 0,
            "accepted_edit_count": 1,
            "accepted_edits": [],
            "style_policy": {"detector_evasion": False},
            "proofreading": {"enabled": False, "issue_count": 0, "issues": []},
            "before": {"ai_style_signal": {"score": 0, "issues": []}},
            "after": {"ai_style_signal": {"score": 0, "issues": []}},
            "issues": [
                {
                    "code": "polish_budget_exhausted",
                    "review_action": "rerun_with_higher_polish_budget",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert review["summary"]["polish_budget_exhausted"] is True
    assert review["summary"]["fact_guard_issues"] == 0
    assert manifest["review"]["manuscript_polish_budget_exhausted"] is True
    assert manifest["review"]["manuscript_polish_fact_guard_issues"] == 0


def test_artifact_package_includes_manuscript_citation_fix_log_summary(tmp_path: Path) -> None:
    project = Project("artifact citation fix log", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# Title\n\n"
            "## Introduction\n\n"
            "Background statement [1] [2].\n\n"
            "## References\n\n"
            "[1] Trial report.\n\n"
            "[2] Guideline report.\n"
        ),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_citation_fixes.json",
        {
            "schema_version": 1,
            "current_revision": 3,
            "entries": [
                {
                    "revision": 1,
                    "action": "citation_patch",
                    "issue_id": "citation_audit:0",
                    "citation": "[2]",
                    "section": "Introduction",
                    "user_id": "reviewer",
                    "quality_delta": {
                        "quality_status_before": "blocked",
                        "quality_status_after": "ready",
                        "reference_entries_added": 1,
                        "resolved_issue_ids": [
                            "citation_audit:0",
                            "primary_result:0:pooled_effect",
                            "claim_support:0:primary_effect",
                        ],
                        "primary_result_mismatched_fields_resolved": 1,
                        "primary_result_failed_issues_resolved": 1,
                        "resolved_primary_result_issue_ids": ["primary_result:0:pooled_effect"],
                        "claim_support_unsupported_claims_resolved": 1,
                        "claim_support_failed_issues_resolved": 1,
                        "resolved_claim_support_issue_ids": ["claim_support:0:primary_effect"],
                    },
                },
                {
                    "revision": 2,
                    "action": "add_reference",
                    "issue_id": "citation_audit:1",
                    "candidate_id": "evimed:guide:hf",
                    "candidate_source": {"source_type": "guideline", "url": "https://example.test/guideline"},
                    "trust": {"status": "needs_review", "requires_human_review": True},
                    "citation": "[2]",
                    "section": "Introduction",
                    "user_id": "reviewer",
                },
                {
                    "revision": 3,
                    "action": "reuse_reference_citation",
                    "issue_id": "citation_audit:2",
                    "candidate_id": "evimed:guide:hf",
                    "candidate_source": {"source_type": "guideline", "url": "https://example.test/guideline"},
                    "trust": {"status": "needs_review", "requires_human_review": True},
                    "citation": "[2]",
                    "reference_added": False,
                    "section": "Discussion",
                    "user_id": "reviewer",
                },
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        log = json.loads(zf.read("manuscript/manuscript_citation_fixes.json"))
        review = json.loads(zf.read("review/manuscript_citation_fixes.json"))
        html = zf.read("review/manuscript_citation_fixes.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "manuscript/manuscript_citation_fixes.json" in names
    assert "review/manuscript_citation_fixes.json" in names
    assert "review/manuscript_citation_fixes.html" in names
    assert log["current_revision"] == 3
    assert review["summary"]["current_revision"] == 3
    assert review["summary"]["reference_add_actions"] == 1
    assert review["summary"]["reference_reuse_actions"] == 1
    assert review["summary"]["reference_add_human_review_required"] == 1
    assert review["summary"]["reference_fix_human_review_required"] == 2
    assert review["summary"]["quality_delta_entries"] == 1
    assert review["summary"]["quality_resolved_issues"] == 3
    assert review["summary"]["quality_reference_entries_added"] == 1
    assert review["summary"]["quality_primary_result_mismatches_resolved"] == 1
    assert review["summary"]["quality_claim_support_unsupported_claims_resolved"] == 1
    assert review["summary"]["quality_primary_result_issues_resolved"] == 1
    assert review["summary"]["quality_claim_support_issues_resolved"] == 1
    assert "MetaAgent Manuscript Citation Fixes" in html
    assert "Quality impact" in html
    assert "blocked to ready" in html
    assert "primary result mismatches -1" in html
    assert "unsupported claims -1" in html
    assert "add_reference" in html
    assert "reuse_reference_citation" in html
    assert "needs_review" in html
    assert "https://example.test/guideline" in html
    assert manifest["review"]["manuscript_citation_fixes_included"] is True
    assert manifest["review"]["manuscript_citation_fixes_html_review"] is True
    assert manifest["review"]["manuscript_citation_fixes_current_revision"] == 3
    assert manifest["review"]["manuscript_citation_fix_entries"] == 3
    assert manifest["review"]["manuscript_citation_patch_actions"] == 1
    assert manifest["review"]["manuscript_reference_add_actions"] == 1
    assert manifest["review"]["manuscript_reference_reuse_actions"] == 1
    assert manifest["review"]["manuscript_reference_add_human_review_required"] == 1
    assert manifest["review"]["manuscript_reference_fix_human_review_required"] == 2
    assert manifest["review"]["manuscript_citation_fix_quality_delta_entries"] == 1
    assert manifest["review"]["manuscript_citation_fix_quality_resolved_issues"] == 3
    assert manifest["review"]["manuscript_citation_fix_quality_reference_entries_added"] == 1
    assert manifest["review"]["manuscript_citation_fix_quality_primary_result_mismatches_resolved"] == 1
    assert manifest["review"]["manuscript_citation_fix_quality_claim_support_unsupported_claims_resolved"] == 1
    assert manifest["review"]["manuscript_citation_fix_quality_primary_result_issues_resolved"] == 1
    assert manifest["review"]["manuscript_citation_fix_quality_claim_support_issues_resolved"] == 1


def test_manuscript_citation_fix_review_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("artifact citation fix zh log", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        (
            "# 中文Meta分析稿件\n\n"
            "## 引言\n\n"
            "既往证据提示需要补充指南引用[1]。\n\n"
            "## 参考文献\n\n"
            "[1] 指南报告。\n"
        ),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_facts.json", {"output_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_citation_fixes.json",
        {
            "schema_version": 1,
            "current_revision": 2,
            "entries": [
                {
                    "revision": 1,
                    "action": "citation_patch",
                    "issue_id": "citation_audit:0",
                    "citation": "[1]",
                    "section": "引言",
                    "user_id": "reviewer",
                },
                {
                    "revision": 2,
                    "action": "add_reference",
                    "issue_id": "citation_audit:1",
                    "candidate_id": "evimed:guide:hf",
                    "candidate_source": {"source_type": "guideline", "url": "https://example.test/guideline"},
                    "trust": {"status": "needs_review", "requires_human_review": True},
                    "citation": "[1]",
                    "section": "引言",
                    "user_id": "reviewer",
                },
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_citation_fixes.json"))
        html = zf.read("review/manuscript_citation_fixes.html").decode("utf-8")

    assert review["language"] == "zh"
    assert '<html lang="zh">' in html
    assert "MetaAgent 稿件引用修复记录" in html
    assert "记录系统补充文内引用和人工复核参考文献候选的过程" in html
    assert "当前修订" in html
    assert "修复记录" in html
    assert "引用补丁" in html
    assert "新增参考文献" in html
    assert "引用复用" in html
    assert "需要人工复核" in html
    assert "修订记录" in html
    assert "修订" in html
    assert "动作" in html
    assert "章节" in html
    assert "引用" in html
    assert "问题" in html
    assert "候选" in html
    assert "信任状态" in html
    assert "来源" in html
    assert "MetaAgent Manuscript Citation Fixes" not in html
    assert "Audit trail for citation patches" not in html


def test_manuscript_citation_fix_review_localizes_chinese_empty_state(tmp_path: Path) -> None:
    project = Project("artifact citation fix zh empty", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# 中文Meta分析稿件\n\n## 引言\n\n已有引用[1]。\n", subdir="manuscript")
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_facts.json", {"output_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_citation_fixes.json",
        {"schema_version": 1, "current_revision": 0, "entries": []},
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        html = zf.read("review/manuscript_citation_fixes.html").decode("utf-8")

    assert "未记录稿件引用修复条目" in html
    assert "No manuscript citation fix entries were recorded" not in html


def test_artifact_package_includes_submission_readiness_review(tmp_path: Path) -> None:
    project = Project("submission readiness package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Submission-ready manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "",
            "## Methods",
            "PubMed Boolean strategy: corticosteroids AND COVID-19 AND mortality [1].",
            "",
            "## Results",
            "One source-verified trial row is shown for this fixture [1].",
            "",
            *_clinical_discussion_fixture(effect="OR 0.66", endpoint="mortality"),
            "## Figures",
            "![Figure 1. Forest plot](../figures/forest_plot.png)",
            "",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "needs_review",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_validation.json",
        {"passed": True, "issues": [], "facts_summary": {"main_word_count": 6123, "report_type": "meta"}},
        subdir="manuscript",
    )
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "all_extractions.json",
        [
            ExtractedStudy(
                characteristics=StudyCharacteristics(study_id="S1", pmid="S1", title="Source Trial"),
                outcomes=[
                    OutcomeData(
                        outcome_name="mortality",
                        outcome_type="dichotomous",
                        events_intervention=1,
                        total_intervention=10,
                        events_control=2,
                        total_control=10,
                        source_quote="Mortality was 1/10 vs 2/10.",
                        source_quote_match="Mortality was 1/10 vs 2/10.",
                        source_quote_verified=True,
                        extraction_confidence="high",
                    )
                ],
            )
        ],
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "Before source. Mortality was 1/10 vs 2/10. After source.", "page_map": []}},
        subdir="papers",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "summary": {"outcomes": 1},
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 1/10 vs 2/10.",
                    "source_quote_match": "Mortality was 1/10 vs 2/10.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "(corticosteroids) AND (COVID-19) AND mortality")
    project.save_text("search_strategy_report.txt", "PubMed Boolean strategy: corticosteroids AND COVID-19 AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        readiness_html = zf.read("review/submission_readiness_review.html").decode("utf-8")

    assert "search_query.txt" in names
    assert "search_strategy_report.txt" in names
    assert "review/submission_readiness_review.json" in names
    assert "review/submission_readiness_review.html" in names
    assert manifest["submission"]["included"] is True
    assert manifest["submission"]["passed"] is True
    assert manifest["submission"]["status"] == "ready"
    assert manifest["submission"]["failed_gates"] == 0
    assert readiness["passed"] is True
    assert readiness["status"] == "ready"
    assert {gate["id"] for gate in readiness["gates"]} >= {
        "manuscript_formats",
        "manuscript_validation",
        "project_submission_quality_gate",
        "evidence_readiness",
        "primary_source_context",
        "references",
            "search_strategy",
            "declarations",
            "benchmark",
            "figures",
        }
    assert all(gate["status"] == "pass" for gate in readiness["gates"])
    assert "MetaAgent Submission Readiness" in readiness_html
    assert "mini_ready" in readiness_html


def test_submission_readiness_blocks_missing_project_submission_gate(tmp_path: Path) -> None:
    project = Project("missing project submission gate", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "# Manuscript\n\n## Abstract\n\nResult.\n\n## Methods\n\nMethod.\n\n## Results\n\nResult.\n\n## Discussion\n\nDiscussion.\n",
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    readiness = _build_submission_readiness_review(
        project,
        pdf_intake_review=None,
        text_source_coverage=None,
        evidence_review={"status": "ready", "summary": {"blockers": 0, "warnings": 0}},
        abstract_audit=None,
        publication_tone_audit=None,
        readability_audit=None,
        clinical_interpretation_audit=None,
        reference_audit=None,
        citation_audit=None,
        prisma_audit=None,
        search_strategy_audit=None,
        figure_audit=None,
        figure_legend_audit=None,
        cross_reference_audit=None,
        table_footnote_audit=None,
        llm_reliability_audit=None,
        risk_of_bias_completeness=None,
        calculation_audit=None,
        primary_source_trace=None,
        primary_result_audit=None,
        claim_support_audit=None,
        benchmark_review=None,
        publication_similarity=None,
    )

    gate = next(item for item in readiness["gates"] if item["id"] == "project_submission_quality_gate")
    assert readiness["status"] == "blocked"
    assert readiness["passed"] is False
    assert gate["status"] == "fail"
    assert "project_submission_quality_gate" in gate["detail"]


def test_submission_readiness_separates_validation_warnings_from_blockers(tmp_path: Path) -> None:
    project = Project("submission validation detail package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "",
            "## Methods",
            "The search strategy used PubMed terms for the target condition [1].",
            "",
            "## Results",
            "The source-verified trial row is retained in the package [1].",
            "",
            "## Discussion",
            "The result should be interpreted with its source context and certainty assessment [1].",
            "",
            "## Declarations",
            "No new participant data were collected.",
            "",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_validation.json",
        {
            "passed": True,
            "issues": [
                {
                    "kind": "limited_text_source_warning_suppressed",
                    "severity": "fixed",
                    "message": "Kept limited source notice in the package review.",
                },
                {
                    "kind": "evidence_readiness_warning",
                    "severity": "warning",
                    "message": "3 retrieved/screened records use limited source text.",
                },
            ],
        },
        subdir="manuscript",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "condition AND trial")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    gate = next(item for item in readiness["gates"] if item["id"] == "manuscript_validation")
    assert gate["status"] == "pass"
    assert gate["passed"] is True
    assert "blocking=0" in gate["detail"]
    assert "warnings=1" in gate["detail"]
    assert "fixed=1" in gate["detail"]
    assert "total=2" in gate["detail"]
    assert "validation issue(s)" not in gate["detail"]


def test_publication_similarity_review_scores_formal_meta_profile_above_threshold(tmp_path: Path) -> None:
    project = Project("publication similarity high", output_dir=tmp_path / uuid4().hex)
    exact_query = "heart failure AND treatment"
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Formal meta-analysis",
            "## Abstract\nImportance: clinical burden. Objective: estimate treatment effect. Data sources: PubMed. Study selection: randomized trials. Results: pooled HR 0.81 (95% CI 0.74 to 0.88). Conclusions: evidence supports individualized care.",
            "## Introduction\n" + " ".join(["Heart failure context and prior evidence [1]."] * 80),
            f"## Methods\nSearch strategy ({exact_query}), eligibility criteria, screening, data extraction, risk of bias, statistical analysis, random effects, inverse variance, GRADE, PRISMA, and publication bias were prespecified [2].",
            "## Results\nPRISMA screening identified records, full texts were assessed, two studies were included, pooled HR was 0.81 (95% CI 0.74 to 0.88), I² was 0%, and Figure 1 and Table 2 report the forest plot and effect estimates [1,2].",
            "## Discussion\n" + "\n\n".join([
                "The pooled result should be interpreted through magnitude, direction, uncertainty, and baseline risk [1].",
                "Absolute benefit depends on comparator event risk and number needed to treat [1].",
                "Composite endpoint components, follow-up, censoring, and event adjudication affect clinical meaning [1].",
                "Safety, adverse events, kidney function, tolerability, and discontinuation shape benefit-harm balance [2].",
                "Applicability depends on age, renal function, comorbidity, ejection fraction, and background therapy [2].",
                "Implementation requires monitoring, access, cost, patient values, and shared decision making [2].",
                "Certainty, heterogeneity, publication bias, and risk of bias should temper recommendations [2].",
            ]),
            "## Conclusion\nThe intervention may reduce the primary composite outcome, but clinical use should reflect baseline risk, safety, certainty, and patient preference [1,2].",
            "## Tables\n| Study | HR |\n| A | 0.82 |",
            "## Figures\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\nNo new participant data were collected.",
            "## References\n" + "\n".join(f"[{idx}] Reference {idx}." for idx in range(1, 24)),
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {"status": "ready", "blockers": [], "selected_primary_rows": [{"row_id": "A:0"}, {"row_id": "B:0"}]},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_text("references.bib", "\n".join(f"@article{{ref{idx},title={{Reference {idx}}}}}" for idx in range(1, 24)))
    project.save_text("search_query.txt", exact_query)
    project.save_text("draft.docx", "docx placeholder", subdir="manuscript")
    project.save_text("draft.pdf", "pdf placeholder", subdir="manuscript")
    review = build_publication_similarity_review(
        project,
        abstract_audit={"summary": {"present_labels": 6, "required_labels": 6, "forbidden_phrase_count": 0}},
        publication_tone_audit={"passed": True},
        readability_audit={"passed": True},
        clinical_interpretation_audit={
            "passed": True,
            "summary": {
                "domain_count": 8,
                "covered_domains": 8,
                "result_context_present": True,
                "discussion_paragraph_count": 7,
                "process_framing_paragraphs": 0,
                "redundant_domain_count": 0,
                "failed_issues": 0,
            },
        },
        citation_audit={
            "summary": {
                "reference_entries": 23,
                "citation_density_per_1000_words": 8.0,
                "introduction_cited_paragraph_rate": 0.8,
                "discussion_cited_paragraph_rate": 0.9,
                "failed_issues": 0,
                "mechanical_citation_density_paragraphs": 0,
                "repeated_large_citation_clusters": 0,
            },
        },
        prisma_audit={"passed": True},
        figure_audit={"passed": True},
        figure_legend_audit={"passed": True},
        cross_reference_audit={"passed": True},
        table_footnote_audit={"passed": True},
        calculation_audit={"passed": True},
        primary_source_trace={"passed": True},
        benchmark_review={
            "benchmark_id": "published_anchor_trial_pair",
            "gates": [
                {"id": "study_count", "passed": True},
                {"id": "participant_count", "passed": True},
                {"id": "primary_effect", "passed": True},
            ],
            "published_anchor": {
                "effect_measure": "HR",
                "n_trials": 2,
                "n_participants": 12251,
                "effect": 0.80,
                "ci_lower": 0.73,
                "ci_upper": 0.87,
            },
            "observed_primary": {
                "effect_measure": "HR",
                "n_studies": 2,
                "total_participants": 12251,
                "effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
        },
    )

    assert review is not None
    assert review["passed"] is True
    assert review["similarity_score"] >= 85
    assert review["summary"]["components_passing"] >= 8
    benchmark_component = next(item for item in review["components"] if item["id"] == "published_benchmark_alignment")
    assert benchmark_component["passed"] is True
    assert benchmark_component["percent"] >= 85


def test_publication_similarity_methods_requires_exact_search_query_in_manuscript(tmp_path: Path) -> None:
    project = Project("publication similarity exact search query", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Formal meta-analysis",
            "## Abstract\nImportance: clinical burden. Objective: estimate treatment effect. Data sources: PubMed. Study selection: randomized trials. Results: pooled OR 0.66 (95% CI 0.53 to 0.82). Conclusions: evidence supports cautious interpretation.",
            "## Introduction\n" + " ".join(["Clinical context and prior evidence were summarized [1]."] * 80),
            "## Methods\nFull search query: (COVID-19 AND corticosteroids) AND mortality. Search strategy, eligibility criteria, screening, data extraction, risk of bias, statistical analysis, random effects, inverse variance, GRADE, PRISMA, and publication bias were prespecified [2].",
            "## Results\nPRISMA screening identified records, full texts were assessed, two studies were included, pooled OR was 0.66 (95% CI 0.53 to 0.82), I² was 0%, and Figure 1 and Table 2 report the forest plot and effect estimates [1,2].",
            "## Discussion\n" + "\n\n".join([
                "The pooled result should be interpreted through magnitude, direction, uncertainty, and baseline risk [1].",
                "Absolute benefit depends on comparator event risk and number needed to treat [1].",
                "Composite endpoint components, follow-up, censoring, and event adjudication affect clinical meaning [1].",
                "Safety, adverse events, kidney function, tolerability, and discontinuation shape benefit-harm balance [2].",
                "Applicability depends on age, comorbidity, baseline risk, and background therapy [2].",
                "Implementation requires monitoring, access, cost, patient values, and shared decision making [2].",
                "Certainty, heterogeneity, publication bias, and risk of bias should temper recommendations [2].",
            ]),
            "## Conclusion\nThe intervention may reduce the primary outcome, but clinical use should reflect baseline risk, safety, certainty, and patient preference [1,2].",
            "## Tables\n| Study | OR |\n| A | 0.66 |",
            "## Figures\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\nNo new participant data were collected.",
            "## References\n" + "\n".join(f"[{idx}] Reference {idx}." for idx in range(1, 24)),
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "selected_primary_rows": [{"row_id": "A:0"}, {"row_id": "B:0"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_text("references.bib", "\n".join(f"@article{{ref{idx},title={{Reference {idx}}}}}" for idx in range(1, 24)))
    project.save_text("search_query.txt", exact_query)

    review = build_publication_similarity_review(project)

    assert review is not None
    methods_component = next(item for item in review["components"] if item["id"] == "methods_specificity")
    assert methods_component["passed"] is False
    assert "exact_query_reproduced=False" in methods_component["details"]
    assert any(
        issue["code"] == "publication_similarity_component_low:methods_specificity"
        for issue in review["issues"]
    )


def test_publication_similarity_citation_profile_penalizes_excessive_density(tmp_path: Path) -> None:
    project = Project("publication similarity excessive citations", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Formal meta-analysis",
            "## Abstract\nImportance: clinical burden. Objective: estimate treatment effect. Data sources: PubMed. Study selection: randomized trials. Results: pooled HR 0.81 (95% CI 0.74 to 0.88). Conclusions: evidence supports individualized care.",
            "## Introduction\n" + " ".join(["Heart failure context and prior evidence [1]."] * 80),
            "## Methods\nFull search query: heart failure AND treatment. Search strategy, eligibility criteria, screening, data extraction, risk of bias, statistical analysis, random effects, inverse variance, GRADE, PRISMA, and publication bias were prespecified [2].",
            "## Results\nPRISMA screening identified records, full texts were assessed, two studies were included, pooled HR was 0.81 (95% CI 0.74 to 0.88), I² was 0%, and Figure 1 and Table 2 report the forest plot and effect estimates [1,2].",
            "## Discussion\n" + "\n\n".join([
                "The pooled result should be interpreted through magnitude, direction, uncertainty, and baseline risk [1].",
                "Absolute benefit depends on comparator event risk and number needed to treat [1].",
                "Composite endpoint components, follow-up, censoring, and event adjudication affect clinical meaning [1].",
                "Safety, adverse events, kidney function, tolerability, and discontinuation shape benefit-harm balance [2].",
                "Applicability depends on age, renal function, comorbidity, ejection fraction, and background therapy [2].",
                "Implementation requires monitoring, access, cost, patient values, and shared decision making [2].",
                "Certainty, heterogeneity, publication bias, and risk of bias should temper recommendations [2].",
            ]),
            "## Conclusion\nThe intervention may reduce the primary composite outcome, but clinical use should reflect baseline risk, safety, certainty, and patient preference [1,2].",
            "## Tables\n| Study | HR |\n| A | 0.82 |",
            "## Figures\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\nNo new participant data were collected.",
            "## References\n" + "\n".join(f"[{idx}] Reference {idx}." for idx in range(1, 24)),
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {"status": "ready", "blockers": [], "selected_primary_rows": [{"row_id": "A:0"}, {"row_id": "B:0"}]},
        },
        subdir="manuscript",
    )
    project.save_text("search_query.txt", "heart failure AND treatment")

    review = build_publication_similarity_review(
        project,
        citation_audit={
            "summary": {
                "reference_entries": 23,
                "citation_density_per_1000_words": 44.8,
                "introduction_cited_paragraph_rate": 1.0,
                "discussion_cited_paragraph_rate": 1.0,
                "failed_issues": 0,
                "mechanical_citation_density_paragraphs": 0,
                "repeated_large_citation_clusters": 0,
                "excessive_citation_density": True,
            },
        },
    )

    assert review is not None
    citation_component = next(item for item in review["components"] if item["id"] == "citation_profile")
    assert citation_component["score"] == 10.0
    assert citation_component["passed"] is False
    assert "excessive_density=True" in citation_component["details"]


def test_publication_similarity_review_flags_published_benchmark_mismatch(tmp_path: Path) -> None:
    project = Project("publication similarity benchmark mismatch", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Formal meta-analysis",
            "## Abstract\nImportance: clinical burden. Objective: estimate treatment effect. Data sources: PubMed. Study selection: randomized trials. Results: pooled HR 0.81 (95% CI 0.74 to 0.88). Conclusions: evidence supports individualized care.",
            "## Introduction\n" + " ".join(["Heart failure context and prior evidence [1]."] * 80),
            "## Methods\nSearch strategy, eligibility criteria, screening, data extraction, risk of bias, statistical analysis, random effects, inverse variance, GRADE, PRISMA, and publication bias were prespecified [2].",
            "## Results\nPRISMA screening identified records, full texts were assessed, two studies were included, pooled HR was 0.81 (95% CI 0.74 to 0.88), I² was 0%, and Figure 1 and Table 2 report the forest plot and effect estimates [1,2].",
            "## Discussion\n" + "\n\n".join([
                "The pooled result should be interpreted through magnitude, direction, uncertainty, and baseline risk [1].",
                "Absolute benefit depends on comparator event risk and number needed to treat [1].",
                "Composite endpoint components, follow-up, censoring, and event adjudication affect clinical meaning [1].",
                "Safety, adverse events, kidney function, tolerability, and discontinuation shape benefit-harm balance [2].",
                "Applicability depends on age, renal function, comorbidity, ejection fraction, and background therapy [2].",
                "Implementation requires monitoring, access, cost, patient values, and shared decision making [2].",
                "Certainty, heterogeneity, publication bias, and risk of bias should temper recommendations [2].",
            ]),
            "## Conclusion\nThe intervention may reduce the primary composite outcome, but clinical use should reflect baseline risk, safety, certainty, and patient preference [1,2].",
            "## Tables\n| Study | HR |\n| A | 0.82 |",
            "## Figures\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\nNo new participant data were collected.",
            "## References\n" + "\n".join(f"[{idx}] Reference {idx}." for idx in range(1, 24)),
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {"status": "ready", "blockers": [], "selected_primary_rows": [{"row_id": "A:0"}, {"row_id": "B:0"}]},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_text("references.bib", "\n".join(f"@article{{ref{idx},title={{Reference {idx}}}}}" for idx in range(1, 24)))
    project.save_text("search_query.txt", "heart failure AND treatment")
    project.save_text("draft.docx", "docx placeholder", subdir="manuscript")
    project.save_text("draft.pdf", "pdf placeholder", subdir="manuscript")

    review = build_publication_similarity_review(
        project,
        abstract_audit={"summary": {"present_labels": 6, "required_labels": 6, "forbidden_phrase_count": 0}},
        publication_tone_audit={"passed": True},
        readability_audit={"passed": True},
        clinical_interpretation_audit={
            "passed": True,
            "summary": {
                "domain_count": 8,
                "covered_domains": 8,
                "result_context_present": True,
                "discussion_paragraph_count": 7,
                "process_framing_paragraphs": 0,
                "redundant_domain_count": 0,
                "failed_issues": 0,
            },
        },
        citation_audit={
            "summary": {
                "reference_entries": 23,
                "citation_density_per_1000_words": 8.0,
                "introduction_cited_paragraph_rate": 0.8,
                "discussion_cited_paragraph_rate": 0.9,
                "failed_issues": 0,
                "mechanical_citation_density_paragraphs": 0,
                "repeated_large_citation_clusters": 0,
            },
        },
        prisma_audit={"passed": True},
        figure_audit={"passed": True},
        figure_legend_audit={"passed": True},
        cross_reference_audit={"passed": True},
        table_footnote_audit={"passed": True},
        calculation_audit={"passed": True},
        primary_source_trace={"passed": True},
        benchmark_review={
            "benchmark_id": "published_anchor_trial_pair",
            "gates": [
                {"id": "study_count", "passed": False},
                {"id": "participant_count", "passed": False},
                {"id": "primary_effect", "passed": False},
            ],
            "published_anchor": {
                "effect_measure": "HR",
                "n_trials": 2,
                "n_participants": 12251,
                "effect": 0.80,
                "ci_lower": 0.73,
                "ci_upper": 0.87,
            },
            "observed_primary": {
                "effect_measure": "RR",
                "n_studies": 3,
                "total_participants": 6000,
                "effect": 1.20,
                "ci_lower": 0.95,
                "ci_upper": 1.48,
            },
        },
    )

    assert review is not None
    benchmark_component = next(item for item in review["components"] if item["id"] == "published_benchmark_alignment")
    assert benchmark_component["passed"] is False
    assert benchmark_component["percent"] < 85
    assert any(
        issue["code"] == "publication_similarity_component_low:published_benchmark_alignment"
        for issue in review["issues"]
    )


def test_publication_similarity_review_flags_draft_below_85_percent(tmp_path: Path) -> None:
    project = Project("publication similarity low", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Short meta-analysis",
            "## Abstract\nThe result was HR 0.81.",
            "## Introduction\nBrief background.",
            "## Methods\nBrief methods.",
            "## Results\nTwo studies were included.",
            "## Discussion\nBrief interpretation.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {"status": "ready", "blockers": [], "selected_primary_rows": [{"row_id": "A:0"}, {"row_id": "B:0"}]},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    review = build_publication_similarity_review(project)

    assert review is not None
    assert review["passed"] is False
    assert review["similarity_score"] < 85
    assert any(issue["code"] == "publication_similarity_below_threshold" for issue in review["issues"])
    assert review["next_actions"]


def test_submission_readiness_blocks_requested_chinese_manuscript_written_in_english(tmp_path: Path) -> None:
    project = Project("language mismatch package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# English manuscript",
            "",
            "## Abstract",
            "This draft is written in English even though the requested output language was Chinese.",
            "",
            "## Introduction",
            "The introduction remains English.",
            "",
            "## Methods",
            "The methods remain English.",
            "",
            "## Results",
            "The results remain English.",
            "",
            "## Discussion",
            "The discussion remains English.",
            "",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "selected_primary_rows": [{"row_id": "S1:0"}, {"row_id": "S2:0"}],
            },
            "writing_constraints": {"publication_min_main_words": 1},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    readiness = _build_submission_readiness_review(
        project,
        pdf_intake_review=None,
        text_source_coverage=None,
        evidence_review=None,
        abstract_audit=None,
        publication_tone_audit=None,
        readability_audit=None,
        clinical_interpretation_audit=None,
        reference_audit=None,
        citation_audit=None,
        prisma_audit=None,
        search_strategy_audit=None,
        figure_audit=None,
        figure_legend_audit=None,
        cross_reference_audit=None,
        table_footnote_audit=None,
        llm_reliability_audit=None,
        risk_of_bias_completeness=None,
        calculation_audit=None,
        primary_source_trace=None,
        primary_result_audit=None,
        claim_support_audit=None,
        benchmark_review=None,
    )

    language_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_language")
    assert language_gate["status"] == "fail"
    assert "expected=zh" in language_gate["detail"]
    assert "detected=en" in language_gate["detail"]
    assert readiness["passed"] is False


def test_submission_readiness_blocks_heading_only_language_translation(tmp_path: Path) -> None:
    project = Project("mixed language package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# SGLT2抑制剂治疗心力衰竭",
            "",
            "## 摘要",
            "This abstract remains in English, although the requested output language was Chinese.",
            "",
            "## 引言",
            "Heart failure with preserved ejection fraction remains a major source of morbidity and recurrent hospitalization.",
            "",
            "## 方法",
            "We searched PubMed, Embase, trial registries, and reference lists using prespecified eligibility criteria.",
            "",
            "## 结果",
            "The pooled hazard ratio favored SGLT2 inhibitors for the primary composite endpoint.",
            "",
            "## 讨论",
            "The findings should be interpreted with attention to kidney function, ejection fraction, and background therapy.",
            "",
            "## 声明",
            "### 伦理批准",
            "No new participant data were collected.",
            "### 数据和代码可用性",
            "Aggregate data are included with the supplementary files.",
            "### 资金",
            "No dedicated funding was recorded.",
            "### 利益冲突",
            "No competing interests were recorded.",
            "",
            "## 参考文献",
            "［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "selected_primary_rows": [{"row_id": "S1:0"}, {"row_id": "S2:0"}],
            },
            "writing_constraints": {"publication_min_main_words": 1},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    readiness = _build_submission_readiness_review(
        project,
        pdf_intake_review=None,
        text_source_coverage=None,
        evidence_review=None,
        abstract_audit=None,
        publication_tone_audit=None,
        readability_audit=None,
        clinical_interpretation_audit=None,
        reference_audit=None,
        citation_audit=None,
        prisma_audit=None,
        search_strategy_audit=None,
        figure_audit=None,
        figure_legend_audit=None,
        cross_reference_audit=None,
        table_footnote_audit=None,
        llm_reliability_audit=None,
        risk_of_bias_completeness=None,
        calculation_audit=None,
        primary_source_trace=None,
        primary_result_audit=None,
        claim_support_audit=None,
        benchmark_review=None,
    )

    language_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_language")
    assert language_gate["status"] == "fail"
    assert "expected=zh" in language_gate["detail"]
    assert "detected=mixed" in language_gate["detail"]
    assert readiness["passed"] is False


def test_review_language_accepts_chinese_medical_abbreviations_and_search_code_blocks() -> None:
    manuscript = "\n".join([
        "# HFmrEF/HFpEF患者中SGLT2抑制剂的Meta分析",
        "",
        "## 摘要",
        "本研究评价SGLT2抑制剂相较于安慰剂对心血管死亡或心力衰竭住院的影响，主要效应量为HR和95% CI。",
        "",
        "## 引言",
        "HFmrEF/HFpEF患者具有明显临床异质性，GRADE评价、PRISMA流程和HR结果均需在中文正文中保留通用英文缩写。",
        "这些英文缩写是医学术语，不代表正文改成英文。",
        "",
        "## 方法",
        "完整检索式如下：",
        "```text",
        "(heart failure[TIAB] OR preserved ejection fraction[TIAB] OR mildly reduced ejection fraction[TIAB]) AND (SGLT2[TIAB] OR dapagliflozin[TIAB] OR empagliflozin[TIAB])",
        "(heart failure hospitalization[TIAB] OR cardiovascular death[TIAB] OR clinical trial[TIAB]) AND randomized controlled trial[TIAB]",
        "(sodium glucose cotransporter two inhibitor[TIAB] OR placebo controlled trial[TIAB]) AND follow-up[TIAB]",
        "```",
        "",
        "## 结果",
        "合并HR为0.81（95% CI 0.74至0.88），异质性较低。",
        "",
        "## 讨论",
        "中文医学稿件应允许必要英文缩写、数据库名称和代码块检索式。",
        "",
        "## 参考文献",
        "［1］ Trial report.",
    ])

    assert _review_language_from_text(manuscript) == "zh"


def test_publication_section_shape_accepts_chinese_headings() -> None:
    manuscript = "\n".join([
        "# 中文稿件",
        "",
        "## 摘要",
        "摘要内容。",
        "",
        "## 引言",
        "引言内容。",
        "",
        "## 方法",
        "方法内容。",
        "",
        "## 结果",
        "结果内容。",
        "",
        "## 讨论",
        "讨论内容。",
    ])

    assert _has_publication_section_shape(manuscript) is True


def test_submission_readiness_warns_when_llm_usage_has_recovered_output_issues(tmp_path: Path) -> None:
    project = Project("llm reliability package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# LLM reliability manuscript",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "## Methods",
            "PubMed Boolean strategy: corticosteroids AND COVID-19 AND mortality.",
            "## Results",
            "One source-verified trial row is shown for this fixture.",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "llm_usage_manifest.json",
        {
            "schema_version": 1,
            "summary": {
                "total_calls": 2,
                "total_tokens": 900,
                "prompt_tokens": 300,
                "completion_tokens": 600,
                "estimated_cost_usd": 0.001,
            },
            "events": [
                {
                    "timestamp": "2026-05-23T00:00:00Z",
                    "model": "qwen3.6-plus",
                    "endpoint": "responses",
                    "prompt_tokens": 100,
                    "completion_tokens": 512,
                    "total_tokens": 612,
                    "max_tokens": 512,
                    "finish_reason": "incomplete",
                    "retryable_output_issue": "status:incomplete",
                    "near_truncation": True,
                },
                {
                    "timestamp": "2026-05-23T00:00:01Z",
                    "model": "qwen3.6-plus",
                    "endpoint": "responses",
                    "prompt_tokens": 200,
                    "completion_tokens": 88,
                    "total_tokens": 288,
                    "max_tokens": 1536,
                    "finish_reason": "completed",
                    "retryable_output_issue": "",
                    "near_truncation": False,
                },
            ],
        },
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        llm_audit = json.loads(zf.read("review/llm_reliability_audit.json"))
        llm_html = zf.read("review/llm_reliability_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    gate = next(gate for gate in readiness["gates"] if gate["id"] == "llm_reliability")
    assert llm_audit["summary"]["retryable_output_issues"] == 1
    assert llm_audit["summary"]["near_truncation_events"] == 1
    assert llm_audit["issues"][0]["code"] == "llm_retryable_output_issue"
    assert gate["status"] == "warn"
    assert readiness["summary"]["warning_gates"] >= 1
    assert "Responses incomplete" in llm_html
    assert manifest["review"]["llm_reliability_audit_included"] is True
    assert manifest["review"]["llm_reliability_retryable_output_issues"] == 1


def test_submission_readiness_passes_when_retryable_llm_issue_recovered_without_truncation(tmp_path: Path) -> None:
    project = Project("llm reliability recovered package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# LLM reliability manuscript",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "## Methods",
            "PubMed Boolean strategy: corticosteroids AND COVID-19 AND mortality.",
            "## Results",
            "One source-verified trial row is shown for this fixture.",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "evidence_gap"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "llm_usage_manifest.json",
        {
            "schema_version": 1,
            "summary": {"total_calls": 2, "total_tokens": 900},
            "events": [
                {
                    "timestamp": "2026-05-23T00:00:00Z",
                    "model": "qwen3.6-plus",
                    "endpoint": "responses",
                    "finish_reason": "incomplete",
                    "retryable_output_issue": "status:incomplete",
                    "near_truncation": False,
                },
                {
                    "timestamp": "2026-05-23T00:00:01Z",
                    "model": "qwen3.6-plus",
                    "endpoint": "responses",
                    "finish_reason": "completed",
                    "retryable_output_issue": "",
                    "near_truncation": False,
                },
            ],
        },
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        llm_audit = json.loads(zf.read("review/llm_reliability_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    gate = next(gate for gate in readiness["gates"] if gate["id"] == "llm_reliability")
    assert llm_audit["summary"]["retryable_output_issues"] == 1
    assert llm_audit["summary"]["near_truncation_events"] == 0
    assert gate["status"] == "pass"


def test_submission_readiness_blocks_primary_meta_with_missing_formal_rob(tmp_path: Path) -> None:
    project = Project("rob completeness package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# RoB completeness manuscript",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "## Methods",
            "PubMed Boolean strategy: corticosteroids AND COVID-19 AND mortality.",
            "## Results",
            "Two studies contributed to the primary meta-analysis.",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report. *Journal*. 2024.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "evidence_readiness": {"status": "ready", "selected_primary_rows": []}},
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "28-day mortality",
                "n_studies": 2,
                "studies": [
                    {"study_id": "S1", "study_label": "Smith 2024"},
                    {"study_id": "S2", "study_label": "Jones 2024"},
                ],
            },
            "secondary_outcomes": [],
        },
        subdir="analysis",
    )
    project.save_json(
        "rob_results.json",
        [
            {
                "study_id": "S1",
                "tool_used": "RoB 2",
                "overall_judgment": "Low risk",
                "is_synthetic": False,
                "domains": [
                    {"domain": "Randomization process", "judgment": "Low risk", "support": "Randomized."},
                    {"domain": "Deviations from intended interventions", "judgment": "Low risk", "support": "Blinded."},
                    {"domain": "Missing outcome data", "judgment": "Low risk", "support": "Complete."},
                    {"domain": "Measurement of the outcome", "judgment": "Low risk", "support": "Adjudicated."},
                    {"domain": "Selection of the reported result", "judgment": "Low risk", "support": "Protocol."},
                ],
            }
        ],
        subdir="risk_of_bias",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        rob_audit = json.loads(zf.read("review/risk_of_bias_completeness.json"))
        rob_audit_html = zf.read("review/risk_of_bias_completeness.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    gate = next(gate for gate in readiness["gates"] if gate["id"] == "risk_of_bias_completeness")
    assert rob_audit["summary"]["primary_contributing_studies"] == 2
    assert rob_audit["summary"]["missing_formal_rob"] == 1
    assert rob_audit["summary"]["synthetic_rob"] == 0
    assert rob_audit["issues"][0]["study_id"] == "S2"
    assert gate["status"] == "fail"
    assert "formal risk-of-bias assessment" in gate["detail"]
    assert "Jones 2024" in rob_audit_html


def test_submission_readiness_blocks_primary_meta_with_synthetic_rob(tmp_path: Path) -> None:
    project = Project("synthetic rob package", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# Synthetic RoB manuscript", subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "evidence_readiness": {"status": "ready", "selected_primary_rows": []}},
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "28-day mortality",
                "n_studies": 1,
                "studies": [{"study_id": "S1", "study_label": "Smith 2024"}],
            },
            "secondary_outcomes": [],
        },
        subdir="analysis",
    )
    project.save_json(
        "rob_results.json",
        [
            {
                "study_id": "S1",
                "tool_used": "RoB 2",
                "overall_judgment": "Not assessed (insufficient information)",
                "is_synthetic": True,
                "domains": [],
            }
        ],
        subdir="risk_of_bias",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        rob_audit = json.loads(zf.read("review/risk_of_bias_completeness.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    gate = next(gate for gate in readiness["gates"] if gate["id"] == "risk_of_bias_completeness")
    assert rob_audit["summary"]["missing_formal_rob"] == 0
    assert rob_audit["summary"]["synthetic_rob"] == 1
    assert rob_audit["issues"][0]["code"] == "primary_study_synthetic_rob"
    assert gate["status"] == "fail"


def test_artifact_package_includes_human_readable_calculation_audit(tmp_path: Path) -> None:
    project = Project("calculation audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Calculation audit manuscript",
            "## Abstract",
            "Primary outcome OR 0.50 (95% CI 0.20 to 1.20).",
            "## Methods",
            "The log odds ratio was calculated from 2 x 2 tables.",
            "## Results",
            "Trial-level estimates are available.",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "mortality",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                        "source_location": "Table 2",
                        "source_quote": "Mortality was 1/10 vs 2/10.",
                        "source_quote_verified": True,
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 1,
                "effect_measure": "OR",
                "pooled_effect": 0.5,
                "ci_lower": 0.2,
                "ci_upper": 1.2,
                "p_value": 0.12,
                "pooled_log": -0.6931471805599453,
                "ci_lower_log": -1.6094379124341003,
                "ci_upper_log": 0.1823215567939546,
                "model": "fixed",
                "q_statistic": 0.0,
                "q_p_value": 1.0,
                "i_squared": 0.0,
                "tau_squared": 0.0,
                "studies": [
                    {
                        "study_id": "S1",
                        "study_label": "Smith 2024",
                        "yi": -0.6931471805599453,
                        "vi": 1.8,
                        "se": 1.3416407864998738,
                        "weight": 100.0,
                    }
                ],
            }
        },
        subdir="analysis",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        audit = json.loads(zf.read("review/calculation_audit.json"))
        audit_html = zf.read("review/calculation_audit.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/calculation_audit.json" in names
    assert "review/calculation_audit.html" in names
    assert audit["summary"]["effect_measure"] == "OR"
    assert audit["summary"]["n_studies"] == 1
    assert audit["formulas"]["OR"]["effect"].startswith("log(")
    assert audit["rows"][0]["study_id"] == "S1"
    assert audit["rows"][0]["effect_original"] == 0.5
    assert audit["rows"][0]["events_intervention"] == 1
    assert audit["rows"][0]["source_location"] == "Table 2"
    assert "Meta-Analysis Calculation Audit" in audit_html
    assert "Smith 2024" in audit_html
    assert manifest["review"]["calculation_audit_included"] is True
    assert manifest["review"]["calculation_audit_rows"] == 1


def test_artifact_package_includes_primary_source_trace_review(tmp_path: Path) -> None:
    project = Project("primary source trace package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Primary source trace manuscript",
            "## Abstract",
            "Primary outcome OR 0.50 (95% CI 0.20 to 1.20).",
            "## Methods",
            "The log odds ratio was calculated from 2 x 2 tables.",
            "## Results",
            "Trial-level estimates are available.",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "mortality",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                        "source_location": "Table 2",
                        "source_page": 5,
                        "source_section": "Results",
                        "source_quote": "Mortality was 1/10 vs 2/10.",
                        "source_quote_verified": True,
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 1,
                "effect_measure": "OR",
                "pooled_effect": 0.5,
                "ci_lower": 0.2,
                "ci_upper": 1.2,
                "model": "fixed",
                "studies": [
                    {
                        "study_id": "S1",
                        "study_label": "Smith 2024",
                        "yi": -0.6931471805599453,
                        "vi": 1.8,
                        "se": 1.3416407864998738,
                        "weight": 100.0,
                    }
                ],
            }
        },
        subdir="analysis",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        trace = json.loads(zf.read("review/primary_source_trace.json"))
        trace_html = zf.read("review/primary_source_trace.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/primary_source_trace.json" in names
    assert "review/primary_source_trace.html" in names
    assert trace["passed"] is True
    assert trace["summary"]["row_count"] == 1
    assert trace["summary"]["source_traceable_rows"] == 1
    assert trace["summary"]["missing_source_quote_rows"] == 0
    assert trace["summary"]["missing_source_location_rows"] == 0
    assert trace["rows"][0]["trace_status"] == "traceable"
    assert trace["rows"][0]["values"]["events_intervention"] == 1
    assert trace["rows"][0]["source"]["page"] == 5
    assert "Primary Source Trace" in trace_html
    assert "Mortality was 1/10 vs 2/10." in trace_html
    assert manifest["review"]["primary_source_trace_included"] is True
    assert manifest["review"]["primary_source_trace_rows"] == 1
    assert manifest["review"]["primary_source_trace_failed_issues"] == 0


def test_primary_source_trace_review_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese primary source trace package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n主要结局RR 0.50（95% CI 0.20至1.20）。",
            "## 方法\n提取主要结局原文来源并复核。",
            "## 结果\n主要分析纳入1项研究。",
            "## 讨论\n主要结果需保留原文溯源。",
            "## 参考文献\n［1］ Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "evidence_readiness": {
                "status": "ready",
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "死亡率",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                        "source_location": "表2",
                        "source_page": 5,
                        "source_section": "结果",
                        "source_quote": "死亡率为1/10对2/10。",
                        "source_quote_verified": True,
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "死亡率",
                "n_studies": 1,
                "effect_measure": "RR",
                "pooled_effect": 0.5,
                "ci_lower": 0.2,
                "ci_upper": 1.2,
                "model": "fixed",
                "studies": [
                    {
                        "study_id": "S1",
                        "study_label": "Smith 2024",
                        "yi": -0.6931471805599453,
                        "vi": 1.8,
                        "se": 1.3416407864998738,
                        "weight": 100.0,
                    }
                ],
            }
        },
        subdir="analysis",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        trace = json.loads(zf.read("review/primary_source_trace.json"))
        trace_html = zf.read("review/primary_source_trace.html").decode("utf-8")

    assert trace["language"] == "zh"
    assert '<html lang="zh">' in trace_html
    assert "MetaAgent 主要结果溯源" in trace_html
    assert "展示主要分析的每个数字行" in trace_html
    assert "结局" in trace_html
    assert "效应量" in trace_html
    assert "可溯源" in trace_html
    assert "缺失原文引用" in trace_html
    assert "主要分析行" in trace_html
    assert "研究" in trace_html
    assert "状态" in trace_html
    assert "来源" in trace_html
    assert "页码" in trace_html
    assert "已核验" in trace_html
    assert "问题" in trace_html
    assert "MetaAgent Primary Source Trace" not in trace_html
    assert "Missing quote" not in trace_html


def test_submission_readiness_blocks_primary_source_trace_gaps(tmp_path: Path) -> None:
    project = Project("primary source trace gap package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    long_text = " ".join(["auditability"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full source trace audit manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text}",
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables using inverse-variance weighting."
            ),
            (
                "## Results\nThe pooled estimate was OR 0.66 (95% CI 0.53 to 0.82). "
                "The primary meta-analysis included 2 trials totaling 60 participants. "
                "There were 4/30 deaths in the corticosteroid groups and 8/30 deaths in the control groups."
            ),
            "### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |\n| Trial 2 | 3/20 vs 6/20 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Discussion\nThe result is interpreted with source verification.",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_location": "Table 2",
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row.get("source_quote") or "",
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "pooled_log": -0.415515,
                "ci_lower_log": -0.634878,
                "ci_upper_log": -0.198451,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n")
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        trace = json.loads(zf.read("review/primary_source_trace.json"))
        trace_html = zf.read("review/primary_source_trace.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    source_trace_gate = next(gate for gate in readiness["gates"] if gate["id"] == "primary_source_trace")
    assert trace["passed"] is False
    assert trace["summary"]["missing_source_quote_rows"] == 1
    assert trace["summary"]["missing_source_location_rows"] == 1
    assert trace["issues"][0]["code"] == "primary_source_quote_missing"
    assert source_trace_gate["status"] == "fail"
    assert "traceable=1/2" in source_trace_gate["detail"]
    assert readiness["status"] == "blocked"
    assert manifest["review"]["primary_source_trace_failed_issues"] == 2
    assert "Missing source quote" in trace_html


def test_submission_readiness_blocks_full_length_meta_without_calculation_audit(tmp_path: Path) -> None:
    project = Project("missing calculation audit package", output_dir=tmp_path / uuid4().hex)
    long_text = " ".join(["auditability"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full meta-analysis",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text}",
            "## Methods\nFull search query: corticosteroids AND COVID-19 AND mortality. The log odds ratio was calculated from 2 x 2 tables using inverse-variance weighting.",
            "## Results\nThe pooled estimate was OR 0.66 (95% CI 0.53 to 0.82).",
            "### Table 1. Trial-level effects\n| Study | OR | Source |\n|---|---:|---|\n| Smith | 0.50 | Table 2 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Discussion\nThe result is interpreted with source verification.",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 5,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 5/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_location": "Table 2",
                    "source_quote_match": row["source_quote"],
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ]
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 5/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}\n@article{jones2024,title={Trial report}}")
    project.save_text("search_query.txt", "corticosteroids AND COVID-19 AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    calculation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "calculation_audit")
    assert "review/calculation_audit.json" not in names
    assert calculation_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["submission"]["passed"] is False
    assert manifest["review"]["calculation_audit_included"] is False


def test_submission_readiness_blocks_meta_when_calculation_audit_rows_are_not_verifiable(tmp_path: Path) -> None:
    project = Project("weak calculation audit package", output_dir=tmp_path / uuid4().hex)
    long_text = " ".join(["auditability"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full meta-analysis",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text}",
            "## Methods\nFull search query: corticosteroids AND COVID-19 AND mortality. The log odds ratio was calculated from 2 x 2 tables using inverse-variance weighting.",
            "## Results\nThe pooled estimate was OR 0.66 (95% CI 0.53 to 0.82).",
            "### Table 1. Trial-level effects\n| Study | OR | Source |\n|---|---:|---|\n| Smith | 0.50 | Table 2 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Discussion\nThe result is interpreted with source verification.",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "total_control": 20,
            "source_quote": "Mortality was reported incompletely.",
            "source_quote_verified": False,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "p_value": 0.001,
                "pooled_log": -0.415515,
                "ci_lower_log": -0.634878,
                "ci_upper_log": -0.198451,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Smith 2024", "yi": -0.693147, "vi": 1.8, "se": 1.341641, "weight": 55.0},
                    {"study_id": "S2", "study_label": "Jones 2024", "yi": -0.223144, "vi": 2.0, "se": 1.414214, "weight": 45.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_location": "Table 2",
                    "source_quote_match": row["source_quote"],
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ]
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was reported incompletely.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}\n@article{jones2024,title={Trial report}}")
    project.save_text("search_query.txt", "corticosteroids AND COVID-19 AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/calculation_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    calculation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "calculation_audit")
    assert audit["summary"]["source_quote_verified_rows"] == 1
    assert audit["summary"]["formula_inputs_complete_rows"] == 1
    assert calculation_gate["status"] == "fail"
    assert "source_quote_verified=1/2" in calculation_gate["detail"]
    assert "formula_inputs_complete=1/2" in calculation_gate["detail"]
    assert readiness["status"] == "blocked"


def test_artifact_package_benchmark_review_embeds_published_alignment_with_calculation_audit(tmp_path: Path) -> None:
    project = Project("benchmark alignment package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Benchmark alignment manuscript",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "## Methods",
            "Full search query: corticosteroids AND COVID-19 AND mortality. The log odds ratio was calculated from 2 x 2 tables.",
            "## Results",
            "The pooled estimate was OR 0.66 (95% CI 0.53 to 0.82).",
            "## Figures",
            "### Figure 1. Forest plot",
            "![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "mortality",
                        "events_intervention": 1,
                        "total_intervention": 10,
                        "events_control": 2,
                        "total_control": 10,
                        "source_quote": "Mortality was 1/10 vs 2/10.",
                        "source_quote_verified": True,
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 1,
                "effect_measure": "OR",
                "pooled_effect": 0.6593,
                "ci_lower": 0.5323,
                "ci_upper": 0.8167,
                "p_value": 0.001,
                "pooled_log": -0.416576,
                "ci_lower_log": -0.630548,
                "ci_upper_log": -0.202483,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Smith 2024", "yi": -0.416576, "vi": 1.8, "se": 1.341641, "weight": 100.0}
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_benchmark",
            "summary_card": {
                "benchmark_id": "mini_benchmark",
                "status": "passed",
                "passed": True,
                "published_anchor": {
                    "n_trials": 1,
                    "n_participants": 20,
                    "effect_measure": "OR",
                    "model_preference": "fixed",
                    "effect": 0.66,
                    "ci_lower": 0.53,
                    "ci_upper": 0.82,
                },
                "observed_primary": {
                    "effect_measure": "OR",
                    "model_preference": "fixed",
                    "n_studies": 1,
                    "effect": 0.6593,
                    "ci_lower": 0.5323,
                    "ci_upper": 0.8167,
                    "total_participants": 20,
                    "participant_difference": 0,
                },
                "gates": [],
                "failing_gates": [],
                "missing_primary_full_texts": [],
                "next_actions": [],
            },
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "corticosteroids AND COVID-19 AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/benchmark_review.json"))
        html = zf.read("review/benchmark_review.html").decode("utf-8")

    alignment = review["benchmark_alignment"]
    assert alignment["published"]["effect"] == 0.66
    assert alignment["observed"]["effect"] == 0.6593
    assert alignment["calculation_audit"]["row_count"] == 1
    assert alignment["calculation_audit"]["source_quote_verified_rows"] == 1
    assert alignment["differences"]["effect"] == -0.0007
    assert "Published vs Reproduced Alignment" in html
    assert "source_quote_verified=1/1" in html


def test_benchmark_review_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese benchmark review package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "## 摘要",
            "主要结局OR 0.66（95% CI 0.53至0.82）。",
            "## 方法",
            "系统将复现结果与已发表锚点进行比较。",
            "## 结果",
            "主要分析需要人工复核。",
            "## 参考文献",
            "［1］ Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "evidence_readiness": {"status": "ready"},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_benchmark",
            "summary_card": {
                "benchmark_id": "mini_benchmark",
                "status": "blocked",
                "passed": False,
                "published_anchor": {
                    "n_trials": 1,
                    "n_participants": 20,
                    "effect_measure": "OR",
                    "model_preference": "fixed",
                    "effect": 0.66,
                    "ci_lower": 0.53,
                    "ci_upper": 0.82,
                },
                "observed_primary": {
                    "effect_measure": "OR",
                    "model_preference": "fixed",
                    "n_studies": 1,
                    "effect": 0.72,
                    "ci_lower": 0.55,
                    "ci_upper": 0.90,
                    "total_participants": 18,
                    "participant_difference": -2,
                },
                "gates": [
                    {
                        "gate": "pooled_effect",
                        "label": "Pooled effect",
                        "passed": False,
                        "participant_difference": -2,
                        "failure_reasons": ["pooled_effect_mismatch"],
                    }
                ],
                "failing_gates": [
                    {
                        "gate": "pooled_effect",
                        "label": "Pooled effect",
                        "participant_difference": -2,
                        "failure_reasons": ["pooled_effect_mismatch"],
                    }
                ],
                "missing_primary_full_texts": [
                    {
                        "trial_id": "T1",
                        "trial_name": "Trial One",
                        "publication_pmids": ["32678530"],
                        "publication_dois": ["10.1000/example"],
                        "registration_id": "NCT00000001",
                    }
                ],
                "next_actions": [
                    {
                        "type": "upload_full_text",
                        "message": "Upload the missing primary full text before submission.",
                    }
                ],
            },
            "primary_analysis": {
                "missing": [
                    {
                        "trial_id": "T1",
                        "trial_name": "Trial One",
                        "registration_id": "NCT00000001",
                    }
                ]
            },
        },
        subdir="benchmark",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/benchmark_review.json"))
        html = zf.read("review/benchmark_review.html").decode("utf-8")

    assert review["language"] == "zh"
    assert '<html lang="zh">' in html
    assert "MetaAgent 基准复现复核" in html
    assert "将本项目结果与已发表基准锚点逐项比较" in html
    assert "基准" in html
    assert "状态" in html
    assert "发表锚点" in html
    assert "复现主要结果" in html
    assert "发表与复现对齐" in html
    assert "未通过质量门" in html
    assert "缺失主要全文" in html
    assert "来源获取任务" in html
    assert "下一步动作" in html
    assert "汇总效应不匹配" in html
    assert "上传缺失的主要全文" in html
    assert "MetaAgent Benchmark Review" not in html
    assert "Published-anchor comparison" not in html


def test_artifact_package_benchmark_alignment_surfaces_equivalent_zero_tau_model_difference(tmp_path: Path) -> None:
    project = Project("benchmark model compatibility package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Benchmark model compatibility manuscript",
            "## Abstract",
            "Primary outcome HR 0.81 (95% CI 0.74 to 0.88).",
            "## Methods",
            "Full search query: heart failure AND SGLT2. Hazard ratios were pooled using inverse-variance methods.",
            "## Results",
            "The pooled estimate was HR 0.81 (95% CI 0.74 to 0.88).",
            "## Figures",
            "### Figure 1. Forest plot",
            "![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "selected_primary_rows": [
                    {
                        "row_id": "S1:0",
                        "study_id": "S1",
                        "outcome_name": "primary composite",
                        "effect_measure": "HR",
                        "effect_value": 0.81,
                        "ci_lower": 0.74,
                        "ci_upper": 0.88,
                        "source_quote": "The hazard ratio was 0.81.",
                        "source_quote_verified": True,
                    }
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "primary composite",
                "n_studies": 2,
                "effect_measure": "HR",
                "pooled_effect": 0.8069,
                "ci_lower": 0.7395,
                "ci_upper": 0.8805,
                "p_value": 0.001,
                "pooled_log": -0.21455,
                "ci_lower_log": -0.30178,
                "ci_upper_log": -0.12727,
                "model": "random",
                "tau_squared": 0.0,
                "i_squared": 0.0,
                "studies": [
                    {"study_id": "S1", "study_label": "Smith 2024", "yi": -0.2, "vi": 0.01, "se": 0.1, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Jones 2024", "yi": -0.22, "vi": 0.01, "se": 0.1, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_benchmark",
            "summary_card": {
                "benchmark_id": "mini_benchmark",
                "status": "passed",
                "passed": True,
                "published_anchor": {
                    "n_trials": 2,
                    "n_participants": 100,
                    "effect_measure": "HR",
                    "model_preference": "fixed",
                    "effect": 0.80,
                    "ci_lower": 0.73,
                    "ci_upper": 0.87,
                },
                "observed_primary": {
                    "effect_measure": "HR",
                    "model_preference": "random",
                    "n_studies": 2,
                    "effect": 0.8069,
                    "ci_lower": 0.7395,
                    "ci_upper": 0.8805,
                    "total_participants": 100,
                    "participant_difference": 0,
                },
                "gates": [
                    {
                        "gate": "pooled_effect",
                        "passed": True,
                        "expected": {"effect_measure": "HR", "model_preference": "fixed"},
                        "observed": {"effect_measure": "HR", "model_preference": "random"},
                        "failure_reasons": [],
                    }
                ],
                "failing_gates": [],
                "missing_primary_full_texts": [],
                "next_actions": [],
            },
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "heart failure AND SGLT2")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/benchmark_review.json"))
        html = zf.read("review/benchmark_review.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))

    alignment = review["benchmark_alignment"]
    assert alignment["model_compatibility_notes"] == ["random_model_equivalent_to_fixed_tau_zero"]
    assert alignment["differences"]["model_preference"] == "published=fixed; observed=random; tau_squared=0.0"
    assert manifest["review"]["benchmark_alignment_model_compatibility_notes"] == 1
    assert "Random-effects model is numerically equivalent to the fixed-effect anchor because tau-squared is 0" in html


def test_submission_readiness_blocks_reference_list_bibtex_mismatch(tmp_path: Path) -> None:
    project = Project("reference audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full reference audit manuscript",
            "## Abstract\n" + " ".join(["abstract"] * 350),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            (
                "## Methods\nFull search query: (COVID-19[tiab] AND corticosteroids[tiab]) "
                "AND (mortality[tiab] OR death[tiab]). The log odds ratio was calculated from 2 x 2 tables, "
                "and inverse-variance fixed-effect meta-analysis was used. "
                + " ".join(["methods"] * 1450)
            ),
            (
                "## Results\nThe pooled estimate was OR 0.66 (95% CI 0.53 to 0.82), with heterogeneity I2 0%. "
                + " ".join(["results"] * 1300)
            ),
            "## Discussion\n" + " ".join(["discussion"] * 1500),
            "## Tables\n### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n\n[2] EU Clinical Trials Register. Trial results.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 6/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row["source_quote"],
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 6/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )
    project.save_text("search_query.txt", "COVID-19 AND corticosteroids AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        reference_audit = json.loads(zf.read("review/reference_audit.json"))
        reference_html = zf.read("review/reference_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/reference_audit.json" in names
    assert "review/reference_audit.html" in names
    assert reference_audit["summary"]["bib_entries"] == 1
    assert reference_audit["summary"]["manuscript_references"] == 2
    assert reference_audit["summary"]["count_mismatch"] is True
    assert reference_audit["issues"][0]["code"] == "reference_count_mismatch"
    reference_gate = next(gate for gate in readiness["gates"] if gate["id"] == "references")
    assert reference_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["reference_audit_count_mismatch"] is True
    assert "Reference Count Mismatch" in reference_html


def test_reference_audit_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese reference audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n主要结局OR 0.66（95% CI 0.53至0.82）。",
            "## 引言\n既往证据来自随机试验［1］。",
            "## 方法\n检索并整理参考文献。",
            "## 结果\n主要分析纳入2项研究［1-2］。",
            "## 讨论\n参考文献清单需要与BibTeX一致。",
            "## 参考文献\n［1］ Smith J. Trial report.\n［2］ Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "output_language": "zh", "primary_effect": {"n_studies": 2}},
        subdir="manuscript",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))
        reference_html = zf.read("review/reference_audit.html").decode("utf-8")

    assert reference_audit["language"] == "zh"
    assert reference_audit["entries"][0]["source_type"] == "journal_article"
    assert '<html lang="zh">' in reference_html
    assert "MetaAgent 参考文献审计" in reference_html
    assert "核对稿件参考文献清单与打包的 BibTeX 文献库是否一致" in reference_html
    assert "稿件参考文献" in reference_html
    assert "BibTeX条目" in reference_html
    assert "数量不一致" in reference_html
    assert "缺失标识符" in reference_html
    assert "缺失期刊" in reference_html
    assert "缺失卷/页码" in reference_html
    assert "参考文献数量不一致" in reference_html
    assert "问题" in reference_html
    assert "严重性" in reference_html
    assert "参考文献条目" in reference_html
    assert "期刊论文" in reference_html
    assert "journal_article" not in reference_html
    assert "稿件编号参考文献清单与 references.bib 的条目数不同" in reference_html
    assert "MetaAgent Reference Audit" not in reference_html
    assert "Reference Count Mismatch" not in reference_html


def test_reference_audit_omits_chinese_count_mismatch_banner_when_counts_match(tmp_path: Path) -> None:
    project = Project("chinese reference audit match package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n主要结局OR 0.66。",
            "## 引言\n既往证据来自随机试验［1］。",
            "## 参考文献\n［1］ Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_facts.json", {"output_language": "zh"}, subdir="manuscript")
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))
        reference_html = zf.read("review/reference_audit.html").decode("utf-8")

    assert reference_audit["summary"]["count_mismatch"] is False
    assert "参考文献数量不一致" not in reference_html
    assert "稿件编号参考文献清单与 references.bib 的条目数不同" not in reference_html


def test_submission_readiness_blocks_journal_references_missing_journal_title(tmp_path: Path) -> None:
    project = Project("reference completeness package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Reference completeness manuscript",
            "## Abstract\nPrimary outcome HR 0.81.",
            "## Introduction\nThe evidence base includes randomized trials [1].",
            "## Methods\nFull search query: heart failure AND SGLT2.",
            "## Results\nThe primary meta-analysis included 2 trials [1].",
            "## Discussion\nThe pooled estimate is interpreted against the source trials [1].",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Anker SD, Butler J. Empagliflozin in Heart Failure with a Preserved Ejection Fraction. 2021. doi: 10.1056/NEJMoa2107038",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {"report_type": "meta", "primary_effect": {"n_studies": 2}},
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_text(
        "references.bib",
        "@article{anker2021,\n"
        "  title = {Empagliflozin in Heart Failure with a Preserved Ejection Fraction},\n"
        "  author = {Anker SD and Butler J},\n"
        "  journal = {},\n"
        "  year = {2021},\n"
        "  doi = {10.1056/NEJMoa2107038},\n"
        "  pmid = {34449189},\n"
        "}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    reference_gate = next(gate for gate in readiness["gates"] if gate["id"] == "references")
    assert reference_audit["summary"]["entries_missing_journal"] == 1
    assert reference_audit["issues"][0]["code"] == "reference_missing_journal"
    assert reference_gate["status"] == "fail"
    assert manifest["review"]["reference_audit_missing_journal"] == 1


def test_reference_audit_does_not_require_volume_pages_for_registry_mirror_urls(tmp_path: Path) -> None:
    project = Project("registry mirror reference package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Registry reference manuscript",
            "## Abstract\nA registry source was used.",
            "## Results\nThe registry source is cited [1].",
            "## References\n[1] Smart Patients registry mirror. Hydrocortisone for COVID-19 and Severe Hypoxia. 2020. https://www.smartpatients.com/trials/NCT04348305",
        ]),
        subdir="manuscript",
    )
    project.save_text(
        "references.bib",
        "@article{registry2020,\n"
        "  title = {Hydrocortisone for COVID-19 and Severe Hypoxia},\n"
        "  author = {},\n"
        "  journal = {Smart Patients registry mirror},\n"
        "  year = {2020},\n"
        "  url = {https://www.smartpatients.com/trials/NCT04348305},\n"
        "}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))

    assert reference_audit["entries"][0]["source_type"] in {"trial_registry", "web_source"}
    assert reference_audit["summary"]["entries_missing_volume_or_pages"] == 0
    assert not any(issue["code"] == "reference_missing_volume_or_pages" for issue in reference_audit["issues"])


def test_reference_audit_does_not_require_volume_pages_for_handbook_urls(tmp_path: Path) -> None:
    project = Project("handbook reference package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Handbook reference manuscript",
            "## Methods\nMethods followed handbook guidance [1].",
            "## References\n[1] Higgins JPT, Thomas J, Chandler J, Cumpston M, Li T, Page MJ. Cochrane Handbook for Systematic Reviews of Interventions. 2023. https://training.cochrane.org/handbook",
        ]),
        subdir="manuscript",
    )
    project.save_text(
        "references.bib",
        "@article{jpt2023,\n"
        "  title = {Cochrane Handbook for Systematic Reviews of Interventions},\n"
        "  author = {Higgins JPT and Thomas J and Chandler J and Cumpston M and Li T and Page MJ},\n"
        "  journal = {Cochrane},\n"
        "  year = {2023},\n"
        "  volume = {},\n"
        "  pages = {},\n"
        "  url = {https://training.cochrane.org/handbook},\n"
        "}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))

    assert reference_audit["entries"][0]["source_type"] == "methods_handbook"
    assert reference_audit["summary"]["entries_missing_volume_or_pages"] == 0
    assert not any(issue["code"] == "reference_missing_volume_or_pages" for issue in reference_audit["issues"])


def test_submission_readiness_blocks_full_length_meta_without_main_text_citations(tmp_path: Path) -> None:
    project = Project("citation audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    long_text = " ".join(["citationaudit"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full citation audit manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text}",
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables."
            ),
            "## Results\nThe primary meta-analysis included 2 trials totaling 60 participants.",
            "## Discussion\nThe result is interpreted in relation to the selected trials.",
            "## Tables\n### Table 1. Trial-level effects\n| Study | OR |\n|---|---:|\n| Trial 1 | 0.50 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"},
        {"row_id": "S2:0", "study_id": "S2", "outcome_name": "mortality"},
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/smith}}\n"
        "@article{jones2024,title={Trial report},doi={10.1000/jones}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        citation_audit = json.loads(zf.read("review/citation_audit.json"))
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    citation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "citation_coverage")
    assert citation_audit["summary"]["reference_entries"] == 2
    assert citation_audit["summary"]["main_text_inline_citations"] == 0
    assert citation_audit["issues"][0]["code"] == "section_citations_missing"
    assert citation_gate["status"] == "fail"
    assert "Introduction=0" in citation_gate["detail"]
    assert readiness["status"] == "blocked"
    assert manifest["review"]["citation_audit_failed_issues"] >= 1
    assert "Citation Coverage Issue" in citation_html


def test_citation_audit_warns_when_formal_manuscript_has_sparse_references(tmp_path: Path) -> None:
    project = Project("sparse citation audit", output_dir=tmp_path / uuid4().hex)
    long_paragraph = " ".join(["background evidence sentence"] * 80)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Sparse reference manuscript",
            f"## Introduction\n{long_paragraph} [1].",
            f"## Methods\n{long_paragraph} [1].",
            f"## Results\n{long_paragraph} [1].",
            f"## Discussion\n{long_paragraph} [1].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Prior review.",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    codes = {issue["code"] for issue in audit["issues"]}
    assert audit["passed"] is True
    assert audit["summary"]["reference_entries"] == 4
    assert audit["summary"]["warning_issues"] >= 3
    assert audit["summary"]["main_text_word_count"] >= 900
    assert "insufficient_reference_count" in codes
    assert "low_unique_cited_references" in codes
    assert "low_citation_density" in codes

    package_path = create_artifact_package(project)
    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")

    assert manifest["review"]["citation_audit_warning_issues"] >= 3
    assert "Warnings:" in citation_html


def test_citation_audit_warns_when_interpretive_section_paragraphs_lack_citations(tmp_path: Path) -> None:
    project = Project("paragraph citation coverage", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 21))
    long_uncited = " ".join(["This interpretive paragraph summarizes disease burden and unresolved uncertainty"] * 12)
    long_cited = " ".join(["Existing trials and reviews motivate the present synthesis"] * 12)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Paragraph citation manuscript",
            (
                "## Introduction\n"
                f"{long_uncited}.\n\n"
                f"{long_cited} [1,2,3].\n\n"
                f"{long_uncited}."
            ),
            "## Methods\nThe review followed PRISMA and GRADE methods [4,5].",
            "## Results\nTwo randomized trials contributed source-verified outcome data [6,7].",
            (
                "## Discussion\n"
                f"{long_uncited}.\n\n"
                f"{long_cited} [8,9,10].\n\n"
                f"{long_uncited}."
            ),
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issues_by_code = {issue["code"]: issue for issue in audit["issues"]}
    assert audit["passed"] is True
    assert audit["summary"]["introduction_substantial_paragraphs"] == 3
    assert audit["summary"]["introduction_cited_substantial_paragraphs"] == 1
    assert audit["summary"]["discussion_substantial_paragraphs"] == 3
    assert audit["summary"]["discussion_cited_substantial_paragraphs"] == 1
    assert "introduction_paragraph_citation_coverage_low" in issues_by_code
    assert "discussion_paragraph_citation_coverage_low" in issues_by_code


def test_citation_audit_warns_about_mechanical_citation_density_in_discussion(tmp_path: Path) -> None:
    project = Project("mechanical citation density", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 21))
    mechanical_discussion = (
        "The pooled estimate should be interpreted against baseline risk and treatment goals [1]. "
        "Sparse trials make heterogeneity and publication-bias checks descriptive rather than definitive [2]. "
        "Safety outcomes such as volume depletion and treatment discontinuation still need separate judgment [3]. "
        "Subgroup effects require cautious interpretation because aggregate data cannot resolve individual effect modification [4]."
    )
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Mechanical citation density manuscript",
            "## Introduction\nHeart failure context is summarized with prior sources [1,2].",
            "## Methods\nThe review followed reporting and certainty methods [3,4].",
            "## Results\nTwo trials contributed to the main synthesis [5,6].",
            f"## Discussion\n{mechanical_discussion}",
            "## Conclusion\nThe result supports a cautious clinical interpretation [7].",
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "mechanical_citation_density")
    assert audit["passed"] is True
    assert audit["summary"]["mechanical_citation_density_paragraphs"] == 1
    assert issue["section"] == "Discussion"
    assert issue["citation_markers"] == ["[1]", "[2]", "[3]", "[4]"]
    assert issue["markers_per_35_text_units"] > issue["maximum_markers_per_35_text_units"]
    assert "baseline risk" in issue["evidence_excerpt"]


def test_citation_audit_warns_about_excessive_global_citation_density(tmp_path: Path) -> None:
    project = Project("excessive citation density", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 31))
    cited_methods_sentences = " ".join(
        f"Methodological detail {i} was prespecified and checked against the protocol [{(i % 30) + 1}]."
        for i in range(1, 46)
    )
    cited_results_sentences = " ".join(
        f"Result detail {i} was reported with its source estimate [{((i + 10) % 30) + 1}]."
        for i in range(1, 36)
    )
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Excessively cited manuscript",
            "## Introduction\nThe clinical context is supported by prior evidence [1,2].",
            f"## Methods\n{cited_methods_sentences}",
            f"## Results\n{cited_results_sentences}",
            "## Discussion\nThe result should be interpreted by baseline risk, endpoint composition, safety, applicability, and certainty [3].",
            "## Conclusion\nThe intervention may be useful in selected patients [4].",
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "excessive_citation_density")
    assert audit["passed"] is True
    assert audit["summary"]["excessive_citation_density"] is True
    assert audit["summary"]["citation_density_per_1000_words"] > audit["summary"]["maximum_citation_density_per_1000_words"]
    assert issue["severity"] == "warn"
    assert issue["section"] == "Main text"


def test_citation_audit_blocks_repeated_large_citation_clusters(tmp_path: Path) -> None:
    project = Project("repeated large citation cluster", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 24))
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Repeated citation cluster manuscript",
            "## Introduction\nThe clinical context is supported by background literature [1,2].",
            "## Methods\nThe review followed established reporting and certainty methods [4,6].",
            "## Results\nTwo source reports contributed the main effect estimate [8,9].",
            (
                "## Discussion\n"
                "The composite endpoint requires interpretation alongside baseline risk [3,5,7,20,23].\n\n"
                "Absolute benefit should be judged against background therapy [3,5,7,20,23].\n\n"
                "Safety interpretation should stay separate from efficacy interpretation [3,5,7,20,23]."
            ),
            "## Conclusion\nThe evidence supports cautious clinical use [10].",
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "repeated_large_citation_cluster")
    assert audit["passed"] is False
    assert issue["severity"] == "fail"
    assert audit["summary"]["failed_issues"] >= 1
    assert audit["summary"]["repeated_large_citation_clusters"] == 1
    assert issue["citation_numbers"] == [3, 5, 7, 20, 23]
    assert issue["occurrences"] == 3
    assert issue["sections"] == ["Discussion"]


def test_citation_audit_warns_when_meta_publication_reference_depth_is_low(tmp_path: Path) -> None:
    project = Project("publication reference depth", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 14))
    cited_numbers = ", ".join(str(i) for i in range(1, 14))
    paragraph = " ".join(["Formal submission text with topic-specific citations"] * 90)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Publication reference depth manuscript",
            f"## Introduction\n{paragraph} [1,2,3,4,5].",
            f"## Methods\n{paragraph} [6,7,8,9].",
            f"## Results\n{paragraph} [10,11,12,13].",
            f"## Discussion\n{paragraph} [{cited_numbers}].",
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1"},
                    {"row_id": "S2:0", "study_id": "S2"},
                ],
            },
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/citation_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    issue_codes = {issue["code"] for issue in audit["issues"]}
    citation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "citation_coverage")
    assert audit["summary"]["reference_entries"] == 13
    # Two pooled studies: the general floor plus one per included study.
    assert audit["summary"]["publication_minimum_reference_entries"] == 14
    assert "publication_reference_count_below_target" in issue_codes
    assert citation_gate["status"] == "warn"
    assert "publication_min_references=14" in citation_gate["detail"]


def test_citation_audit_warns_when_formal_long_draft_lacks_publication_reference_depth_without_facts(tmp_path: Path) -> None:
    project = Project("formal citation depth without facts", output_dir=tmp_path / uuid4().hex)
    references = "\n".join(f"[{i}] Reference {i}." for i in range(1, 15))
    paragraph = " ".join(["Formal submission text with source anchored clinical context"] * 330)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Formal citation depth manuscript",
            "## Abstract\nThe abstract summarizes a source-anchored meta-analysis.",
            (
                "## Introduction\n"
                f"{paragraph} [1,2].\n\n"
                f"{paragraph} [3,4]."
            ),
            (
                "## Methods\n"
                f"{paragraph} [5,6].\n\n"
                f"{paragraph} [7,8]."
            ),
            f"## Results\n{paragraph} [9,10].",
            (
                "## Discussion\n"
                f"{paragraph} [11,12].\n\n"
                f"{paragraph} [13,14]."
            ),
            "## Conclusion\nThe findings remain source anchored [13,14].",
            f"## References\n{references}",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue_codes = {issue["code"] for issue in audit["issues"]}
    assert audit["summary"]["main_text_word_count"] >= 6000
    assert audit["summary"]["reference_entries"] == 14
    assert audit["summary"]["publication_reference_depth_required"] is True
    assert audit["summary"]["publication_minimum_reference_entries"] == 20
    assert "publication_reference_count_below_target" in issue_codes


def test_citation_audit_warns_when_section_citations_use_wrong_source_types(tmp_path: Path) -> None:
    project = Project("section citation source audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Source-aware citation manuscript",
            "## Introduction\nThe condition has a substantial clinical burden and prior reviews remain inconclusive [1].",
            "## Methods\nThe review followed a predefined systematic-review protocol and assessed certainty [1].",
            "## Results\nTwo randomized trials contributed source-verified primary outcome data [1].",
            "## Discussion\nThe findings should be interpreted alongside existing clinical guidance [1].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Heart failure clinical guideline.",
            "[5] Prior systematic review.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "[2]", "source_type": "reporting_guideline"},
                {"study_id": "methodology:grade", "citation": "[3]", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "evidence:guideline", "citation": "[4]", "source_type": "guideline"},
                {"study_id": "evidence:review", "citation": "[5]", "source_type": "prior_review"},
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    issues_by_code = {issue["code"]: issue for issue in audit["issues"]}
    assert audit["passed"] is True
    assert "introduction_background_citations_missing" in issues_by_code
    assert "methods_methodology_citations_missing" in issues_by_code
    assert audit["summary"]["introduction_background_inline_citations"] == 0
    assert audit["summary"]["methods_methodology_inline_citations"] == 0
    assert issues_by_code["introduction_background_citations_missing"]["recommended_citations"] == [4, 5]
    assert issues_by_code["methods_methodology_citations_missing"]["recommended_citations"] == [2, 3]


def test_citation_audit_warns_when_numeric_effect_sentence_cites_only_methodology_source(tmp_path: Path) -> None:
    project = Project("numeric effect source citation audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Numeric source-aware citation manuscript",
            "## Introduction\nBackground rationale is supported by current guidance [4].",
            "## Methods\nThe review followed PRISMA and assessed certainty [2,3].",
            "## Results\nThe pooled HR was 0.81 (95% CI 0.74 to 0.88) [3].",
            "## Discussion\nThe result was interpreted with certainty guidance [3].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Heart failure clinical guideline.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "trial:primary", "citation": "[1]", "source_type": "trial_report"},
                {"study_id": "evidence:guideline", "citation": "[4]", "source_type": "guideline"},
            ]
        },
        subdir="search",
    )
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "[2]", "source_type": "reporting_guideline"},
                {"study_id": "methodology:grade", "citation": "[3]", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "numeric_effect_claim_lacks_source_citation")
    assert audit["summary"]["numeric_effect_claims_without_source_citations"] == 1
    assert issue["section"] == "Results"
    assert issue["existing_citations"] == [3]
    assert issue["recommended_citations"] == [1]
    assert "pooled HR was 0.81" in issue["evidence_excerpt"]

    package_path = create_artifact_package(project)
    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        html = zf.read("review/citation_audit.html").decode("utf-8")
    assert manifest["review"]["citation_audit_numeric_effect_claims_without_source_citations"] == 1
    assert "numeric_effect_claim_lacks_source_citation" in html
    assert "Recommended citations: [1]" in html


def test_citation_audit_uses_reference_list_trial_entries_when_evidence_context_is_missing(tmp_path: Path) -> None:
    project = Project("numeric effect bibliography fallback audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Numeric source-aware citation manuscript",
            "## Introduction\nBackground rationale is supported by current guidance [4].",
            "## Methods\nThe review followed PRISMA and assessed certainty [2,3].",
            "## Results\nThe pooled HR was 0.81 (95% CI 0.74 to 0.88) [3].",
            "## Discussion\nThe result was interpreted with certainty guidance [3].",
            "## References",
            "[1] Smith J. Randomized clinical trial report of dapagliflozin in heart failure.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Heart failure clinical guideline.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "[2]", "source_type": "reporting_guideline"},
                {"study_id": "methodology:grade", "citation": "[3]", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "numeric_effect_claim_lacks_source_citation")
    assert audit["summary"]["numeric_effect_claims_without_source_citations"] == 1
    assert issue["existing_citations"] == [3]
    assert issue["recommended_citations"] == [1]


def test_citation_audit_warns_when_contextual_section_citation_depth_is_low(tmp_path: Path) -> None:
    project = Project("section citation depth audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Source-depth citation manuscript",
            "## Introduction\nThe condition has a substantial clinical burden and prior reviews remain inconclusive [4].",
            "## Methods\nThe review followed a predefined systematic-review protocol and assessed certainty [2].",
            "## Results\nTwo randomized trials contributed source-verified primary outcome data [1].",
            "## Discussion\nThe findings should be interpreted alongside existing clinical guidance and certainty limits [4].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] Cochrane Handbook.",
            "[4] Heart failure clinical guideline.",
            "[5] Prior systematic review.",
            "[6] GRADE guidance.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "[2]", "source_type": "reporting_guideline"},
                {"study_id": "methodology:cochrane", "citation": "[3]", "source_type": "methods_handbook"},
                {"study_id": "methodology:grade", "citation": "[6]", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "evidence:guideline", "citation": "[4]", "source_type": "guideline"},
                {"study_id": "evidence:review", "citation": "[5]", "source_type": "prior_review"},
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    issues_by_code = {issue["code"]: issue for issue in audit["issues"]}
    assert audit["passed"] is True
    assert audit["summary"]["introduction_background_inline_citations"] == 1
    assert audit["summary"]["methods_methodology_inline_citations"] == 1
    assert audit["summary"]["discussion_context_inline_citations"] == 1
    assert audit["summary"]["minimum_introduction_background_citations"] == 2
    assert audit["summary"]["minimum_methods_methodology_citations"] == 2
    assert audit["summary"]["minimum_discussion_context_citations"] == 2
    assert issues_by_code["introduction_background_citation_count_low"]["recommended_citations"] == [5]
    assert issues_by_code["methods_methodology_citation_count_low"]["recommended_citations"] == [3, 6]
    assert issues_by_code["discussion_context_citation_count_low"]["recommended_citations"] == [5, 6]


def test_citation_audit_html_surfaces_recommended_context_citations(tmp_path: Path) -> None:
    project = Project("section citation source html audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Source-aware citation manuscript",
            "## Introduction\nThe condition has a substantial clinical burden and prior reviews remain inconclusive [1].",
            "## Methods\nThe review followed a predefined systematic-review protocol and assessed certainty [1].",
            "## Results\nTwo randomized trials contributed source-verified primary outcome data [1].",
            "## Discussion\nThe findings should be interpreted alongside existing clinical guidance [1].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Heart failure clinical guideline.",
            "[5] Prior systematic review.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "[2]", "source_type": "reporting_guideline"},
                {"study_id": "methodology:grade", "citation": "[3]", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "evidence:guideline", "citation": "[4]", "source_type": "guideline"},
                {"study_id": "evidence:review", "citation": "[5]", "source_type": "prior_review"},
            ]
        },
        subdir="search",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")

    assert "Recommended citations: [4], [5]" in citation_html
    assert "Recommended citations: [2], [3]" in citation_html


def test_chinese_citation_audit_html_surfaces_recommended_context_citations(tmp_path: Path) -> None:
    project = Project("section citation source zh html audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 引言\n疾病负担和既往证据仍需结合背景来源说明［1］。",
            "## 方法\n本研究遵循系统综述和证据确定性评价流程［1］。",
            "## 结果\n两项随机试验贡献了主要结局数据［1］。",
            "## 讨论\n结果应结合既往指南和综述背景解释［1］。",
            "## 参考文献",
            "［1］ Trial report.",
            "［2］ PRISMA 2020 statement.",
            "［3］ GRADE guidance.",
            "［4］ Heart failure clinical guideline.",
            "［5］ Prior systematic review.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "methodology_context.json",
        {
            "references": [
                {"study_id": "methodology:prisma", "citation": "［2］", "source_type": "reporting_guideline"},
                {"study_id": "methodology:grade", "citation": "［3］", "source_type": "certainty_framework"},
            ]
        },
        subdir="search",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "evidence:guideline", "citation": "［4］", "source_type": "guideline"},
                {"study_id": "evidence:review", "citation": "［5］", "source_type": "prior_review"},
            ]
        },
        subdir="search",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")

    assert "建议补充引用：[4], [5]" in citation_html
    assert "建议补充引用：[2], [3]" in citation_html
    assert "Recommended citations" not in citation_html


def test_chinese_citation_audit_html_localizes_issue_labels(tmp_path: Path) -> None:
    project = Project("citation issue zh labels", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 引言\n疾病负担需要结合背景来源说明。",
            "## 方法\n本研究遵循系统综述流程［1］。",
            "## 结果\n两项随机试验贡献主要结局数据［1］。",
            "## 讨论\n结果应结合既往指南解释［1］。",
            "## 参考文献",
            "［1］ Trial report.",
            "［2］ Prior systematic review.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        citation_audit = json.loads(zf.read("review/citation_audit.json"))
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")

    assert any(issue.get("code") == "section_citations_missing" for issue in citation_audit["issues"])
    assert "失败" in citation_html
    assert "缺少章节引用" in citation_html
    assert ">fail</td>" not in citation_html
    assert ">section_citations_missing</td>" not in citation_html


def test_citation_audit_counts_evimed_guide_and_paper_references_as_background_context(tmp_path: Path) -> None:
    project = Project("evimed source type citation audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Evimed citation manuscript",
            "## Introduction\nThe condition and prior evidence are summarized with Evimed context [4,5].",
            "## Methods\nThe review followed PRISMA methods [1].",
            "## Results\nTwo randomized trials contributed outcome data [1].",
            "## Discussion\nThe findings were interpreted alongside existing context [4,5].",
            "## References",
            "[1] Trial report.",
            "[2] PRISMA 2020 statement.",
            "[3] GRADE guidance.",
            "[4] Evimed guideline.",
            "[5] Evimed systematic review paper.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {"study_id": "evimed:guide:g1", "citation": "[4]", "source_type": "guide"},
                {"study_id": "evimed:paper:p1", "citation": "[5]", "source_type": "paper"},
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    assert audit["summary"]["introduction_background_inline_citations"] == 2
    assert audit["summary"]["discussion_context_inline_citations"] == 2
    assert "introduction_background_citations_missing" not in {issue["code"] for issue in audit["issues"]}


def test_citation_audit_warns_when_background_reference_is_topic_mismatched(tmp_path: Path) -> None:
    project = Project("citation background relevance audit", output_dir=tmp_path / uuid4().hex)
    project.save_json(
        "protocol.json",
        {
            "research_question": (
                "In adults with heart failure with preserved ejection fraction, do sodium-glucose cotransporter-2 inhibitors "
                "reduce cardiovascular death or hospitalization for heart failure?"
            ),
            "pico": {
                "population": "Adults with HFpEF or HFmrEF",
                "intervention": "Sodium-glucose cotransporter-2 inhibitors including dapagliflozin and empagliflozin",
                "comparator": "Placebo",
                "outcome_primary": "Cardiovascular death or hospitalization for heart failure",
            },
        },
    )
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Topic-mismatch citation manuscript",
            "## Introduction\nExisting SGLT2 and HFpEF background evidence is summarized here [2,3].",
            "## Methods\nThe review followed PRISMA methods [1].",
            "## Results\nTwo randomized trials contributed outcome data [1].",
            "## Discussion\nThe findings were interpreted alongside heart failure guidance [3].",
            "## References",
            "[1] Trial report.",
            "[2] SGLT2 inhibitors for type 2 diabetes: network meta-analysis.",
            "[3] Heart failure with preserved ejection fraction: everything the clinician needs to know.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {
                    "study_id": "pubmed_background:diabetes",
                    "citation": "[2]",
                    "source_type": "pubmed_background",
                    "paper": {"title": "SGLT2 inhibitors for type 2 diabetes: network meta-analysis"},
                },
                {
                    "study_id": "pubmed_background:hfpef",
                    "citation": "[3]",
                    "source_type": "pubmed_background",
                    "paper": {"title": "Heart failure with preserved ejection fraction: everything the clinician needs to know"},
                },
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "background_reference_topic_mismatch")
    assert audit["passed"] is True
    assert audit["summary"]["background_reference_topic_mismatch_count"] == 1
    assert issue["severity"] == "warn"
    assert issue["citation_numbers"] == [2]
    assert "type 2 diabetes" in issue["titles"][0].lower()


def test_citation_audit_topic_mismatch_recognizes_spelled_out_sglt2_protocol(tmp_path: Path) -> None:
    project = Project("citation spelled-out sglt2 relevance audit", output_dir=tmp_path / uuid4().hex)
    project.save_json(
        "protocol.json",
        {
            "research_question": (
                "In adults with heart failure with preserved ejection fraction, do sodium-glucose "
                "cotransporter-2 inhibitors reduce cardiovascular death or hospitalization for heart failure?"
            ),
            "pico": {
                "population": "Adults with HFpEF or HFmrEF",
                "intervention": "Sodium-glucose cotransporter-2 inhibitors",
                "comparator": "Placebo",
                "outcome_primary": "Cardiovascular death or hospitalization for heart failure",
            },
        },
    )
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Topic-mismatch citation manuscript",
            "## Introduction\nExisting background evidence is summarized here [2].",
            "## Methods\nThe review followed PRISMA methods [1].",
            "## Results\nTwo randomized trials contributed outcome data [1].",
            "## Discussion\nThe findings were interpreted alongside background evidence [2].",
            "## References",
            "[1] Trial report.",
            "[2] Sodium-glucose cotransporter-2 inhibitors for type 2 diabetes: network meta-analysis.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "evidence_context.json",
        {
            "references": [
                {
                    "study_id": "pubmed_background:diabetes",
                    "citation": "[2]",
                    "source_type": "pubmed_background",
                    "paper": {
                        "title": "Sodium-glucose cotransporter-2 inhibitors for type 2 diabetes: network meta-analysis"
                    },
                },
            ]
        },
        subdir="search",
    )

    audit = _build_citation_audit_review(project)

    assert any(issue["code"] == "background_reference_topic_mismatch" for issue in audit["issues"])


def test_submission_readiness_warns_when_formal_manuscript_has_sparse_citations(tmp_path: Path) -> None:
    project = Project("submission readiness sparse citations", output_dir=tmp_path / uuid4().hex)
    long_paragraph = " ".join(["background evidence sentence"] * 80)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Sparse citation manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_paragraph} [1].",
            f"## Methods\nPubMed Boolean strategy: treatment AND mortality. {long_paragraph} [1].",
            f"## Results\nTwo source-verified trial rows are shown. {long_paragraph} [1].",
            f"## Discussion\n{long_paragraph} [1].",
            "## Figures\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Trial report.\n[2] PRISMA 2020 statement.\n[3] GRADE guidance.\n[4] Prior review.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "writing_constraints": {"publication_min_main_words": 500},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"},
                    {"row_id": "S2:0", "study_id": "S2", "outcome_name": "mortality"},
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 1/10 vs 2/10.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                },
                {
                    "row_id": "S2:0",
                    "study_id": "S2",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 3/20 vs 5/20.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                },
            ],
        },
        subdir="extraction",
    )
    project.save_text(
        "references.bib",
        "\n".join([
            "@article{trial,title={Trial report},journal={Journal},year={2024},doi={10.1000/trial}}",
            "@article{prisma,title={PRISMA 2020 statement},journal={BMJ},year={2021},doi={10.1136/bmj.n71}}",
            "@article{grade,title={GRADE guidance},journal={J Clin Epidemiol},year={2011},doi={10.1016/j.jclinepi.2010.04.026}}",
            "@article{review,title={Prior review},journal={Journal},year={2023},doi={10.1000/review}}",
        ])
    )
    project.save_text("search_query.txt", "treatment AND mortality")
    project.save_text("search_strategy_report.txt", "PubMed Boolean strategy: treatment AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    citation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "citation_coverage")
    assert citation_gate["status"] == "warn"
    assert "warning_issues=" in citation_gate["detail"]
    assert "density=" in citation_gate["detail"]


def test_reference_and_citation_audits_recognize_bibliography_heading(tmp_path: Path) -> None:
    project = Project("bibliography heading citation audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Bibliography heading manuscript",
            "## Introduction\nBackground statement [1].",
            "## Methods\nWe followed reporting guidance [2].",
            "## Results\nThe pooled estimate favored treatment [1].",
            "## Discussion\nThe result was interpreted with certainty guidance [2].",
            "## Bibliography\n[1] Smith J. Trial report.\n[2] Page MJ. PRISMA 2020 statement.",
        ]),
        subdir="manuscript",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},journal={Example Journal},year={2024},volume={1},pages={1-9},doi={10.1000/smith}}\n"
        "@article{page2021,title={PRISMA 2020 statement},journal={BMJ},year={2021},volume={372},pages={n71},doi={10.1136/bmj.n71}}\n",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        reference_audit = json.loads(zf.read("review/reference_audit.json"))
        citation_audit = json.loads(zf.read("review/citation_audit.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert reference_audit["summary"]["manuscript_references"] == 2
    assert reference_audit["summary"]["count_mismatch"] is False
    assert citation_audit["summary"]["reference_entries"] == 2
    assert citation_audit["summary"]["main_text_inline_citations"] == 4
    assert citation_audit["summary"]["unique_cited_reference_numbers"] == 2
    assert manifest["manuscript"]["has_references_section"] is True


def test_citation_audit_recognizes_full_width_chinese_reference_markers(tmp_path: Path) -> None:
    project = Project("full width citation audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n主要结局提示干预可能降低死亡风险［1］。",
            "## 引言\n既往随机试验和综述提示该问题仍有不确定性［1，2］。",
            "## 方法\n根据PRISMA 2020构建检索、筛选和数据提取流程［2］。",
            "## 结果\n主要Meta分析纳入2项研究，总计60名参与者［1-2］。",
            "## 讨论\n结果需要结合证据确定性和外部可推广性解释［1，2］。",
            "## 参考文献\n［1］ Smith J. Trial report.\n［2］ Page MJ. PRISMA 2020 statement.",
        ]),
        subdir="manuscript",
    )

    citation_audit = _build_citation_audit_review(project)
    summary = _manuscript_content_summary(project.load_text("draft.md", subdir="manuscript"))

    assert citation_audit is not None
    assert citation_audit["passed"] is True
    assert citation_audit["summary"]["reference_entries"] == 2
    assert citation_audit["summary"]["main_text_inline_citations"] == 8
    assert citation_audit["summary"]["unique_cited_reference_numbers"] == 2
    assert citation_audit["summary"]["introduction_inline_citations"] == 2
    assert citation_audit["summary"]["methods_inline_citations"] == 1
    assert citation_audit["summary"]["results_inline_citations"] == 2
    assert citation_audit["summary"]["discussion_inline_citations"] == 2
    assert summary["reference_count"] == 2


def test_citation_audit_warns_when_chinese_draft_uses_ascii_numeric_markers_outside_code(tmp_path: Path) -> None:
    project = Project("mixed chinese citation style audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 引言\n既往研究提示该问题有临床意义 [1,2]。",
            "## 方法\n检索式保留在代码块中：\n```text\n(\"HFpEF\"[tiab] AND \"SGLT2\"[tiab])\n```",
            "## 结果\n主要Meta分析纳入2项研究［1，2］。",
            "## 讨论\n结果需要结合证据确定性解释［2］。",
            "## 参考文献\n[1] Smith J. Trial report.\n[2] Page MJ. PRISMA 2020 statement.",
        ]),
        subdir="manuscript",
    )

    citation_audit = _build_citation_audit_review(project)

    assert citation_audit is not None
    issue = next(
        item for item in citation_audit["issues"]
        if item["code"] == "chinese_ascii_numeric_citation_marker_style"
    )
    assert issue["severity"] == "warn"
    assert issue["ascii_numeric_citation_markers"] == 4
    assert citation_audit["summary"]["ascii_numeric_citation_markers_in_chinese"] == 4
    assert citation_audit["summary"]["warning_issues"] >= 1


def test_exported_quality_reviews_localize_chinese_citation_and_polish_html(tmp_path: Path) -> None:
    project = Project("chinese localized quality reviews", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 引言\n背景证据需要引用［1］。本研究显示疗效降低。本研究显示疗效稳定。本研究显示疗效可信。本研究显示疗效相关。",
            "## 方法\n根据PRISMA 2020构建检索、筛选和数据提取流程。本研究显示疗效降低。本研究显示疗效稳定。本研究显示疗效可信。本研究显示疗效相关。",
            "## 结果\n主要Meta分析纳入2项研究［1］。本研究显示疗效降低。本研究显示疗效稳定。本研究显示疗效可信。本研究显示疗效相关。",
            "## 讨论\n结果结合GRADE证据确定性解释［1］。本研究显示疗效降低。本研究显示疗效稳定。本研究显示疗效可信。本研究显示疗效相关。",
            "## 参考文献\n［1］ Trial report.\n［2］ PRISMA 2020 statement.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "rewrite_scope": "targeted",
            "accepted_chunks": 0,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "before": {"language": "zh", "ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"language": "zh", "ai_style_signal": {"score": 1, "issues": [{"code": "low_sentence_length_variation"}]}},
            "issues": [
                {
                    "code": "citations_changed",
                    "heading": "讨论",
                    "message": "Polish rewrite changed citation markers.",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        citation_audit = json.loads(zf.read("review/citation_audit.json"))
        citation_html = zf.read("review/citation_audit.html").decode("utf-8")
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        readiness_html = zf.read("review/submission_readiness_review.html").decode("utf-8")

    assert citation_audit["language"] == "zh"
    assert '<html lang="zh">' in citation_html
    assert "引用覆盖审计" in citation_html
    assert "正式稿正文应引用来源报告" in citation_html
    assert "方法部分" in citation_html
    assert '<html lang="zh">' in polish_html
    assert "稿件润色审计" in polish_html
    assert "事实保护闸" in polish_html
    assert "剩余风格信号" in polish_html
    assert readiness["language"] == "zh"
    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")
    citation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "citation_coverage")
    assert polish_gate["label_localized"] == "稿件润色保护闸"
    assert citation_gate["label_localized"] == "正文引用覆盖"
    assert any("人工确认" in action for action in polish_gate["next_actions"])
    assert not any("human confirms" in action for action in polish_gate["next_actions"])
    assert '<html lang="zh">' in readiness_html
    assert "MetaAgent 提交就绪性" in readiness_html
    assert "下一步动作" in readiness_html
    assert "事实保护问题=1" in readiness_html
    assert "剩余风格信号=2" in readiness_html
    assert "复核队列状态=需人工复核" in readiness_html
    assert "人工复核项=3" in readiness_html
    assert "fact_guard_issues" not in readiness_html
    assert "remaining_style_issues" not in readiness_html
    assert "review_queue_status" not in readiness_html
    assert "human_review_required" not in readiness_html
    assert "Submission Readiness" not in readiness_html
    assert "Next actions" not in readiness_html


def test_manuscript_polish_audit_localizes_chinese_proofreading_headers(tmp_path: Path) -> None:
    project = Project("chinese proofreading polish review", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 引言\n既往证据需要结合正式审校意见解释［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "rewrite_scope": "targeted",
            "accepted_chunks": 0,
            "rejected_chunks": 0,
            "unchanged_chunks": 1,
            "before": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [],
            "proofreading": {
                "enabled": True,
                "status": "ok",
                "provider": "languagetool",
                "language_code": "zh-CN",
                "issue_count": 1,
                "issues": [
                    {
                        "rule_id": "ZH_STYLE",
                        "category": "STYLE",
                        "message": "建议调整措辞。",
                        "replacements": ["更正式的措辞"],
                    }
                ],
            },
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    assert '<html lang="zh">' in polish_html
    assert "<th>规则</th>" in polish_html
    assert "<th>类别</th>" in polish_html
    assert "<th>审校信息</th>" in polish_html
    assert "<th>建议</th>" in polish_html
    assert "<th>Rule</th>" not in polish_html
    assert "<th>Category</th>" not in polish_html
    assert "<th>Message</th>" not in polish_html
    assert "<th>Suggestion</th>" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_review_actions(tmp_path: Path) -> None:
    project = Project("chinese polish review actions", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 讨论\n合并效应提示治疗可能降低风险［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "accepted_chunks": 1,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "accepted_edit_count": 1,
            "accepted_edits": [
                {
                    "heading": "讨论",
                    "original_text": "需要指出的是，合并效应提示治疗可能降低风险［1］。",
                    "candidate_text": "合并效应提示治疗可能降低风险［1］。",
                    "diff": "--- original\n+++ polished\n-需要指出的是，合并效应提示治疗可能降低风险［1］。\n+合并效应提示治疗可能降低风险［1］。",
                    "review_action": "accepted_fact_preserving_polish",
                }
            ],
            "before": {"language": "zh", "ai_style_signal": {"score": 1, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [
                {
                    "code": "citations_changed",
                    "heading": "讨论",
                    "message": "Polish rewrite changed citation markers.",
                    "original_text": "合并效应提示治疗可能降低风险［1］。",
                    "candidate_text": "合并效应提示治疗可能降低风险。",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    assert "已接受事实保护润色" in polish_html
    assert "需人工复核" in polish_html
    assert "accepted_fact_preserving_polish" not in polish_html
    assert "manual_review_required" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_study_design_rejection(tmp_path: Path) -> None:
    project = Project("chinese polish study design rejection", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 方法\n符合条件的证据来自随机对照试验，并采用盲法结局判定［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "accepted_chunks": 0,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "before": {"language": "zh", "ai_style_signal": {"score": 1, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [
                {
                    "code": "study_design_changed",
                    "heading": "方法",
                    "message": "Polish rewrite changed study design terms.",
                    "original_text": "符合条件的证据来自随机对照试验，并采用盲法结局判定［1］。",
                    "candidate_text": "符合条件的证据来自观察性队列研究，并采用盲法结局判定［1］。",
                    "original_study_design_terms": ["randomized_trial", "blinded"],
                    "candidate_study_design_terms": ["observational_study", "cohort_study", "blinded"],
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    rejected = review["rejected_edits"][0]
    assert rejected["original_study_design_terms"] == ["randomized_trial", "blinded"]
    assert rejected["candidate_study_design_terms"] == ["observational_study", "cohort_study", "blinded"]
    assert "研究设计术语变更" in polish_html
    assert "润色候选修改了研究设计术语，需人工复核。" in polish_html
    assert "study_design_changed" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_language_change_rejection(tmp_path: Path) -> None:
    project = Project("chinese polish language rejection", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 方法\n符合条件的证据来自随机对照试验，并重复提取预先指定的结局［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "accepted_chunks": 0,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "before": {"language": "zh", "ai_style_signal": {"score": 1, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [
                {
                    "code": "language_changed",
                    "heading": "方法",
                    "message": "Polish rewrite changed the manuscript output language.",
                    "original_text": "符合条件的证据来自随机对照试验，并重复提取预先指定的结局［1］。",
                    "candidate_text": (
                        "Eligible evidence came from randomized trials, and prespecified "
                        "outcomes were extracted in duplicate [1]."
                    ),
                    "original_language": "zh",
                    "candidate_language": "en",
                    "original_language_counts": {"language": "zh", "cjk_chars": 32, "latin_words": 0},
                    "candidate_language_counts": {"language": "en", "cjk_chars": 0, "latin_words": 13},
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review = json.loads(zf.read("review/manuscript_polish_audit.json"))
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    rejected = review["rejected_edits"][0]
    assert rejected["original_language"] == "zh"
    assert rejected["candidate_language"] == "en"
    assert rejected["original_language_counts"]["language"] == "zh"
    assert rejected["candidate_language_counts"]["language"] == "en"
    assert "稿件输出语言变更" in polish_html
    assert "润色候选修改了稿件输出语言，需人工复核。" in polish_html
    assert "language_changed" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_review_queue_status(tmp_path: Path) -> None:
    project = Project("chinese polish review queue status", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 讨论\n合并效应提示治疗可能降低风险［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "accepted_chunks": 0,
            "rejected_chunks": 1,
            "unchanged_chunks": 0,
            "before": {"language": "zh", "ai_style_signal": {"score": 1, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [
                {
                    "code": "citations_changed",
                    "heading": "讨论",
                    "message": "Polish rewrite changed citation markers.",
                    "original_text": "合并效应提示治疗可能降低风险［1］。",
                    "candidate_text": "合并效应提示治疗可能降低风险。",
                    "review_action": "manual_review_required",
                }
            ],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    assert "状态: <strong>事实保护闸已拒绝候选，无需人工复核</strong>" in polish_html
    assert "polish_guard_discarded_candidates_no_review_required" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_policy_metadata(tmp_path: Path) -> None:
    project = Project("chinese polish policy metadata", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 讨论\n合并效应提示治疗可能降低风险［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "accepted_chunks": 0,
            "rejected_chunks": 0,
            "unchanged_chunks": 1,
            "before": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "style_policy": {
                "name": "MetaAgent conservative scholarly polish",
                "detector_evasion": False,
                "external_proofreader_role": "review_only",
            },
            "issues": [],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    assert "MetaAgent 保守学术润色策略" in polish_html
    assert "检测器规避：禁用" in polish_html
    assert "外部审校角色：仅复核" in polish_html
    assert "MetaAgent conservative scholarly polish" not in polish_html
    assert "review_only" not in polish_html


def test_manuscript_polish_audit_localizes_chinese_boolean_stats(tmp_path: Path) -> None:
    project = Project("chinese polish boolean stats", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 讨论\n合并效应提示治疗可能降低风险［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "zh",
            "rewrite_scope": "targeted",
            "accepted_chunks": 0,
            "rejected_chunks": 0,
            "unchanged_chunks": 1,
            "polish_budget_exhausted": False,
            "before": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "after": {"language": "zh", "ai_style_signal": {"score": 0, "issues": []}},
            "issues": [],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        polish_html = zf.read("review/manuscript_polish_audit.html").decode("utf-8")

    assert "已启用: 是" in polish_html
    assert "润色范围: 定向润色" in polish_html
    assert "审校器: 未启用" in polish_html
    assert "预算耗尽: 否" in polish_html
    assert "已启用: True" not in polish_html
    assert "润色范围: targeted" not in polish_html
    assert "审校器: none" not in polish_html
    assert "预算耗尽: False" not in polish_html


def test_extraction_review_html_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese extraction review", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n目的：评价治疗效果。",
            "## 方法\n我们复核原文来源并提取结局数据［1］。",
            "## 结果\n主要结局来自1项研究［1］。",
            "## 讨论\n需要人工复核来源片段［1］。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "grade": {
                "outcomes": [
                    {
                        "outcome_name": "28天死亡率",
                        "n_studies": 1,
                        "certainty": "低",
                        "effect_summary": "RR 0.55 (95% CI 0.30 to 0.99)",
                        "domains": [
                            {
                                "domain": "indirectness",
                                "rating": "serious",
                                "rationale": "比较组来源需要复核。",
                                "details": {
                                    "method": "rule_based_pico_directness_v1",
                                    "n_contributing": 1,
                                    "target_outcome": "28天死亡率",
                                    "protocol_primary_outcome": "死亡率",
                                    "source_verified_direct_rows": 1,
                                    "surrogate_outcome": False,
                                    "dimensions": {
                                        "comparator": {"unverified": 1, "total": 1},
                                        "design": {"non_randomized": 0, "total": 1},
                                    },
                                },
                            },
                            {
                                "domain": "imprecision",
                                "rating": "no concern",
                                "rationale": "样本量足够。",
                                "details": {},
                            },
                        ],
                    }
                ]
            },
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "28天死亡率"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "all_extractions.json",
        [
            ExtractedStudy(
                characteristics=StudyCharacteristics(
                    study_id="S1",
                    pmid="12345",
                    title="中文来源试验",
                    pdf_path="/tmp/chinese-trial.pdf",
                ),
                outcomes=[
                    OutcomeData(
                        outcome_name="28天死亡率",
                        outcome_type="dichotomous",
                        events_intervention=11,
                        total_intervention=75,
                        events_control=20,
                        total_control=73,
                        source_location="表2",
                        source_page=5,
                        source_quote="28天死亡为11/75对20/73。",
                        source_quote_match="28天死亡为11/75对20/73。",
                        source_quote_verified=True,
                        extraction_confidence="medium",
                    )
                ],
            ).model_dump()
        ],
        subdir="extraction",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "summary": {"outcomes": 1},
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "28天死亡率",
                    "outcome_type": "dichotomous",
                    "source_location": "表2",
                    "source_page": 5,
                    "source_quote": "28天死亡为11/75对20/73。",
                    "source_quote_match": "28天死亡为11/75对20/73。",
                    "source_quote_verified": True,
                    "extraction_confidence": "medium",
                    "requires_review": True,
                    "conflicts": [{"field": "total_intervention", "message": "需核对分母。"}],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "12345": {
                "full_text": "[PAGE 5]\n表2之前。28天死亡为11/75对20/73。表2之后。",
                "page_map": [{"page_number": 5, "start_char": 0, "end_char": 40}],
            }
        },
        subdir="papers",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        review_html = zf.read("review/extraction_review.html").decode("utf-8")

    assert '<html lang="zh">' in review_html
    assert "MetaAgent 数据提取复核" in review_html
    assert "请在外部使用前" in review_html
    assert "报告类型" in review_html
    assert "需要复核" in review_html
    assert "GRADE质量复核" in review_html
    assert "<td>间接性</td><td>严重</td>" in review_html
    assert "<td>不精确性</td><td>无明显问题</td>" in review_html
    assert "规则PICO直接性检查" in review_html
    assert "贡献研究数: 1" in review_html
    assert "目标结局: 28天死亡率" in review_html
    assert "比较组" in review_html
    assert "未核验: 1" in review_html
    assert "非随机研究: 0" in review_html
    assert "<td>indirectness</td>" not in review_html
    assert ">serious<" not in review_html
    assert ">no concern<" not in review_html
    assert "rule_based_pico_directness_v1" not in review_html
    assert "unverified: 1" not in review_html
    assert "non_randomized: 0" not in review_html
    assert "提取值" in review_html
    assert "来源" in review_html
    assert "信任状态" in review_html
    assert "来源定位" in review_html
    assert "原文位置" in review_html
    assert "提取置信度" in review_html
    assert "<td>提取置信度</td><td>中等</td>" in review_html
    assert "中等" in review_html
    assert "引用已核验: 是" in review_html
    assert "需人工复核；需复核" in review_html
    assert "高亮" in review_html
    assert "来源上下文" in review_html
    assert "复核决定载荷种子" in review_html
    assert "确认该行" in review_html
    assert "source location" not in review_html
    assert "extraction confidence" not in review_html
    assert "Trust status: needs_review" not in review_html
    assert "MetaAgent Extraction Review" not in review_html
    assert "Review Decision Payload Seed" not in review_html


def test_citation_audit_applies_formal_reference_threshold_to_long_chinese_text(tmp_path: Path) -> None:
    project = Project("long chinese citation audit", output_dir=tmp_path / uuid4().hex)
    long_chinese = "背景证据显示该临床问题仍存在重要不确定性。" * 180
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            f"## 引言\n{long_chinese}［1，2］。",
            f"## 方法\n{long_chinese}［2］。",
            f"## 结果\n{long_chinese}［1-2］。",
            f"## 讨论\n{long_chinese}［1，2］。",
            "## 参考文献\n［1］ Smith J. Trial report.\n［2］ Page MJ. PRISMA 2020 statement.",
        ]),
        subdir="manuscript",
    )

    citation_audit = _build_citation_audit_review(project)

    assert citation_audit is not None
    assert citation_audit["summary"]["main_text_word_count"] >= 500
    assert "insufficient_reference_count" in {issue["code"] for issue in citation_audit["issues"]}


def test_readiness_audits_recognize_chinese_table_and_figure_labels(tmp_path: Path) -> None:
    project = Project("chinese table figure audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n主要结局 OR 0.66 (95% CI 0.53 to 0.82)。",
            "## 引言\n背景证据支持开展系统评价 [1,2]。",
            (
                "## 方法\n完整检索式: (COVID-19 AND corticosteroids) AND mortality。"
                "使用2 x 2表计算log odds ratio，并采用inverse-variance fixed-effect meta-analysis。"
            ),
            "## 结果\n主要Meta分析纳入2项研究，总计60名参与者 [1,2]。图1显示森林图，表1总结研究层面效应量。",
            "## 讨论\n结果结合GRADE证据确定性解释 [2]。",
            (
                "## 表格\n### 表1. 研究层面效应量\n"
                "| 研究 | OR | 95% CI |\n|---|---:|---|\n| 研究1 | 0.50 | 0.20 to 1.20 |\n\n"
                "注：OR=优势比；CI=置信区间。"
            ),
            (
                "## 图表\n### 图1. 森林图\n"
                "![图1. 森林图](../figures/forest_plot.png)\n\n"
                "图注：森林图展示各研究效应量及其95%置信区间。"
            ),
            "## 声明\n未收集新的个体参与者数据。",
            "## 参考文献\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    cross_ref = _build_cross_reference_audit_review(project)
    figure_assets = _build_figure_audit_review(project)
    figure_legend = _build_figure_legend_audit_review(project)
    table_footnote = _build_table_footnote_audit_review(project)
    summary = _manuscript_content_summary(project.load_text("draft.md", subdir="manuscript"))

    assert cross_ref is not None
    assert cross_ref["passed"] is True
    assert cross_ref["summary"]["defined_tables"] == 1
    assert cross_ref["summary"]["defined_figures"] == 1
    assert cross_ref["summary"]["main_text_referenced_tables"] == 1
    assert cross_ref["summary"]["main_text_referenced_figures"] == 1
    assert figure_assets is not None
    assert figure_assets["passed"] is True
    assert figure_assets["summary"]["figure_headings"] == 1
    assert figure_assets["summary"]["referenced_images"] == 1
    assert figure_assets["summary"]["missing_referenced_images"] == 0
    assert figure_legend is not None
    assert figure_legend["summary"]["figures_with_legends"] == 1
    assert table_footnote is not None
    assert table_footnote["summary"]["tables_with_notes"] == 1
    assert summary["table_count"] == 1
    assert summary["figure_count"] == 1


def test_submission_readiness_blocks_unreferenced_main_figures(tmp_path: Path) -> None:
    project = Project("cross reference audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    long_text = " ".join(["crossrefaudit"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full cross-reference audit manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text} [1,2].",
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables."
            ),
            "## Results\nThe primary meta-analysis included 2 trials totaling 60 participants [1,2]. Table 1 summarizes the trial-level effects.",
            "## Discussion\nThe result is interpreted in relation to the selected trials [1,2].",
            "## Tables\n### Table 1. Trial-level effects\n| Study | OR |\n|---|---:|\n| Trial 1 | 0.50 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"},
                    {"row_id": "S2:0", "study_id": "S2", "outcome_name": "mortality"},
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/smith}}\n"
        "@article{jones2024,title={Trial report},doi={10.1000/jones}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        cross_ref = json.loads(zf.read("review/cross_reference_audit.json"))
        cross_ref_html = zf.read("review/cross_reference_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    cross_ref_gate = next(gate for gate in readiness["gates"] if gate["id"] == "cross_references")
    assert cross_ref["summary"]["defined_figures"] == 1
    assert cross_ref["summary"]["main_text_referenced_figures"] == 0
    assert cross_ref["issues"][0]["code"] == "figure_unreferenced_in_main_text"
    assert cross_ref_gate["status"] == "fail"
    assert "figures=0/1" in cross_ref_gate["detail"]
    assert readiness["status"] == "blocked"
    assert manifest["review"]["cross_reference_audit_unreferenced_figures"] == 1
    assert "Cross-Reference Issue" in cross_ref_html


def test_cross_reference_audit_recognizes_chinese_figure_and_table_ranges(tmp_path: Path) -> None:
    project = Project("chinese cross reference ranges", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 结果\n图1至图3展示研究流程、主要效应和偏倚风险。表1至表3列出研究特征、效应量和GRADE摘要。",
            "## 讨论\n结果需要结合表1至表3和图1至图3解释。",
            "## 表格",
            "### 表1. 研究特征\n| 研究 | N |\n|---|---:|\n| A | 10 |",
            "### 表2. 效应量\n| 研究 | HR |\n|---|---:|\n| A | 0.80 |",
            "### 表3. GRADE摘要\n| 结局 | 确定性 |\n|---|---|\n| 主要结局 | 高 |",
            "## 图表",
            "### 图1. PRISMA流程图\n![图1](../figures/prisma.png)",
            "### 图2. 森林图\n![图2](../figures/forest.png)",
            "### 图3. 偏倚风险概要\n![图3](../figures/rob.png)",
        ]),
        subdir="manuscript",
    )

    cross_ref = _build_cross_reference_audit_review(project)

    assert cross_ref is not None
    assert cross_ref["passed"] is True
    assert cross_ref["summary"]["main_text_referenced_tables"] == 3
    assert cross_ref["summary"]["main_text_referenced_figures"] == 3
    assert cross_ref["unreferenced_tables"] == []
    assert cross_ref["unreferenced_figures"] == []


def test_writer_backfills_missing_figure_reference_after_fact_repair_removes_text() -> None:
    manuscript = "\n\n".join([
        "# Meta-analysis",
        "## Results\nFigure 1 shows the study flow. Figures 3 and Figure 4 show sensitivity and risk-of-bias information.",
        "## Figures",
        "### Figure 1. PRISMA flow diagram\n![Figure 1](../figures/prisma_diagram.png)",
        "### Figure 2. Forest plot\n![Figure 2](../figures/forest_plot.png)",
        "### Figure 3. Sensitivity plot\n![Figure 3](../figures/sensitivity.png)",
        "### Figure 4. Risk-of-bias summary\n![Figure 4](../figures/rob_summary.png)",
    ])

    repaired = WritingAgent._backfill_after_fact_repair(manuscript)
    main_text = WritingAgent._main_text_before_figures(repaired)

    assert 2 in WritingAgent._numbered_label_refs(main_text, "Figure")


def test_submission_readiness_blocks_figures_without_legends(tmp_path: Path) -> None:
    project = Project("figure legend audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    long_text = " ".join(["figurelegendaudit"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full figure legend audit manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text} [1,2].",
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables."
            ),
            "## Results\nThe primary meta-analysis included 2 trials totaling 60 participants [1,2]. Table 1 summarizes the effects. Figure 1 shows the forest plot.",
            "## Discussion\nThe result is interpreted in relation to the selected trials [1,2].",
            (
                "## Tables\n### Table 1. Trial-level effects\n"
                "| Study | OR | 95% CI |\n|---|---:|---|\n| Trial 1 | 0.50 | 0.20 to 1.20 |\n\n"
                "Note: OR=odds ratio; CI=confidence interval."
            ),
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"},
                    {"row_id": "S2:0", "study_id": "S2", "outcome_name": "mortality"},
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/smith}}\n"
        "@article{jones2024,title={Trial report},doi={10.1000/jones}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        legend_audit = json.loads(zf.read("review/figure_legend_audit.json"))
        legend_html = zf.read("review/figure_legend_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    legend_gate = next(gate for gate in readiness["gates"] if gate["id"] == "figure_legends")
    assert legend_audit["summary"]["figure_count"] == 1
    assert legend_audit["summary"]["figures_with_legends"] == 0
    assert legend_audit["issues"][0]["code"] == "figure_legend_missing"
    assert legend_gate["status"] == "fail"
    assert "legends=0/1" in legend_gate["detail"]
    assert manifest["review"]["figure_legend_audit_missing_legends"] == 1
    assert "Figure Legend Issue" in legend_html


def test_submission_readiness_blocks_tables_without_footnotes(tmp_path: Path) -> None:
    project = Project("table footnote audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    long_text = " ".join(["footnoteaudit"] * 6200)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full table footnote audit manuscript",
            "## Abstract\nPrimary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            f"## Introduction\n{long_text} [1,2].",
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables."
            ),
            "## Results\nThe primary meta-analysis included 2 trials totaling 60 participants [1,2]. Table 1 summarizes the trial-level effects. Figure 1 shows the forest plot.",
            "## Discussion\nThe result is interpreted in relation to the selected trials [1,2].",
            "## Tables\n### Table 1. Trial-level effects\n| Study | OR | 95% CI |\n|---|---:|---|\n| Trial 1 | 0.50 | 0.20 to 1.20 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.\n[2] Jones J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"},
                    {"row_id": "S2:0", "study_id": "S2", "outcome_name": "mortality"},
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/smith}}\n"
        "@article{jones2024,title={Trial report},doi={10.1000/jones}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        table_audit = json.loads(zf.read("review/table_footnote_audit.json"))
        table_html = zf.read("review/table_footnote_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    table_gate = next(gate for gate in readiness["gates"] if gate["id"] == "table_footnotes")
    assert table_audit["summary"]["table_count"] == 1
    assert table_audit["summary"]["tables_with_notes"] == 0
    assert table_audit["issues"][0]["code"] == "table_footnote_missing"
    assert table_gate["status"] == "fail"
    assert "notes=0/1" in table_gate["detail"]
    assert manifest["review"]["table_footnote_audit_missing_notes"] == 1
    assert "Table Footnote Issue" in table_html


def test_submission_readiness_blocks_prisma_flow_mismatch(tmp_path: Path) -> None:
    project = Project("prisma audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full PRISMA audit manuscript",
            "## Abstract\n" + " ".join(["abstract"] * 350),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            (
                "## Methods\nFull search query: (COVID-19[tiab] AND corticosteroids[tiab]) "
                "AND (mortality[tiab] OR death[tiab]). The log odds ratio was calculated from 2 x 2 tables, "
                "and inverse-variance fixed-effect meta-analysis was used. "
                + " ".join(["methods"] * 1450)
            ),
            (
                "## Results\nThe search identified 99 records; 20 duplicates were removed, leaving 79 records for screening. "
                "The review screened 79 title/abstract records and assessed 8 full-text records. "
                "The primary meta-analysis included 2 trials totaling 60 participants. "
                "The pooled estimate was OR 0.66 (95% CI 0.53 to 0.82), with heterogeneity I2 0%. "
                + " ".join(["results"] * 1250)
            ),
            "## Discussion\n" + " ".join(["discussion"] * 1500),
            "## Tables\n### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |",
            "## Figures\n### Figure 1. PRISMA flow diagram\n![Figure 1. PRISMA flow diagram](../figures/prisma_diagram.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 6/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 109,
                "duplicates_removed": 79,
                "records_after_dedup": 30,
            },
            "screening": {"title_abstract_screened": 30},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row["source_quote"],
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 6/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )
    project.save_text("search_query.txt", "COVID-19 AND corticosteroids AND mortality")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "prisma_diagram.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        prisma_audit = json.loads(zf.read("review/prisma_audit.json"))
        prisma_html = zf.read("review/prisma_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/prisma_audit.json" in names
    assert "review/prisma_audit.html" in names
    assert prisma_audit["summary"]["expected_fields"] == 6
    assert prisma_audit["summary"]["matched_fields"] == 1
    assert prisma_audit["summary"]["mismatched_fields"] == 5
    assert prisma_audit["issues"][0]["code"] == "prisma_field_mismatch"
    prisma_gate = next(gate for gate in readiness["gates"] if gate["id"] == "prisma_flow")
    assert prisma_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["prisma_audit_mismatched_fields"] == 5
    assert "PRISMA Flow Mismatch" in prisma_html


def test_prisma_audit_extracts_chinese_flow_counts(tmp_path: Path) -> None:
    project = Project("chinese prisma audit", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文PRISMA审计",
            "",
            "## 结果",
            "本次检索识别781条记录，删除761条重复记录，去重后20条记录进入题名/摘要筛选。",
            "题名/摘要筛选20条记录，全文评估10篇，最终纳入2项研究进入主要Meta分析。",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 781,
                "duplicates_removed": 761,
                "records_after_dedup": 20,
            },
            "screening": {"title_abstract_screened": 20},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )

    audit = _build_prisma_audit_review(project)

    assert audit["passed"] is True
    assert audit["summary"]["matched_fields"] == 6
    assert audit["summary"]["failed_issues"] == 0


def test_prisma_audit_extracts_natural_language_title_abstract_screening_count(tmp_path: Path) -> None:
    project = Project("english prisma audit natural wording", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# PRISMA audit",
            "",
            "## Results",
            "The search identified 781 records. After removing 761 duplicates, 20 records remained for screening.",
            "We screened 20 titles and abstracts and assessed 10 full-text articles.",
            "The primary meta-analysis included 2 studies.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 781,
                "duplicates_removed": 761,
                "records_after_dedup": 20,
            },
            "screening": {"title_abstract_screened": 20},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )

    audit = _build_prisma_audit_review(project)

    assert audit["passed"] is True
    assert audit["summary"]["matched_fields"] == 6
    field = next(item for item in audit["fields"] if item["field"] == "title_abstract_screened")
    assert field["reported_values"] == [20]


def test_citation_audit_recommends_bibliography_trials_for_uncited_results_claim(tmp_path: Path) -> None:
    project = Project("citation bibliography trial recommendations", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Manuscript",
            "## Results\nTwo large trials contributed to the primary synthesis.",
            "## References\n"
            "[1] Anker SD, Butler J, Filippatos G, et al. Empagliflozin in Heart Failure with a Preserved Ejection Fraction. N Engl J Med. 2021;385:1451-1461.\n"
            "[2] Solomon SD, McMurray JJV, Claggett B, et al. Dapagliflozin in Heart Failure with Mildly Reduced or Preserved Ejection Fraction. N Engl J Med. 2022;387:1089-1098.",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "uncited_results_study_data_claim")
    assert issue["recommended_citations"][:2] == [1, 2]


def test_citation_audit_recommends_bibliography_trials_for_uncited_discussion_claim(tmp_path: Path) -> None:
    project = Project("citation bibliography discussion recommendations", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Manuscript",
            "## Discussion\nThe treatment appears to reduce the chance of worsening heart failure events.",
            "## References\n"
            "[1] Anker SD, Butler J, Filippatos G, et al. Empagliflozin in Heart Failure with a Preserved Ejection Fraction. N Engl J Med. 2021;385:1451-1461.\n"
            "[2] Solomon SD, McMurray JJV, Claggett B, et al. Dapagliflozin in Heart Failure with Mildly Reduced or Preserved Ejection Fraction. N Engl J Med. 2022;387:1089-1098.",
        ]),
        subdir="manuscript",
    )

    audit = _build_citation_audit_review(project)

    issue = next(item for item in audit["issues"] if item["code"] == "uncited_discussion_result_claim")
    assert issue["recommended_citations"][:2] == [1, 2]


def test_submission_readiness_blocks_when_exact_search_query_is_not_in_manuscript(tmp_path: Path) -> None:
    project = Project("search audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full search audit manuscript",
            "## Abstract\n" + " ".join(["abstract"] * 350),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            (
                "## Methods\nFull search query: (COVID-19 AND corticosteroids) AND mortality. "
                "The log odds ratio was calculated from 2 x 2 tables, and inverse-variance fixed-effect "
                "meta-analysis was used. " + " ".join(["methods"] * 1450)
            ),
            (
                "## Results\nThe search identified 109 records; 79 duplicates were removed, leaving 30 records for screening. "
                "The review screened 30 title/abstract records and assessed 10 full-text records. "
                "The review included 2 studies. The primary meta-analysis included 2 trials totaling 60 participants. "
                "The pooled estimate was OR 0.66 (95% CI 0.53 to 0.82), with heterogeneity I2 0%. "
                + " ".join(["results"] * 1250)
            ),
            "## Discussion\n" + " ".join(["discussion"] * 1500),
            "## Tables\n### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 6/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 109,
                "duplicates_removed": 79,
                "records_after_dedup": 30,
            },
            "screening": {"title_abstract_screened": 30},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row["source_quote"],
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 6/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        search_audit = json.loads(zf.read("review/search_strategy_audit.json"))
        search_html = zf.read("review/search_strategy_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/search_strategy_audit.json" in names
    assert "review/search_strategy_audit.html" in names
    assert search_audit["summary"]["exact_query_reproduced"] is False
    assert search_audit["summary"]["query_chars"] == len(exact_query)
    assert search_audit["issues"][0]["code"] == "search_query_not_reproduced"
    search_gate = next(gate for gate in readiness["gates"] if gate["id"] == "search_strategy")
    assert search_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["search_strategy_audit_exact_query_reproduced"] is False
    assert "Search Query Mismatch" in search_html


def test_submission_readiness_blocks_missing_referenced_figure_asset(tmp_path: Path) -> None:
    project = Project("figure audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full figure audit manuscript",
            "## Abstract\n" + " ".join(["abstract"] * 350),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables, and inverse-variance fixed-effect "
                "meta-analysis was used. " + " ".join(["methods"] * 1450)
            ),
            (
                "## Results\nThe search identified 109 records; 79 duplicates were removed, leaving 30 records for screening. "
                "The review screened 30 title/abstract records and assessed 10 full-text records. "
                "The review included 2 studies. The primary meta-analysis included 2 trials totaling 60 participants. "
                "The pooled estimate was OR 0.66 (95% CI 0.53 to 0.82), with heterogeneity I2 0%. "
                + " ".join(["results"] * 1250)
            ),
            "## Discussion\n" + " ".join(["discussion"] * 1500),
            "## Tables\n### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/missing_forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 6/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 109,
                "duplicates_removed": 79,
                "records_after_dedup": 30,
            },
            "screening": {"title_abstract_screened": 30},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row["source_quote"],
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 6/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "other_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        figure_audit = json.loads(zf.read("review/figure_audit.json"))
        figure_html = zf.read("review/figure_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/figure_audit.json" in names
    assert "review/figure_audit.html" in names
    assert figure_audit["summary"]["referenced_images"] == 1
    assert figure_audit["summary"]["missing_referenced_images"] == 1
    assert figure_audit["issues"][0]["code"] == "figure_image_missing"
    figures_gate = next(gate for gate in readiness["gates"] if gate["id"] == "figures")
    assert figures_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["figure_audit_missing_referenced_images"] == 1
    assert "Missing Figure Asset" in figure_html


def test_submission_readiness_blocks_primary_result_mismatch(tmp_path: Path) -> None:
    project = Project("primary result audit package", output_dir=tmp_path / uuid4().hex)
    exact_query = '("COVID-19"[tiab] AND corticosteroids[tiab]) AND mortality[tiab]'
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full primary result audit manuscript",
            "## Abstract\n" + " ".join(["abstract"] * 350),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            (
                f"## Methods\nFull search query: {exact_query}. "
                "The log odds ratio was calculated from 2 x 2 tables, and inverse-variance fixed-effect "
                "meta-analysis was used. " + " ".join(["methods"] * 1450)
            ),
            (
                "## Results\nThe search identified 109 records; 79 duplicates were removed, leaving 30 records for screening. "
                "The review screened 30 title/abstract records and assessed 10 full-text records. "
                "The review included 2 studies. The primary meta-analysis included 2 trials totaling 60 participants. "
                "There were 4/30 deaths in the corticosteroid groups and 8/30 deaths in the control groups. "
                "The pooled estimate was OR 0.72 (95% CI 0.53 to 0.82), with heterogeneity I2 0%. "
                + " ".join(["results"] * 1250)
            ),
            "## Discussion\n" + " ".join(["discussion"] * 1500),
            "## Tables\n### Table 1. Trial-level mortality counts\n| Study | Events |\n|---|---|\n| Trial 1 | 1/10 vs 2/10 |",
            "## Figures\n### Figure 1. Forest plot\n![Figure 1. Forest plot](../figures/forest_plot.png)",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    selected_rows = [
        {
            "row_id": "S1:0",
            "study_id": "S1",
            "outcome_name": "mortality",
            "events_intervention": 1,
            "total_intervention": 10,
            "events_control": 2,
            "total_control": 10,
            "source_quote": "Mortality was 1/10 vs 2/10.",
            "source_quote_verified": True,
        },
        {
            "row_id": "S2:0",
            "study_id": "S2",
            "outcome_name": "mortality",
            "events_intervention": 3,
            "total_intervention": 20,
            "events_control": 6,
            "total_control": 20,
            "source_quote": "Mortality was 3/20 vs 6/20.",
            "source_quote_verified": True,
        },
    ]
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": selected_rows,
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "prisma_flow.json",
        {
            "identification": {
                "records_identified": 109,
                "duplicates_removed": 79,
                "records_after_dedup": 30,
            },
            "screening": {"title_abstract_screened": 30},
            "eligibility": {"full_text_assessed": 10},
            "included": {"studies_included": 2},
        },
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    **row,
                    "outcome_index": 0,
                    "source_quote_match": row["source_quote"],
                    "requires_review": False,
                    "conflicts": [],
                }
                for row in selected_rows
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []},
            "S2": {"full_text": "Mortality was 3/20 vs 6/20.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "meta_results.json",
        {
            "primary_outcome": {
                "outcome_name": "mortality",
                "n_studies": 2,
                "effect_measure": "OR",
                "pooled_effect": 0.66,
                "ci_lower": 0.53,
                "ci_upper": 0.82,
                "model": "fixed",
                "studies": [
                    {"study_id": "S1", "study_label": "Trial 1", "yi": -0.69, "vi": 1.8, "se": 1.34, "weight": 50.0},
                    {"study_id": "S2", "study_label": "Trial 2", "yi": -0.72, "vi": 0.9, "se": 0.95, "weight": 50.0},
                ],
            }
        },
        subdir="analysis",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text(
        "references.bib",
        "@article{smith2024,title={Trial report},doi={10.1000/trial},pmid={12345}}\n",
    )
    project.save_text("search_query.txt", exact_query)
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        primary_audit = json.loads(zf.read("review/primary_result_audit.json"))
        primary_html = zf.read("review/primary_result_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/primary_result_audit.json" in names
    assert "review/primary_result_audit.html" in names
    assert primary_audit["summary"]["expected_fields"] >= 6
    assert primary_audit["summary"]["mismatched_fields"] == 1
    assert primary_audit["issues"][0]["code"] == "primary_result_field_missing"
    assert primary_audit["issues"][0]["field"] == "pooled_effect"
    primary_gate = next(gate for gate in readiness["gates"] if gate["id"] == "primary_result")
    assert primary_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["primary_result_audit_mismatched_fields"] == 1
    assert "Primary Result Mismatch" in primary_html


def test_submission_readiness_blocks_unsupported_primary_effect_claim(tmp_path: Path) -> None:
    project = Project("claim support package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Claim support manuscript",
            "",
            "## Abstract",
            "The pooled HR was 0.90 (95% CI 0.80 to 1.01), suggesting no clear reduction.",
            "",
            "## Results",
            "The pooled HR was 0.81 (95% CI 0.74 to 0.88), and certainty was High.",
            "",
            "## Discussion",
            "The result was interpreted against the prespecified primary outcome.",
            "",
            "## References",
            "[1] Example trial.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "primary_effect": {
                "outcome_name": "Composite endpoint",
                "effect_measure": "HR",
                "n_studies": 2,
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "grade": {
                "outcomes": [
                    {
                        "outcome_name": "Composite endpoint",
                        "certainty": "High",
                        "effect_summary": "HR 0.81 (95% CI: 0.74 to 0.88)",
                    }
                ]
            },
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        audit = json.loads(zf.read("review/claim_support_audit.json"))
        audit_html = zf.read("review/claim_support_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/claim_support_audit.json" in names
    assert "review/claim_support_audit.html" in names
    assert audit["summary"]["checked_claims"] == 3
    assert audit["summary"]["unsupported_claims"] == 1
    unsupported = [row for row in audit["claims"] if row["status"] == "unsupported"]
    assert len(unsupported) == 1
    assert unsupported[0]["claim_type"] == "primary_effect"
    assert "0.90" in unsupported[0]["sentence"]
    assert "expected HR 0.81" in unsupported[0]["reason"]
    assert audit["summary"]["supported_claims"] == 2
    assert "Claim support manuscript" in audit_html
    assert "0.90" in audit_html
    claim_gate = next(gate for gate in readiness["gates"] if gate["id"] == "claim_support")
    assert claim_gate["status"] == "fail"
    assert "unsupported=1" in claim_gate["detail"]
    assert manifest["review"]["claim_support_audit_included"] is True
    assert manifest["review"]["claim_support_unsupported_claims"] == 1


def test_claim_support_audit_recognizes_chinese_primary_effect_claim(tmp_path: Path) -> None:
    project = Project("claim support zh package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析",
            "",
            "## 结果",
            "主要分析显示，合并HR为0.81（95% CI 0.74–0.88），GRADE证据确定性评为高。",
            "",
            "## 参考文献",
            "[1] 示例研究。",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "output_language": "zh",
            "primary_effect": {
                "effect_measure": "HR",
                "n_studies": 2,
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "grade": {
                "outcomes": [
                    {"outcome_name": "主要结局", "certainty": "High"}
                ]
            },
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/claim_support_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert audit["summary"]["checked_claims"] == 2
    assert audit["summary"]["supported_claims"] == 2
    assert audit["summary"]["unsupported_claims"] == 0
    assert {claim["claim_type"] for claim in audit["claims"]} == {"primary_effect", "grade_certainty"}
    claim_gate = next(gate for gate in readiness["gates"] if gate["id"] == "claim_support")
    assert claim_gate["status"] == "pass"


def test_claim_support_audit_accepts_less_than_high_for_moderate_certainty(tmp_path: Path) -> None:
    project = Project("claim support non-high certainty package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis",
            "",
            "## Discussion",
            "The certainty profile therefore reports both the statistical result and the reasons confidence may remain less than high [1].",
            "",
            "## References",
            "[1] GRADE summary.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "grade": {"outcomes": [{"outcome_name": "mortality", "certainty": "Moderate"}]},
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/claim_support_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert audit["summary"]["checked_claims"] == 1
    assert audit["summary"]["supported_claims"] == 1
    assert audit["summary"]["unsupported_claims"] == 0
    claim_gate = next(gate for gate in readiness["gates"] if gate["id"] == "claim_support")
    assert claim_gate["status"] == "pass"


def test_claim_support_audit_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("claim support zh handoff package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析",
            "",
            "## 摘要",
            "主要分析显示，合并HR为0.90（95% CI 0.80至1.01），GRADE证据确定性评为低。",
            "",
            "## 结果",
            "主要分析显示，合并HR为0.81（95% CI 0.74至0.88），GRADE证据确定性评为高。",
            "",
            "## 参考文献",
            "[1] 示例研究。",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "primary_effect": {
                "effect_measure": "HR",
                "n_studies": 2,
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "grade": {
                "outcomes": [
                    {"outcome_name": "主要结局", "certainty": "High"}
                ]
            },
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/claim_support_audit.json"))
        html = zf.read("review/claim_support_audit.html").decode("utf-8")

    assert audit["language"] == "zh"
    assert '<html lang="zh">' in html
    assert "MetaAgent 正文主张支持审计" in html
    assert "核对主要效应和GRADE确定性主张是否与结构化事实源一致" in html
    assert "已检查" in html
    assert "已支持" in html
    assert "未支持" in html
    assert "失败问题" in html
    assert "不受支持的正文主张" in html
    assert "主效应主张" in html
    assert "证据确定性主张" in html
    assert "已支持" in html
    assert "不支持" in html
    assert any(claim["support_source"] == "manuscript_facts.primary_effect" for claim in audit["claims"])
    assert any(claim["support_source"] == "manuscript_facts.grade" for claim in audit["claims"])
    assert "结构化主效应事实" in html
    assert "结构化GRADE事实" in html
    assert "manuscript_facts" not in html
    assert "预期 HR 0.81" in html
    assert "MetaAgent Claim Support Audit" not in html
    assert "Unsupported Manuscript Claim" not in html
    assert "Checked Claims" not in html


def test_claim_support_audit_omits_chinese_unsupported_banner_when_all_claims_supported(tmp_path: Path) -> None:
    project = Project("claim support zh clean package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析",
            "",
            "## 结果",
            "主要分析显示，合并HR为0.81（95% CI 0.74至0.88），GRADE证据确定性评为高。",
            "",
            "## 参考文献",
            "[1] 示例研究。",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
            "primary_effect": {
                "effect_measure": "HR",
                "n_studies": 2,
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "grade": {
                "outcomes": [
                    {"outcome_name": "主要结局", "certainty": "High"}
                ]
            },
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/claim_support_audit.json"))
        html = zf.read("review/claim_support_audit.html").decode("utf-8")

    assert audit["summary"]["unsupported_claims"] == 0
    assert "<h2>不受支持的正文主张</h2>" not in html
    assert "效应量或证据确定性主张与结构化事实源不一致" not in html


def test_claim_support_audit_ignores_background_trial_effect_claims(tmp_path: Path) -> None:
    project = Project("claim support background package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Background effect manuscript",
            "",
            "## Introduction",
            "A prior trial reported HR 0.79 (95% CI 0.69 to 0.90) for a related outcome.",
            "",
            "## References",
            "[1] Prior trial.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "primary_effect": {
                "effect_measure": "HR",
                "n_studies": 2,
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert "review/claim_support_audit.json" not in names
    assert not any(gate["id"] == "claim_support" and gate["status"] == "fail" for gate in readiness["gates"])


def test_claim_support_audit_ignores_method_descriptions_without_result_values(tmp_path: Path) -> None:
    project = Project("claim support method package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Method text manuscript",
            "",
            "## Methods",
            "Study-level HRs and standard errors were pooled using inverse-variance methods.",
            "Risk of bias and certainty of evidence were assessed after extraction.",
            "",
            "## References",
            "[1] Methods source.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "evidence_gap",
            "primary_effect": {
                "effect_measure": "HR",
                "pooled_effect": 0.81,
                "ci_lower": 0.74,
                "ci_upper": 0.88,
            },
            "grade": {"outcomes": [{"outcome_name": "Primary", "certainty": "High"}]},
            "evidence_readiness": {"status": "ready", "blockers": []},
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())

    assert "review/claim_support_audit.json" not in names


def test_submission_readiness_blocks_internal_abstract_notes(tmp_path: Path) -> None:
    project = Project("abstract polish audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Abstract",
            "**Importance:** The clinical topic needs synthesis.",
            "**Objective:** To estimate the treatment effect.",
            "**Data sources:** PubMed and trial registries.",
            "**Study selection:** Randomized trials were eligible.",
            "**Data extraction and synthesis:** Data were extracted from source records.",
            "**Main outcome and measures:** Mortality.",
            "**Results:** The pooled effect was OR 0.66 (95% CI 0.53 to 0.82).",
            "**Conclusions and relevance:** The intervention was associated with lower mortality. "
            "Supplementary source context: 3 retrieved/screened record(s) use limited source text.",
            "",
            "## Introduction",
            "Brief rationale.",
            "",
            "## Methods",
            "The Boolean search strategy was mortality AND treatment.",
            "",
            "## Results",
            "The primary analysis is summarized above.",
            "",
            "## Discussion",
            "Brief interpretation.",
            "",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        abstract_audit = json.loads(zf.read("review/abstract_audit.json"))
        abstract_html = zf.read("review/abstract_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/abstract_audit.json" in names
    assert "review/abstract_audit.html" in names
    assert abstract_audit["summary"]["failed_issues"] == 1
    assert abstract_audit["issues"][0]["code"] == "abstract_internal_note"
    abstract_gate = next(gate for gate in readiness["gates"] if gate["id"] == "abstract_polish")
    assert abstract_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["abstract_audit_failed_issues"] == 1
    assert "Abstract Polish Issue" in abstract_html


def test_submission_readiness_blocks_internal_publication_tone(tmp_path: Path) -> None:
    project = Project("publication tone audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Abstract",
            "**Importance:** The clinical topic needs synthesis.",
            "**Objective:** To estimate the treatment effect.",
            "**Data sources:** PubMed and trial registries.",
            "**Study selection:** Randomized trials were eligible.",
            "**Data extraction and synthesis:** Data were extracted from source records.",
            "**Main outcome and measures:** Mortality.",
            "**Results:** The pooled effect was OR 0.66 (95% CI 0.53 to 0.82).",
            "**Conclusions and relevance:** The intervention was associated with lower mortality.",
            "",
            "## Introduction",
            "Brief rationale.",
            "",
            "## Methods",
            "The Boolean search strategy was mortality AND treatment.",
            "",
            "## Results",
            "### Evidence-readiness status",
            "The trial-level table should be read as the audit trail for the pooled result. "
            "The structured data files remain important and the manuscript is internally consistent.",
            "",
            "## Discussion",
            "Brief interpretation.",
            "",
            "## Supplementary Materials",
            "### Appendix 2. Source audit for selected primary rows",
            "This appendix can retain source audit terminology.",
            "",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        tone_audit = json.loads(zf.read("review/publication_tone_audit.json"))
        tone_html = zf.read("review/publication_tone_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/publication_tone_audit.json" in names
    assert "review/publication_tone_audit.html" in names
    assert tone_audit["summary"]["failed_issues"] == 4
    assert {issue["code"] for issue in tone_audit["issues"]} == {"publication_internal_tone"}
    assert all("Source audit for selected primary rows" not in issue.get("excerpt", "") for issue in tone_audit["issues"])
    tone_gate = next(gate for gate in readiness["gates"] if gate["id"] == "publication_tone")
    assert tone_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["publication_tone_audit_failed_issues"] == 4
    assert "Publication Tone Issue" in tone_html


def test_publication_tone_audit_blocks_internal_database_source_names(tmp_path: Path) -> None:
    project = Project("publication tone internal database", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "## Abstract",
            "The review synthesized randomized evidence.",
            "## Introduction",
            "The condition requires evidence synthesis.",
            "## Methods",
            "The search covered an internal literature database and OpenAlex.",
            "## Results",
            "Two randomized trials contributed to the primary synthesis.",
            "## Discussion",
            "The result should be interpreted according to baseline risk.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        tone_audit = json.loads(zf.read("review/publication_tone_audit.json"))

    labels = {issue["label"] for issue in tone_audit["issues"]}
    assert tone_audit["summary"]["failed_issues"] >= 1
    assert "Internal database source label" in labels


def test_publication_tone_audit_blocks_main_text_workflow_jargon(tmp_path: Path) -> None:
    project = Project("publication tone workflow jargon", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "## Abstract",
            "The review synthesized randomized evidence.",
            "## Introduction",
            "The condition requires evidence synthesis.",
            "## Methods",
            "Search records, full-text availability, and retrieval warnings were retained.",
            "For each candidate primary row, the extraction record captured the source excerpt.",
            "A PDF parser split some text, so source checking remained visible.",
            "The manuscript therefore avoids converting a single p value into the whole interpretation.",
            "It is safer for a manuscript to preserve a serious concern for expert review.",
            "## Results",
            "Table 1 lists the documentation status of selected rows.",
            "## Discussion",
            "The result should be interpreted according to baseline risk.",
            "## Supplement",
            "The supplementary extraction review may contain workflow details.",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        tone_audit = json.loads(zf.read("review/publication_tone_audit.json"))

    labels = {issue["label"] for issue in tone_audit["issues"]}
    assert tone_audit["summary"]["failed_issues"] >= 5
    assert "Retrieval-warning wording" in labels
    assert "Extraction-record wording" in labels
    assert "PDF-parser wording" in labels
    assert "Source-checking wording" in labels
    assert "Documentation-status wording" in labels
    assert "Manuscript-self-reference wording" in labels


def test_publication_tone_audit_blocks_chinese_internal_workflow_language(tmp_path: Path) -> None:
    project = Project("chinese publication tone audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "## 摘要",
            "主要结局提示治疗与较低风险相关。",
            "",
            "## 引言",
            "本稿采用事实锁定写作，把选定主要行、来源摘录、效应量计算和GRADE判断连接起来。",
            "",
            "## 方法",
            "研究者按照PRISMA流程完成系统综述。",
            "",
            "## 结果",
            "结构化数据文件证明摘要、结果表和图形使用同一套事实。",
            "",
            "## 讨论",
            "这种写法的核心价值是可审计性，审稿意见能定位至具体数据行。",
            "来源核验字段、结构化证据表和写作模块暴露了提取复核界面的数据重新生成流程。",
            "",
            "## 补充材料",
            "### 来源审计附录",
            "补充材料可以保留来源审计术语。",
            "",
            "## 参考文献",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        tone_audit = json.loads(zf.read("review/publication_tone_audit.json"))

    labels = {issue["label"] for issue in tone_audit["issues"]}
    assert tone_audit["summary"]["failed_issues"] >= 3
    assert "Chinese fact-locked wording" in labels
    assert "Chinese structured-data-file wording" in labels
    assert "Chinese auditability wording" in labels
    assert "Chinese workflow wording" in labels
    assert all("来源审计附录" not in issue.get("excerpt", "") for issue in tone_audit["issues"])


def test_publication_tone_audit_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("publication tone zh handoff package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "## 摘要",
            "主要结局提示治疗与较低风险相关。",
            "",
            "## 引言",
            "本稿采用事实锁定写作，把选定主要行、来源摘录、效应量计算和GRADE判断连接起来。",
            "",
            "## 讨论",
            "来源核验字段、结构化证据表和写作模块暴露了提取复核界面的数据重新生成流程。",
            "",
            "## 补充材料",
            "### 来源审计附录",
            "补充材料可以保留来源审计术语。",
            "",
            "## 参考文献",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_facts.json", {"output_language": "zh"}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/publication_tone_audit.json"))
        html = zf.read("review/publication_tone_audit.html").decode("utf-8")

    assert audit["language"] == "zh"
    assert audit["summary"]["failed_issues"] >= 1
    assert '<html lang="zh">' in html
    assert "MetaAgent 投稿语气审计" in html
    assert "检查主稿正文中是否残留内部系统、工程或复核流程用语" in html
    assert "扫描词数" in html
    assert "禁用短语" in html
    assert "失败问题" in html
    assert "投稿语气问题" in html
    assert "主稿正文包含内部或工程化措辞" in html
    assert "问题" in html
    assert "严重性" in html
    assert "标签" in html
    assert "命中文本" in html
    assert "上下文" in html
    assert "MetaAgent Publication Tone Audit" not in html
    assert "Publication Tone Issue" not in html


def test_publication_tone_audit_omits_chinese_issue_banner_when_clean(tmp_path: Path) -> None:
    project = Project("publication tone zh clean package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "## 摘要",
            "主要结局提示治疗与较低风险相关。",
            "",
            "## 引言",
            "本研究总结随机试验结果并评价主要结局。",
            "",
            "## 讨论",
            "结果应结合样本量、随访时间和证据确定性解释。",
            "",
            "## 参考文献",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_facts.json", {"output_language": "zh"}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        audit = json.loads(zf.read("review/publication_tone_audit.json"))
        html = zf.read("review/publication_tone_audit.html").decode("utf-8")

    assert audit["summary"]["failed_issues"] == 0
    assert "未记录投稿语气问题" in html
    assert "<h2>投稿语气问题</h2>" not in html
    assert "主稿正文包含内部或工程化措辞" not in html


def test_clinical_interpretation_audit_blocks_process_only_discussion(tmp_path: Path) -> None:
    project = Project("clinical interpretation process-only package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Abstract",
            "**Importance:** The clinical topic needs synthesis.",
            "**Objective:** To estimate treatment effects.",
            "**Data sources:** PubMed.",
            "**Study selection:** Randomized trials.",
            "**Data extraction and synthesis:** Aggregate data were synthesized.",
            "**Main outcome and measures:** Cardiovascular events.",
            "**Results:** Two studies contributed data.",
            "**Conclusions and relevance:** Treatment may reduce events.",
            "",
            "## Introduction",
            "The clinical question needs a systematic review.",
            "",
            "## Methods",
            "Randomized trials were synthesized.",
            "",
            "## Results",
            "The pooled estimate favored treatment.",
            "",
            "## Discussion",
            "The main value of this review is transparent traceability from extracted rows to the final manuscript. "
            "Readers can inspect the source audit, calculation files, and generated tables to confirm that the "
            "same numeric fields were used across sections.",
            "",
            "## Conclusion",
            "The manuscript provides a transparent summary.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        audit = json.loads(zf.read("review/clinical_interpretation_audit.json"))
        html = zf.read("review/clinical_interpretation_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/clinical_interpretation_audit.json" in names
    assert "review/clinical_interpretation_audit.html" in names
    assert audit["passed"] is False
    assert audit["summary"]["covered_domains"] < audit["summary"]["minimum_domains"]
    assert "absolute_risk_translation" in audit["summary"]["missing_domains"]
    assert "benefit_harm_safety" in audit["summary"]["missing_domains"]
    assert audit["issues"][0]["code"] == "clinical_interpretation_depth_low"
    interpretation_gate = next(gate for gate in readiness["gates"] if gate["id"] == "clinical_interpretation")
    assert interpretation_gate["status"] == "fail"
    assert manifest["review"]["clinical_interpretation_audit_failed_issues"] == 1
    assert "Clinical Interpretation Issue" in html


def test_clinical_interpretation_audit_accepts_multidomain_discussion(tmp_path: Path) -> None:
    project = Project("clinical interpretation multidomain package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Discussion",
            "The pooled HR was 0.81 (95% CI 0.74 to 0.88), which indicates fewer cardiovascular events.",
            "The same relative effect should be translated through baseline risk, absolute benefit, and NNT.",
            "Because the primary outcome is a composite endpoint, hospitalization and mortality should be interpreted separately.",
            "Safety, adverse events, renal function, volume status, and ketoacidosis risk affect the benefit-harm balance.",
            "Applicability depends on subgroup, kidney function, diabetes status, ejection fraction, and comorbidity.",
            "Implementation requires monitoring, follow-up, adherence support, cost discussion, and patient preference.",
            "Certainty, heterogeneity, publication bias, and limitations should temper the strength of inference.",
            "",
            "## Conclusion",
            "Treatment appears clinically useful for eligible patients when the absolute benefit outweighs harms.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")

    audit = _build_clinical_interpretation_audit_review(project)

    assert audit["passed"] is True
    assert audit["summary"]["failed_issues"] == 0
    assert audit["summary"]["covered_domains"] >= audit["summary"]["minimum_domains"]
    assert audit["summary"]["result_context_present"] is True
    assert audit["summary"]["missing_domains"] == []


def test_clinical_interpretation_audit_blocks_process_framing_even_when_domains_present(tmp_path: Path) -> None:
    project = Project("clinical interpretation process-framed package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Discussion",
            "The pooled HR was 0.81 (95% CI 0.74 to 0.88), indicating fewer cardiovascular events.",
            "The same relative effect should be translated through baseline risk, absolute benefit, and NNT.",
            "Because the primary outcome is a composite endpoint, hospitalization and mortality should be interpreted separately.",
            "Safety, adverse events, renal function, volume status, and ketoacidosis risk affect the benefit-harm balance.",
            "Applicability depends on subgroup, kidney function, diabetes status, ejection fraction, and comorbidity.",
            "Implementation requires monitoring, follow-up, adherence support, cost discussion, and patient preference.",
            "Certainty, heterogeneity, publication bias, and limitations should temper the strength of inference.",
            "The main value of this review is transparent traceability from extracted rows, source audit files, "
            "calculation files, and generated tables to the final manuscript.",
            "",
            "## Conclusion",
            "The manuscript is useful because readers can inspect the source audit and calculation files.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")

    audit = _build_clinical_interpretation_audit_review(project)

    codes = {issue["code"] for issue in audit["issues"]}
    assert audit["passed"] is False
    assert "clinical_discussion_process_framing" in codes
    assert audit["summary"]["process_framing_paragraphs"] >= 1


def test_clinical_interpretation_audit_blocks_single_process_framed_main_text_note(tmp_path: Path) -> None:
    project = Project("single process-framed discussion note", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "",
            "## 讨论",
            "合并HR为0.81，95% CI为0.74至0.88，提示SGLT2抑制剂可降低心血管死亡或心衰住院风险。",
            "该结果需要结合基线风险、绝对获益、复合终点组成、安全性、适用性、监测随访和GRADE证据确定性解释。",
            "来源提示：3条检索或筛选记录使用了受限来源文本或元数据，但这些记录仅用于筛选或背景上下文。",
            "",
            "## 结论",
            "SGLT2抑制剂对符合纳入标准的HFmrEF/HFpEF患者可能具有明确临床价值。",
            "",
            "## 参考文献",
            "［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta", "output_language": "zh"}, subdir="manuscript")

    audit = _build_clinical_interpretation_audit_review(project)

    codes = {issue["code"] for issue in audit["issues"]}
    assert audit["passed"] is False
    assert "clinical_discussion_process_framing" in codes
    assert audit["summary"]["process_framing_paragraphs"] == 1


def test_clinical_interpretation_audit_blocks_overlong_repetitive_discussion(tmp_path: Path) -> None:
    project = Project("clinical interpretation repetitive package", output_dir=tmp_path / uuid4().hex)
    repeated_paragraphs = [
        (
            f"Clinical interpretation paragraph {i} restates that the pooled HR was 0.81 and that baseline risk, "
            "absolute benefit, composite endpoint components, safety, applicability, monitoring, follow-up, "
            "patient preference, cost, certainty, heterogeneity, and limitations should all be considered."
        )
        for i in range(1, 29)
    ]
    project.save_text(
        "draft.md",
        (
            "# Meta-analysis manuscript\n\n"
            "## Discussion\n\n"
            + "\n\n".join(repeated_paragraphs)
            + "\n\n## Conclusion\n\n"
            "Treatment appears clinically useful when absolute benefit outweighs harms.\n\n"
            "## References\n\n"
            "[1] Smith J. Trial report."
        ),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")

    audit = _build_clinical_interpretation_audit_review(project)

    codes = {issue["code"] for issue in audit["issues"]}
    assert audit["passed"] is False
    assert "clinical_discussion_too_long" in codes
    assert "clinical_discussion_redundant_domains" in codes
    assert audit["summary"]["discussion_paragraph_count"] == 28
    assert audit["summary"]["redundant_domain_count"] >= 1


def test_chinese_clinical_interpretation_audit_accepts_multidomain_discussion(tmp_path: Path) -> None:
    project = Project("clinical interpretation zh multidomain package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "",
            "## 讨论",
            "合并HR为0.81，95% CI为0.74至0.88，提示心血管事件风险降低。",
            "临床解释应结合基线风险、绝对获益和获益需治数。",
            "复合终点需要区分心衰住院和心血管死亡等组成事件。",
            "安全性、不良事件、肾功能、容量状态和酮症酸中毒风险会影响获益风险平衡。",
            "适用性取决于亚组、射血分数、糖尿病状态、合并症和肾功能。",
            "实施时需要监测、随访、依从性支持、费用讨论和患者偏好。",
            "证据确定性、异质性、发表偏倚和局限性应限制推断强度。",
            "",
            "## 结论",
            "在绝对获益超过潜在伤害时，治疗对适用患者具有临床价值。",
            "",
            "## 参考文献",
            "［1］ Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta", "output_language": "zh"}, subdir="manuscript")

    audit = _build_clinical_interpretation_audit_review(project)

    assert audit["language"] == "zh"
    assert audit["passed"] is True
    assert audit["summary"]["covered_domains"] >= audit["summary"]["minimum_domains"]
    assert audit["summary"]["result_context_present"] is True


def test_submission_readiness_blocks_verbose_pico_in_interpretive_sections(tmp_path: Path) -> None:
    project = Project("readability audit package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Meta-analysis manuscript",
            "",
            "## Abstract",
            "**Importance:** HFpEF needs evidence synthesis.",
            "**Objective:** To estimate the treatment effect.",
            "**Data sources:** PubMed and trial registries.",
            "**Study selection:** Randomized trials were eligible.",
            "**Data extraction and synthesis:** Data were extracted from source records.",
            "**Main outcome and measures:** Heart failure hospitalization.",
            "**Results:** The pooled effect was HR 0.81 (95% CI 0.74 to 0.88).",
            "**Conclusions and relevance:** SGLT2 inhibitors were associated with fewer events.",
            "",
            "## Introduction",
            "Patients with HFmrEF/HFpEF are clinically heterogeneous.",
            "",
            "## Methods",
            "Eligible records enrolled adults with HFpEF confirmed by echocardiography, cardiac MRI, "
            "or radionuclide ventriculography and compared treatment with placebo. "
            "This detailed eligibility wording should be summarized in Methods and moved to protocol details.",
            "",
            "## Results",
            "The pooled estimate favored SGLT2 inhibitors.",
            "",
            "## Discussion",
            "Clinically, the result should be translated into decisions only after considering whether "
            "Adults (>=18 years) with heart failure and left ventricular ejection fraction (LVEF) >=40%, "
            "confirmed by echocardiography, cardiac MRI, or radionuclide ventriculography, resemble "
            "the local patient population.",
            "",
            "## Conclusion",
            "SGLT2 inhibitors were associated with fewer cardiovascular events.",
            "",
            "## Supplementary Materials",
            "The exact search query may contain long technical syntax.",
            "",
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = zf.namelist()
        readability_audit = json.loads(zf.read("review/readability_audit.json"))
        readability_html = zf.read("review/readability_audit.html").decode("utf-8")
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "review/readability_audit.json" in names
    assert "review/readability_audit.html" in names
    assert readability_audit["summary"]["failed_issues"] == 4
    assert readability_audit["issues"][0]["code"] == "verbose_pico_fragment"
    assert {issue["section"] for issue in readability_audit["issues"]} == {"Methods", "Discussion"}
    readability_gate = next(gate for gate in readiness["gates"] if gate["id"] == "readability")
    assert readability_gate["status"] == "fail"
    assert readiness["status"] == "blocked"
    assert manifest["review"]["readability_audit_failed_issues"] == 4
    assert "Readability Issue" in readability_html


def test_submission_readiness_blocks_chinese_verbose_pico_in_interpretive_sections(tmp_path: Path) -> None:
    project = Project("readability audit chinese package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文Meta分析稿件",
            "",
            "## 摘要",
            "**重要性：** HFpEF需要证据综合。",
            "**目的：** 估计治疗效应。",
            "**资料来源：** PubMed和试验注册库。",
            "**研究选择：** 纳入随机试验。",
            "**结果：** 合并效应为HR 0.81（95%CI 0.74至0.88）。",
            "**结论和意义：** SGLT2抑制剂与较少心血管事件相关。",
            "",
            "## 引言",
            "HFmrEF/HFpEF患者具有临床异质性。",
            "",
            "## 方法",
            "合格研究纳入经超声心动图、心脏磁共振或核素心室造影确认的HFpEF成人，"
            "并比较任何获批剂量的治疗与安慰剂、无药物治疗或假干预。"
            "这些详细资格条件应在方法部分概括，完整细节放入方案或附录。",
            "",
            "## 结果",
            "合并估计支持SGLT2抑制剂。",
            "",
            "## 讨论",
            "解释结果时，应考虑本地人群是否符合经超声心动图、心脏磁共振或核素心室造影确认的HFpEF成人。",
            "",
            "## 结论",
            "SGLT2抑制剂与较少心血管事件相关。",
            "",
            "## 参考文献",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta", "language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readability_audit = json.loads(zf.read("review/readability_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert readability_audit["summary"]["failed_issues"] == 6
    assert {issue["section"] for issue in readability_audit["issues"]} == {"方法", "讨论"}
    assert {
        "Detailed Chinese diagnostic-confirmation phrase",
        "Detailed Chinese diagnostic-modality phrase",
        "Detailed Chinese intervention dose phrase",
        "Detailed Chinese comparator phrase",
    }.issubset({issue["label"] for issue in readability_audit["issues"]})
    readability_gate = next(gate for gate in readiness["gates"] if gate["id"] == "readability")
    assert readability_gate["status"] == "fail"


def test_submission_readiness_blocks_overlong_english_interpretive_sentences(tmp_path: Path) -> None:
    project = Project("readability long sentence package", output_dir=tmp_path / uuid4().hex)
    long_discussion_sentence = (
        "For clinical interpretation the pooled estimate should be considered together with baseline risk, "
        "the component structure of the composite endpoint, follow-up duration, renal function, volume status, "
        "background therapy, patient preferences, drug cost, local monitoring capacity, adverse-event counseling, "
        "and the possibility that hospitalization thresholds differ across health systems even when the underlying "
        "biology and randomized trial effect are directionally consistent."
    )
    long_methods_sentence = (
        "Eligible records enrolled adults with heart failure, compared treatment with placebo, reported a compatible "
        "time-to-event endpoint, and were reconciled before synthesis according to prespecified source hierarchy rules."
    )
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Introduction",
            "HFmrEF/HFpEF requires clinically grounded synthesis.",
            "",
            "## Methods",
            long_methods_sentence,
            "",
            "## Discussion",
            long_discussion_sentence,
            "",
            "## Conclusion",
            "SGLT2 inhibitors were associated with fewer cardiovascular events.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readability_audit = json.loads(zf.read("review/readability_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert readability_audit["summary"]["overlong_sentences"] == 1
    assert readability_audit["summary"]["failed_issues"] == 1
    assert readability_audit["issues"][0]["code"] == "overlong_sentence"
    assert readability_audit["issues"][0]["section"] == "Discussion"
    assert "Methods" not in {issue.get("section") for issue in readability_audit["issues"]}
    readability_gate = next(gate for gate in readiness["gates"] if gate["id"] == "readability")
    assert readability_gate["status"] == "fail"


def test_submission_readiness_blocks_overlong_chinese_interpretive_sentences(tmp_path: Path) -> None:
    project = Project("readability long chinese sentence package", output_dir=tmp_path / uuid4().hex)
    long_discussion_sentence = (
        "临床解释时应同时考虑基线风险、复合终点组成、随访时间、肾功能、容量状态、背景治疗、"
        "患者偏好、药物费用、本地监测能力、不良事件宣教以及不同医疗体系中心衰住院阈值可能不同，"
        "否则同一个合并HR在实际诊疗中的意义可能被过度简化并影响医患共同决策。"
    )
    project.save_text(
        "draft.md",
        "\n".join([
            "# 中文稿件",
            "",
            "## 引言",
            "HFmrEF/HFpEF需要结合临床语境解释。",
            "",
            "## 方法",
            "符合条件的研究需报告兼容的时间到事件终点，并按预设来源层级协调。",
            "",
            "## 讨论",
            long_discussion_sentence,
            "",
            "## 结论",
            "SGLT2抑制剂与较少心血管事件相关。",
            "",
            "## 参考文献",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_facts.json", {"report_type": "meta", "language": "zh"}, subdir="manuscript")
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readability_audit = json.loads(zf.read("review/readability_audit.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    assert readability_audit["summary"]["overlong_sentences"] == 1
    assert readability_audit["summary"]["failed_issues"] == 1
    assert readability_audit["issues"][0]["code"] == "overlong_sentence"
    assert readability_audit["issues"][0]["section"] == "讨论"
    readability_gate = next(gate for gate in readiness["gates"] if gate["id"] == "readability")
    assert readability_gate["status"] == "fail"


def test_submission_readiness_warns_when_no_benchmark_is_attached(tmp_path: Path) -> None:
    project = Project("submission readiness no benchmark", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66. Primary outcome RR 0.74. Primary outcome HR 0.81. Primary outcome MD -1.0.",
            "",
            *_clinical_discussion_fixture(effect="OR 0.66", endpoint="mortality"),
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 1/10 vs 2/10.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []}},
        subdir="papers",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "mortality AND treatment")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    benchmark_gate = next(gate for gate in readiness["gates"] if gate["id"] == "benchmark")
    assert benchmark_gate["status"] == "warn"
    assert readiness["passed"] is True
    assert readiness["status"] == "ready_with_warnings"
    assert manifest["submission"]["passed"] is True
    assert manifest["submission"]["status"] == "ready_with_warnings"
    assert manifest["submission"]["warning_gates"] == 1


def test_submission_readiness_warns_when_manuscript_polish_guard_rejected_edits(tmp_path: Path) -> None:
    project = Project("submission readiness polish guard", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66. Primary outcome RR 0.74. Primary outcome HR 0.81. Primary outcome MD -1.0.",
            "",
            *_clinical_discussion_fixture(effect="OR 0.66", endpoint="mortality"),
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 1,
            "before": {"ai_style_signal": {"score": 3, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"ai_style_signal": {"score": 1, "issues": [{"code": "low_sentence_length_variation"}]}},
            "proofreading": {
                "enabled": True,
                "status": "ok",
                "provider": "languagetool",
                "issue_count": 2,
                "issues": [
                    {"rule_id": "STYLE_PASSIVE", "message": "Consider direct phrasing."},
                    {"rule_id": "EN_AGREEMENT", "message": "Possible agreement issue."},
                ],
            },
            "issues": [{"code": "numeric_tokens_changed", "review_action": "manual_review_required"}],
        },
        subdir="manuscript",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 1/10 vs 2/10.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []}},
        subdir="papers",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "mortality AND treatment")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")
    assert polish_gate["status"] == "warn"
    assert "fact_guard_issues=1" in polish_gate["detail"]
    assert "remaining_style_issues=0" in polish_gate["detail"]
    assert "proofreading_issues=2" in polish_gate["detail"]
    assert any("human confirms" in action for action in polish_gate["next_actions"])
    assert any("proofreading" in action for action in polish_gate["next_actions"])
    assert readiness["passed"] is True
    assert readiness["status"] == "ready_with_warnings"
    assert manifest["submission"]["passed"] is True
    assert manifest["submission"]["status"] == "ready_with_warnings"
    assert manifest["review"]["manuscript_polish_fact_guard_issues"] == 1
    assert manifest["review"]["manuscript_polish_proofreading_issues"] == 2
    assert manifest["review"]["manuscript_polish_review_queue_status"] == "human_review_required"
    assert manifest["review"]["manuscript_polish_manual_review_items"] == 3
    assert any(
        "human confirms" in action
        for action in manifest["review"]["manuscript_polish_next_actions"]
    )


def test_submission_readiness_passes_for_rejected_polish_candidate_when_final_text_is_safe(tmp_path: Path) -> None:
    project = Project("submission readiness safe rejected polish", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66 (95% CI 0.53 to 0.82).",
            "",
            *_clinical_discussion_fixture(effect="OR 0.66", endpoint="mortality"),
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [{"row_id": "S1:0", "study_id": "S1", "outcome_name": "mortality"}],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 1,
            "before": {"ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {"ai_style_signal": {"score": 0, "issues": []}},
            "proofreading": {"enabled": False, "status": "disabled", "provider": "none", "issue_count": 0},
            "issues": [{"code": "protected_terms_changed", "review_action": "manual_review_required"}],
        },
        subdir="manuscript",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "mortality",
                    "source_quote": "Mortality was 1/10 vs 2/10.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "Mortality was 1/10 vs 2/10.", "page_map": []}},
        subdir="papers",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "mortality AND treatment")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")
    assert polish_gate["status"] == "pass"
    assert "fact_guard_issues=1" in polish_gate["detail"]
    assert "manual_review_items=0" in polish_gate["detail"]
    assert polish_gate.get("next_actions", []) == []
    assert readiness["passed"] is True
    assert readiness["status"] == "ready_with_warnings"
    assert manifest["submission"]["status"] == "ready_with_warnings"
    assert manifest["review"]["manuscript_polish_fact_guard_issues"] == 1


def test_submission_readiness_warns_for_low_weight_remaining_polish_style_issue(tmp_path: Path) -> None:
    project = Project("submission readiness minor polish style", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome OR 0.66. Primary outcome RR 0.74. Primary outcome HR 0.81. Primary outcome MD -1.0.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "manuscript_polish_audit.json",
        {
            "enabled": True,
            "language": "en",
            "accepted_chunks": 1,
            "rejected_chunks": 0,
            "before": {"ai_style_signal": {"score": 2, "issues": [{"code": "template_phrase_hits"}]}},
            "after": {
                "ai_style_signal": {
                    "score": 1,
                    "issues": [{"code": "repeated_sentence_starts", "weight": 1}],
                }
            },
            "proofreading": {"enabled": False, "status": "disabled", "provider": "none", "issue_count": 0},
            "issues": [],
        },
        subdir="manuscript",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        manifest = json.loads(zf.read("package_manifest.json"))

    polish_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_polish")
    assert polish_gate["status"] == "warn"
    assert "remaining_style_issues=1" in polish_gate["detail"]
    assert "manual_review_items=1" in polish_gate["detail"]
    assert any("style signals" in action for action in polish_gate["next_actions"])
    assert manifest["review"]["manuscript_polish_remaining_ai_style_issues"] == 1


def test_submission_readiness_does_not_warn_for_non_primary_limited_sources(tmp_path: Path) -> None:
    project = Project("submission readiness non primary context", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n".join([
            "# Manuscript",
            "",
            "## Abstract",
            "Primary outcome HR 0.81 (95% CI 0.74 to 0.88).",
            "",
            "## Results",
            "The primary meta-analysis used two source-verified trial rows [1].",
            "",
            *_clinical_discussion_fixture(effect="HR 0.81", endpoint="primary composite"),
            "## Declarations",
            "### Ethics approval",
            "No new participant data were collected.",
            "### Data and code availability",
            "Aggregate data are included with the supplementary files.",
            "### Funding",
            "No dedicated external funding was recorded.",
            "### Competing interests",
            "No competing interests were recorded.",
            "",
            "## References",
            "[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [
                    {
                        "code": "limited_text_sources_present",
                        "message": "3 retrieved/screened record(s) use limited source text (3 abstract-only).",
                        "scope": "non_primary_records",
                        "action_required": False,
                    }
                ],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "primary composite"}
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    _save_passing_submission_quality_gate(project)
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "primary composite",
                    "source_quote": "The primary composite hazard ratio was 0.81.",
                    "source_quote_match": "The primary composite hazard ratio was 0.81.",
                    "source_quote_verified": True,
                    "extraction_confidence": "high",
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "The primary composite hazard ratio was 0.81.", "page_map": []}},
        subdir="papers",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "heart failure AND SGLT2")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))
        evidence_review = json.loads(zf.read("review/evidence_readiness_review.json"))

    evidence_gate = next(gate for gate in readiness["gates"] if gate["id"] == "evidence_readiness")
    assert evidence_review["summary"]["warnings"] == 1
    assert evidence_gate["status"] == "pass"
    assert readiness["status"] == "ready"
    assert manifest["submission"]["status"] == "ready"
    assert manifest["submission"]["warning_gates"] == 0


def test_submission_readiness_warns_about_short_meta_manuscript_without_blocking(tmp_path: Path) -> None:
    project = Project("short manuscript package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Short meta-analysis",
            "## Abstract\nThe primary outcome HR was 0.81.",
            "## Introduction\nBrief rationale.",
            "## Methods\nBrief methods with a Boolean search strategy.",
            "## Results\nTwo source-verified trials were included.",
            "## Discussion\nBrief interpretation.",
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n[1] Smith J. Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "primary composite"}
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_validation.json",
        {"passed": True, "issues": [], "facts_summary": {"main_word_count": 6123, "report_type": "meta"}},
        subdir="manuscript",
    )
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "primary composite",
                    "source_quote": "The primary composite hazard ratio was 0.81.",
                    "source_quote_match": "The primary composite hazard ratio was 0.81.",
                    "source_quote_verified": True,
                    "requires_review": False,
                    "conflicts": [],
                }
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {"S1": {"full_text": "The primary composite hazard ratio was 0.81.", "page_map": []}},
        subdir="papers",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "heart failure AND SGLT2")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    length_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_length")
    failed_gate_ids = {gate["id"] for gate in readiness["gates"] if not gate.get("passed")}
    assert length_gate["status"] == "warn"
    assert length_gate["passed"] is True
    assert "manuscript_length" not in failed_gate_ids
    assert manifest["submission"]["warning_gates"] >= 1


def test_submission_readiness_blocks_full_length_meta_manuscript_missing_core_article_content(tmp_path: Path) -> None:
    project = Project("missing manuscript content package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# Full length but incomplete meta-analysis",
            "## Abstract\n" + " ".join(["abstract"] * 300),
            "## Introduction\n" + " ".join(["introduction"] * 1500),
            "## Methods\n" + " ".join(["methods"] * 1500),
            "## Results\n" + " ".join(["results"] * 1300),
            "## Discussion\n" + " ".join(["discussion"] * 1600),
            "## Declarations\n### Ethics approval\nNo new participant data were collected.\n### Data and code availability\nAggregate data are included with the supplementary files.\n### Funding\nNo dedicated external funding was recorded.\n### Competing interests\nNo competing interests were recorded.",
            "## References\n",
        ]),
        subdir="manuscript",
    )
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "primary_effect": {"n_studies": 2},
            "evidence_readiness": {
                "status": "ready",
                "blockers": [],
                "warnings": [],
                "selected_primary_rows": [
                    {"row_id": "S1:0", "study_id": "S1", "outcome_name": "primary composite"},
                    {"row_id": "S2:0", "study_id": "S2", "outcome_name": "primary composite"},
                ],
            },
        },
        subdir="manuscript",
    )
    project.save_json("manuscript_validation.json", {"passed": True, "issues": []}, subdir="manuscript")
    project.save_json(
        "extraction_audit.json",
        {
            "rows": [
                {
                    "row_id": "S1:0",
                    "study_id": "S1",
                    "outcome_index": 0,
                    "outcome_name": "primary composite",
                    "source_quote": "Trial one hazard ratio was 0.81.",
                    "source_quote_match": "Trial one hazard ratio was 0.81.",
                    "source_quote_verified": True,
                    "requires_review": False,
                    "conflicts": [],
                },
                {
                    "row_id": "S2:0",
                    "study_id": "S2",
                    "outcome_index": 0,
                    "outcome_name": "primary composite",
                    "source_quote": "Trial two hazard ratio was 0.79.",
                    "source_quote_match": "Trial two hazard ratio was 0.79.",
                    "source_quote_verified": True,
                    "requires_review": False,
                    "conflicts": [],
                },
            ],
        },
        subdir="extraction",
    )
    project.save_json(
        "parsed_papers.json",
        {
            "S1": {"full_text": "Trial one hazard ratio was 0.81.", "page_map": []},
            "S2": {"full_text": "Trial two hazard ratio was 0.79.", "page_map": []},
        },
        subdir="papers",
    )
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_ready",
            "summary_card": {"benchmark_id": "mini_ready", "status": "passed", "passed": True, "failing_gates": []},
        },
        subdir="benchmark",
    )
    project.save_text("references.bib", "@article{smith2024,title={Trial report}}")
    project.save_text("search_query.txt", "heart failure AND SGLT2")
    figures_dir = project.base_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    (figures_dir / "forest_plot.png").write_bytes(MINIMAL_PNG)

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        readiness = json.loads(zf.read("review/submission_readiness_review.json"))

    content_gate = next(gate for gate in readiness["gates"] if gate["id"] == "manuscript_content")
    assert content_gate["status"] == "fail"
    assert manifest["manuscript"]["reference_count"] == 0
    assert manifest["manuscript"]["table_count"] == 0
    assert manifest["manuscript"]["figure_count"] == 0
    assert manifest["manuscript"]["has_search_query_in_manuscript"] is False
    assert manifest["manuscript"]["has_calculation_detail"] is False
    assert readiness["status"] == "blocked"
    assert manifest["submission"]["passed"] is False


def test_manuscript_content_summary_counts_unique_numbered_figures() -> None:
    manuscript = """
## Results

### Figure 1. Forest plot

![Figure 1. Forest plot](figures/forest_plot.png)

### Figure 2. Funnel plot

![Funnel plot](figures/funnel_plot.png)
"""

    summary = _manuscript_content_summary(manuscript)

    assert summary["figure_count"] == 2


def test_artifact_package_manifest_counts_protocol_adjudication_tasks(tmp_path: Path) -> None:
    project = Project("protocol adjudication package", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# Manuscript", subdir="manuscript")
    project.save_json(
        "benchmark_report.json",
        {
            "benchmark_id": "mini_benchmark",
            "summary_card": {
                "benchmark_id": "mini_benchmark",
                "status": "failed",
                "passed": False,
                "published_anchor": {
                    "effect_measure": "OR",
                    "model_preference": "fixed",
                    "effect": 0.66,
                    "ci_lower": 0.53,
                    "ci_upper": 0.82,
                },
                "observed_primary": {
                    "effect_measure": "RR",
                    "model_preference": "random",
                    "effect": 0.73,
                    "ci_lower": 0.62,
                    "ci_upper": 0.86,
                },
                "failing_gates": [
                    {
                        "gate": "pooled_effect",
                        "passed": False,
                        "failure_reasons": ["effect_measure_mismatch", "model_preference_mismatch"],
                    }
                ],
            },
            "pooled_effect": {
                "expected_effect_measure": "OR",
                "observed_effect_measure": "RR",
                "expected_model_preference": "fixed",
                "observed_model_preference": "random",
                "failure_reasons": ["effect_measure_mismatch", "model_preference_mismatch"],
            },
        },
        subdir="benchmark",
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        manifest = json.loads(zf.read("package_manifest.json"))
        benchmark_review = json.loads(zf.read("review/benchmark_review.json"))

    assert manifest["review"]["benchmark_protocol_adjudication_tasks"] == 1
    assert benchmark_review["protocol_adjudication_tasks"][0]["suggested_protocol_patch"]["fields"] == {
        "effect_measure": "OR",
        "model_preference": "fixed",
    }


def test_artifact_package_includes_pdf_intake_review_for_failed_uploads(tmp_path: Path) -> None:
    project = Project("pdf intake package", output_dir=tmp_path / uuid4().hex)
    project.save_text("draft.md", "# Manuscript", subdir="manuscript")
    project.save_json(
        "pdf_intake_manifest.json",
        {
            "session_id": "upload-session-1",
            "created_at": "2026-05-22T12:00:00+00:00",
            "files": [
                {
                    "filename": "good.pdf",
                    "local_path": "/tmp/good.pdf",
                    "file_size_bytes": 1234,
                    "sha256": "abc",
                    "parse_status": "ok",
                    "parse_error": None,
                    "parser_used": "pdf_parser",
                    "parser_cache_version": "v1",
                    "cache_hit": False,
                    "page_count": 8,
                    "text_chars": 12000,
                    "table_count": 3,
                    "empty_pages": [],
                    "matched_pmid": "12345",
                    "matched_title": "Good Trial",
                    "match_score": 0.98,
                    "match_method": "filename_pmid",
                    "requires_user_review": False,
                },
                {
                    "filename": "broken.pdf",
                    "local_path": "/tmp/broken.pdf",
                    "file_size_bytes": 50,
                    "sha256": "def",
                    "parse_status": "failed",
                    "parse_error": "PDF syntax error",
                    "parser_used": "pdf_parser",
                    "parser_cache_version": "v1",
                    "cache_hit": False,
                    "page_count": 0,
                    "text_chars": 0,
                    "table_count": 0,
                    "empty_pages": [],
                    "matched_pmid": None,
                    "matched_title": None,
                    "match_score": None,
                    "match_method": None,
                    "requires_user_review": True,
                },
            ],
        },
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        names = set(zf.namelist())
        intake_review = json.loads(zf.read("review/pdf_intake_review.json"))
        intake_html = zf.read("review/pdf_intake_review.html").decode("utf-8")
        manifest = json.loads(zf.read("package_manifest.json"))

    assert "pdf_intake_manifest.json" in names
    assert "review/pdf_intake_review.json" in names
    assert "review/pdf_intake_review.html" in names
    assert manifest["review"]["pdf_intake_included"] is True
    assert manifest["review"]["pdf_intake_total_files"] == 2
    assert manifest["review"]["pdf_intake_failed"] == 1
    assert manifest["review"]["pdf_intake_requires_review"] == 1
    assert intake_review["summary"]["ok"] == 1
    assert intake_review["summary"]["failed"] == 1
    assert intake_review["files"][1]["parse_error"] == "PDF syntax error"
    assert "MetaAgent PDF Intake Review" in intake_html
    assert "broken.pdf" in intake_html
    assert "PDF syntax error" in intake_html


def test_pdf_intake_review_localizes_chinese_handoff_surface(tmp_path: Path) -> None:
    project = Project("chinese pdf intake package", output_dir=tmp_path / uuid4().hex)
    project.save_text(
        "draft.md",
        "\n\n".join([
            "# 中文Meta分析稿件",
            "## 摘要\n评价治疗是否改善主要结局。",
            "## 方法\n用户上传全文后，需要确认 PDF 是否解析成功。",
            "## 结果\n两篇原文进入数据提取。",
            "## 讨论\n解析失败的原文必须人工复核。",
            "## 参考文献\n［1］ Trial report.",
        ]),
        subdir="manuscript",
    )
    project.save_json("manuscript_output_language.json", {"expected_language": "zh"}, subdir="manuscript")
    project.save_json(
        "manuscript_facts.json",
        {
            "report_type": "meta",
            "output_language": "zh",
        },
        subdir="manuscript",
    )
    project.save_json(
        "pdf_intake_manifest.json",
        {
            "session_id": "upload-session-zh",
            "created_at": "2026-05-24T12:00:00+08:00",
            "files": [
                {
                    "filename": "parsed.pdf",
                    "local_path": "/tmp/parsed.pdf",
                    "download_status": "ok",
                    "download_error": None,
                    "parse_status": "ok",
                    "parse_error": None,
                    "parser_used": "pdf_parser",
                    "cache_hit": True,
                    "ocr_used": False,
                    "page_count": 12,
                    "text_chars": 32000,
                    "table_count": 4,
                    "matched_pmid": "123",
                    "matched_title": "Parsed Trial",
                    "match_score": 0.99,
                    "match_method": "filename_title",
                    "requires_user_review": False,
                },
                {
                    "filename": "failed.pdf",
                    "local_path": "/tmp/failed.pdf",
                    "download_status": "ok",
                    "download_error": None,
                    "parse_status": "failed",
                    "parse_error": "PDF syntax error",
                    "parser_used": "pdf_parser",
                    "cache_hit": False,
                    "ocr_used": True,
                    "page_count": 0,
                    "text_chars": 0,
                    "table_count": 0,
                    "matched_pmid": None,
                    "matched_title": None,
                    "match_score": None,
                    "match_method": None,
                    "requires_user_review": True,
                },
            ],
        },
    )

    package_path = create_artifact_package(project)

    with zipfile.ZipFile(package_path) as zf:
        intake_review = json.loads(zf.read("review/pdf_intake_review.json"))
        intake_html = zf.read("review/pdf_intake_review.html").decode("utf-8")

    assert intake_review["language"] == "zh"
    assert '<html lang="zh">' in intake_html
    assert "MetaAgent PDF 原文接收复核" in intake_html
    assert "展示用户上传 PDF 的下载、解析、OCR、缓存和人工复核状态" in intake_html
    assert "文件总数" in intake_html
    assert "解析成功" in intake_html
    assert "需人工复核" in intake_html
    assert "缓存命中" in intake_html
    assert "文件明细" in intake_html
    assert "下载" in intake_html
    assert "解析" in intake_html
    assert "匹配标题" in intake_html
    assert "是" in intake_html
    assert "否" in intake_html
    assert "MetaAgent PDF Intake Review" not in intake_html
    assert "Cache hits" not in intake_html
