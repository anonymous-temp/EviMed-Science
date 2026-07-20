"""Submission document export helpers."""
from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Inches

from new_meta.core.project import Project


IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")


def export_manuscript_docx(project: Project, output_name: str = "draft.docx") -> Path | None:
    """Export manuscript/draft.md to a Word document for submission handoff."""
    draft_path = project.base_dir / "manuscript" / "draft.md"
    if not draft_path.exists() or draft_path.stat().st_size <= 0:
        return None

    document = Document()
    _apply_document_defaults(document)
    lines = draft_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        heading = _heading_level(line)
        if heading:
            level, text = heading
            document.add_heading(_clean_inline(text), level=level)
            i += 1
            continue

        if _is_table_start(lines, i):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table(document, table_lines)
            continue

        image = IMAGE_RE.search(line)
        if image:
            _add_image(document, project, image.group("path"), image.group("alt"))
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if (
                not next_line.strip()
                or _heading_level(next_line)
                or _is_table_start(lines, i)
                or IMAGE_RE.search(next_line)
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        document.add_paragraph(_clean_inline(" ".join(part.strip() for part in paragraph_lines)))

    output_path = project.base_dir / "manuscript" / output_name
    document.save(output_path)
    return output_path


def export_manuscript_pdf(project: Project, output_name: str = "draft.pdf") -> Path | None:
    """Export manuscript/draft.md to a lightweight PDF handoff file."""
    draft_path = project.base_dir / "manuscript" / "draft.md"
    if not draft_path.exists() or draft_path.stat().st_size <= 0:
        return None

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return None

    output_path = project.base_dir / "manuscript" / output_name
    styles = getSampleStyleSheet()
    for name in ("Normal", "BodyText"):
        styles[name].fontName = "Times-Roman"
        styles[name].fontSize = 10
        styles[name].leading = 13
    for name in ("Title", "Heading1", "Heading2", "Heading3", "Heading4"):
        if name in styles:
            styles[name].fontName = "Times-Bold"

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=_clean_inline(project.topic or "MetaAgent manuscript"),
    )

    story = []
    lines = draft_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        heading = _heading_level(line)
        if heading:
            level, text = heading
            style_name = "Title" if level == 1 else f"Heading{min(level, 4)}"
            story.append(Paragraph(_pdf_text(text), styles[style_name]))
            story.append(Spacer(1, 0.10 * inch))
            i += 1
            continue

        if _is_table_start(lines, i):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_pdf_markdown_table(story, table_lines, styles)
            continue

        image = IMAGE_RE.search(line)
        if image:
            _add_pdf_image(
                story,
                project,
                image.group("path"),
                image.group("alt"),
                styles,
                max_width=6.0 * inch,
                max_height=7.0 * inch,
            )
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if (
                not next_line.strip()
                or _heading_level(next_line)
                or _is_table_start(lines, i)
                or IMAGE_RE.search(next_line)
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        story.append(Paragraph(_pdf_text(" ".join(part.strip() for part in paragraph_lines)), styles["BodyText"]))
        story.append(Spacer(1, 0.08 * inch))

    document.build(story)
    return output_path


def _apply_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = None


def _heading_level(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
    if not match:
        return None
    return min(len(match.group(1)), 4), match.group(2)


def _is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return lines[index].strip().startswith("|") and bool(
        re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[index + 1])
    )


def _add_markdown_table(document: Document, table_lines: list[str]) -> None:
    rows = [_split_table_row(line) for line in table_lines]
    rows = [row for idx, row in enumerate(rows) if idx != 1 and row]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            table.cell(row_idx, col_idx).text = _clean_inline(row[col_idx]) if col_idx < len(row) else ""


def _add_pdf_markdown_table(story: list, table_lines: list[str], styles) -> None:
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer, Table as PdfTable, TableStyle

    rows = [_split_table_row(line) for line in table_lines]
    rows = [row for idx, row in enumerate(rows) if idx != 1 and row]
    if not rows:
        return
    width = max(len(row) for row in rows)
    data = []
    for row in rows:
        data.append([
            Paragraph(_pdf_text(row[col_idx]), styles["BodyText"]) if col_idx < len(row) else ""
            for col_idx in range(width)
        ])
    table = PdfTable(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.12 * inch))


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _add_image(document: Document, project: Project, raw_path: str, alt: str) -> None:
    image_path = _resolve_markdown_path(project, raw_path)
    if image_path and image_path.exists() and image_path.stat().st_size > 0:
        try:
            document.add_picture(str(image_path), width=Inches(6.0))
        except Exception:
            document.add_paragraph(f"[Image could not be embedded: {_clean_inline(alt or raw_path)}]")
    if alt:
        document.add_paragraph(_clean_inline(alt))


def _add_pdf_image(
    story: list,
    project: Project,
    raw_path: str,
    alt: str,
    styles,
    *,
    max_width: float,
    max_height: float,
) -> None:
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as PdfImage, Paragraph, Spacer

    image_path = _resolve_markdown_path(project, raw_path)
    if image_path and image_path.exists() and image_path.stat().st_size > 0:
        try:
            image = PdfImage(str(image_path))
            if image.drawWidth > 0 and image.drawHeight > 0:
                scale = min(1.0, max_width / image.drawWidth, max_height / image.drawHeight)
                image.drawWidth *= scale
                image.drawHeight *= scale
            story.append(image)
        except Exception:
            story.append(Paragraph(_pdf_text(f"[Image could not be embedded: {alt or raw_path}]"), styles["BodyText"]))
    if alt:
        story.append(Paragraph(_pdf_text(alt), styles["BodyText"]))
    story.append(Spacer(1, 0.10 * inch))


def _resolve_markdown_path(project: Project, raw_path: str) -> Path | None:
    path_text = raw_path.strip().split("#", 1)[0]
    if not path_text:
        return None
    if path_text.startswith("data:image/"):
        return None
    path = Path(path_text)
    if path.is_absolute():
        return path
    return (project.base_dir / "manuscript" / path).resolve()


def _clean_inline(text: str) -> str:
    cleaned = text.replace("**", "").replace("__", "")
    cleaned = re.sub(r"(?<!\*)\*(?!\*)(.*?)\*(?!\*)", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", cleaned)
    return cleaned.strip()


def _pdf_text(text: str) -> str:
    return xml_escape(_clean_inline(text)).replace("\n", "<br/>")
