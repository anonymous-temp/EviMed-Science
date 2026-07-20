# [IN] paper/sections.py, paper/references.py, models.py
# [OUT] Complete paper document
# [POS] mr_agent/paper/generator.py - Full paper generation engine
"""Full academic paper generator for MR studies."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Callable

from mr_agent.llm.client import LLMClient
from mr_agent.models import MRAnalysisResult, PaperReference, SessionState, find_ivw
from mr_agent.paper import references, sections
from mr_agent.tools import pubmed

logger = logging.getLogger(__name__)

ProgressCallback = Callable[[str, float], None]

SECTION_ORDER = [
    "title", "abstract", "introduction", "methods", "results",
    "discussion", "limitations", "conclusion",
    "data_availability", "ethics_statement",
    "table1", "table2", "references",
]


class PaperGenerator:
    """Generates a complete STROBE-MR compliant paper."""

    def __init__(
        self,
        llm: LLMClient,
        state: SessionState,
        on_progress: ProgressCallback | None = None,
        language: str = "en",
    ):
        self.llm = llm
        self.state = state
        self.on_progress = on_progress or (lambda msg, pct: None)
        self.language = language
        self.strobe_report = None

    def generate_outline(self, feedback: str = "") -> dict:
        """Generate paper outline for user confirmation."""
        exp, out = self.state.slots.exposure, self.state.slots.outcome
        results = self.state.analysis_results
        return sections.generate_outline(
            self.llm, exp, out, results,
            bidirectional=self.state.slots.bidirectional,
            feedback=feedback,
        )

    def format_outline(self, outline: dict) -> str:
        """Format outline dict as readable text."""
        lines = [f"Title: {outline.get('title', '')}"]
        for sec in outline.get("sections", []):
            lines.append(f"\n## {sec['name']}")
            for point in sec.get("key_points", []):
                lines.append(f"  - {point}")
        if outline.get("tables"):
            lines.append("\n## Tables")
            for t in outline["tables"]:
                lines.append(f"  - {t}")
        if outline.get("figures"):
            lines.append("\n## Figures")
            for f in outline["figures"]:
                lines.append(f"  - {f}")
        return "\n".join(lines)

    def generate(self) -> dict[str, str]:
        """Generate all paper sections with multi-pass review."""
        self.on_progress("收集参考文献...", 0.0)
        refs = self._gather_references()
        self.state.references = refs
        body = self._write_body_sections(refs)
        tables = self._write_tables()
        self.on_progress("一致性审查...", 0.82)
        body = self._review_coherence(body)
        self.on_progress("STROBE-MR合规检查...", 0.85)
        body = self._fix_strobe_gaps(body)
        abstract = self._write_connected_abstract(body)
        extras = self._write_extra_sections()
        paper = self._assemble_dict(abstract, body, tables, extras, refs)
        paper = self._enforce_numerical_consistency(paper)
        # The LLM is useful for narrative synthesis, but Methods, Results, core
        # tables, and Limitations are release-critical factual surfaces.  Build
        # those sections from the structured runtime result so a fluent rewrite
        # cannot invent cohort overlap, intermediate SNP counts, sensitivity
        # analyses, or metadata that the analysis did not produce.
        paper = self._enforce_structured_grounding(paper)
        paper = references.apply_citation_numbers(paper, refs)
        # Fallback: if all LLM calls failed, generate deterministic report from data
        non_empty = sum(1 for v in paper.values() if v and v.strip())
        if non_empty <= 1:
            logger.warning("所有 LLM 论文生成调用均失败，使用确定性数据回退报告")
            paper = self._deterministic_fallback_paper(paper)
        self.state.paper_sections = paper
        self.on_progress("论文撰写完成!", 1.0)
        return paper

    def _write_body_sections(self, refs: list[PaperReference]) -> dict[str, str]:
        """Write intro, methods, results, discussion, limitations, conclusion."""
        exp, out = self.state.slots.exposure, self.state.slots.outcome
        results = self.state.analysis_results
        lang = self.language
        self.on_progress("撰写引言...", 0.10)
        intro = self._safe_write(
            sections.write_introduction, self.llm, exp, out, refs, language=lang
        )
        self.on_progress("撰写方法...", 0.25)
        methods = self._safe_write(
            sections.write_methods, self.llm, results, exp, out, language=lang
        )
        self.on_progress("撰写结果...", 0.40)
        results_text = self._safe_write(
            sections.write_results, self.llm, results, language=lang
        )
        self.on_progress("撰写讨论...", 0.55)
        discussion = self._safe_write(
            sections.write_discussion, self.llm, exp, out, results, refs, language=lang,
        )
        self.on_progress("撰写局限性...", 0.65)
        populations = list(dict.fromkeys(
            str(result.exposure_metadata.get("population") or "").strip()
            for result in results
            if result.exposure_metadata.get("population")
        ))
        limitations = self._safe_write(
            sections.write_limitations,
            self.llm, exp, out, results,
            bidirectional=self.state.slots.bidirectional,
            population=self.state.slots.population or ", ".join(populations) or "N/A",
            language=lang,
        )
        self.on_progress("撰写结论...", 0.72)
        conclusion = self._safe_write(
            sections.write_conclusion, self.llm, exp, out, results, language=lang,
        )
        return {
            "introduction": intro, "methods": methods,
            "results": results_text, "discussion": discussion,
            "limitations": limitations, "conclusion": conclusion,
        }

    def _zh_system(self) -> str:
        """Build system prompt with language requirement enforced."""
        from mr_agent.llm.prompts import SYSTEM_MR_SCIENTIST
        if self.language == "zh":
            return (
                SYSTEM_MR_SCIENTIST
                + "\n\n【强制语言要求】你必须用中文撰写所有正文内容。"
                "以下专有名词保留英文：基因名、统计方法缩写（IVW、MR-Egger、SNP、OR、CI、"
                "beta、p值、F统计量等）、GWAS数据集ID。"
                "除上述专有名词外，所有句子、段落必须为中文，禁止出现完整英文句子。"
            )
        return (
            SYSTEM_MR_SCIENTIST
            + "\n\nLANGUAGE REQUIREMENT (strictly enforced): "
            "Write ENTIRELY in English. "
            "Do NOT include any Chinese characters anywhere in your response."
        )

    def _zh_prompt(self, prompt: str) -> str:
        """Append language instruction to prompt."""
        if self.language == "zh":
            return (
                prompt
                + "\n\n【重要】请完全用中文撰写，专有名词（IVW、SNP、OR、CI等缩写）保留英文，"
                "其余所有正文内容必须为中文，不得出现完整英文句子。"
            )
        return (
            prompt
            + "\n\nIMPORTANT: Write entirely in English. No Chinese characters."
        )

    def _llm_chat(self, prompt: str, max_tokens: int = 3000) -> str:
        """LLM call with language-appropriate system prompt and instruction."""
        from mr_agent.llm.prompts import SYSTEM_MR_SCIENTIST
        return self.llm.chat(
            messages=[{"role": "user", "content": self._zh_prompt(prompt)}],
            system=self._zh_system(),
            max_tokens=max_tokens,
        )

    def _safe_write(self, fn, *args, **kwargs) -> str:
        """Call a section writer; return empty string on failure."""
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            logger.warning(f"Section write failed ({fn.__name__}): {e}")
            return ""

    def _write_tables(self) -> dict[str, str]:
        """Write Table 1 and Table 2."""
        exp, out = self.state.slots.exposure, self.state.slots.outcome
        results = self.state.analysis_results
        lang = self.language
        self.on_progress("生成表格...", 0.78)
        table1 = self._safe_write(
            sections.write_table1, self.llm, results, exp, out, language=lang,
        )
        table2 = self._safe_write(
            sections.write_table2, self.llm, results, language=lang,
        )
        return {"table1": table1, "table2": table2}

    def _write_connected_abstract(self, body: dict[str, str]) -> str:
        """Write abstract from actual section content (not hardcoded briefs)."""
        self.on_progress("撰写摘要...", 0.88)
        from mr_agent.paper.sections import _format_key_findings
        results = self.state.analysis_results
        key_data = _format_key_findings(results)
        return self._safe_write(
            sections.write_abstract_connected,
            self.llm,
            intro_text=body.get("introduction", "")[:500],
            methods_text=body.get("methods", "")[:500],
            results_text=body.get("results", "")[:500],
            conclusion_text=body.get("conclusion", "")[:300],
            key_data=key_data,
            language=self.language,
        )

    def _review_coherence(self, body: dict[str, str]) -> dict[str, str]:
        """Multi-pass: review paper for coherence, apply feedback."""
        from mr_agent.llm.prompts import PAPER_COHERENCE_REVIEW
        try:
            sections_text = "\n\n".join(
                f"## {k.upper()}\n{v[:800]}" for k, v in body.items() if v
            )
            feedback = self._llm_chat(
                PAPER_COHERENCE_REVIEW.format(sections_text=sections_text),
                max_tokens=1000,
            )
            if "no inconsistencies" in feedback.lower():
                return body
            logger.info(f"Coherence review feedback: {feedback[:300]}")
            affected = self._detect_affected_sections(feedback, body)
            for section_name in affected:
                body[section_name] = self._revise_section(
                    section_name, body[section_name], feedback,
                )
        except Exception as e:
            logger.warning(f"Coherence review failed, skipping: {e}")
        return body

    # Sections that need numerical consistency checking (NOT results/table2)
    _NUM_CHECK_SECTIONS = ("abstract", "discussion", "conclusion", "limitations")

    def _deterministic_fallback_paper(self, paper: dict[str, str]) -> dict[str, str]:
        """Generate a minimal but complete paper from raw analysis data (no LLM).

        Used when all LLM calls fail (e.g. API unreachable).
        """
        from mr_agent.paper.sections import _build_results_summary, _format_key_findings
        exp = self.state.slots.exposure or "exposure"
        out = self.state.slots.outcome or "outcome"
        results = self.state.analysis_results
        zh = self.language == "zh"

        if zh:
            title = f"{exp}与{out}的因果关系：一项孟德尔随机化研究"
            n_pairs = len(results)
            n_ivs = sum(r.n_instruments for r in results) if results else 0
            abstract = (
                f"**背景**：探讨{exp}与{out}之间的因果关系。\n"
                f"**方法**：采用两样本孟德尔随机化（Two-Sample MR）方法，"
                f"共纳入{n_pairs}个暴露-结局分析对，使用{n_ivs}个工具变量（SNP）。"
                f"主要分析方法包括IVW、MR-Egger和加权中位数法。\n"
                f"**结果**：{_format_key_findings(results) if results else '详见结果章节。'}\n"
                f"**结论**：基于孟德尔随机化分析的证据提示了{exp}对{out}的潜在因果关系。"
            )
            introduction = (
                f"孟德尔随机化（Mendelian Randomization, MR）是一种利用遗传变异作为工具变量"
                f"来评估暴露因素与结局之间因果关系的流行病学方法。由于遗传变异在受孕时随机分配，"
                f"且不随环境因素改变，MR分析可以有效避免传统观察性研究中的混杂因素和反向因果关系。"
                f"\n\n本研究旨在利用两样本MR方法，评估{exp}对{out}的因果效应。"
            )
            methods = (
                f"本研究采用两样本孟德尔随机化方法。\n\n"
                f"**数据来源**：GWAS汇总统计数据来自公开数据库（IEU OpenGWAS）。\n"
                f"**工具变量选择**：以P < 5×10⁻⁸为阈值筛选与{exp}显著相关的SNP作为工具变量，"
                f"并去除连锁不平衡（LD clumping, r² < 0.001）。\n"
                f"**分析方法**：主要采用逆方差加权法（IVW），辅以MR-Egger回归、"
                f"加权中位数法（Weighted Median）进行敏感性分析。\n"
                f"**异质性与多效性检验**：使用Cochran's Q检验评估异质性，"
                f"MR-Egger截距检验评估水平多效性，MR-PRESSO检测潜在异常值。"
            )
        else:
            title = (
                f"Causal relationship between {exp} and {out}: "
                f"a Mendelian randomization study"
            )
            n_pairs = len(results)
            n_ivs = sum(r.n_instruments for r in results) if results else 0
            abstract = (
                f"**Background**: To investigate the causal relationship between "
                f"{exp} and {out}.\n"
                f"**Methods**: Two-sample Mendelian randomization (MR) was performed "
                f"using {n_pairs} exposure-outcome pairs with {n_ivs} instrumental variables (SNPs). "
                f"Primary analysis used IVW, with MR-Egger and Weighted Median as sensitivity analyses.\n"
                f"**Results**: {_format_key_findings(results) if results else 'See Results section.'}\n"
                f"**Conclusion**: MR evidence suggests a potential causal effect of "
                f"{exp} on {out}."
            )
            introduction = (
                f"Mendelian Randomization (MR) is an epidemiological method that uses "
                f"genetic variants as instrumental variables to assess causal relationships "
                f"between exposures and outcomes. Since genetic variants are randomly assigned "
                f"at conception and remain constant throughout life, MR can avoid confounding "
                f"and reverse causation inherent in observational studies.\n\n"
                f"This study aims to evaluate the causal effect of {exp} on {out} "
                f"using a two-sample MR approach."
            )
            methods = (
                f"A two-sample Mendelian randomization approach was employed.\n\n"
                f"**Data sources**: GWAS summary statistics were obtained from "
                f"public databases (IEU OpenGWAS).\n"
                f"**Instrument selection**: SNPs significantly associated with {exp} "
                f"(P < 5×10⁻⁸) were selected as instrumental variables, with LD clumping "
                f"(r² < 0.001) to ensure independence.\n"
                f"**Analysis methods**: Inverse variance weighted (IVW) was used as the "
                f"primary method, supplemented by MR-Egger regression and Weighted Median "
                f"for sensitivity analysis.\n"
                f"**Heterogeneity and pleiotropy**: Cochran's Q test assessed heterogeneity, "
                f"MR-Egger intercept test evaluated horizontal pleiotropy, and MR-PRESSO "
                f"detected potential outliers."
            )

        results_text = _build_results_summary(results) if results else ""
        has_significant_ivw = any(
            (ivw := find_ivw(result.mr_results or [])) is not None
            and ivw.pval is not None
            and ivw.pval < 0.05
            for result in results
        )

        if zh:
            discussion = (
                f"本研究通过孟德尔随机化方法评估了{exp}对{out}的因果效应。"
                f"IVW分析结果显示{'存在' if has_significant_ivw else '未发现'}显著的因果关联。\n\n"
                f"本研究结果需要结合现有流行病学证据进行综合解读。"
                f"MR分析的优势在于利用遗传变异作为工具变量，减少了混杂因素的影响；"
                f"但同时也存在局限性，包括潜在的水平多效性、工具变量强度不足等。"
            )
            conclusion = (
                f"基于孟德尔随机化分析，本研究{'发现' if has_significant_ivw else '未发现'}"
                f"{exp}对{out}{'存在显著的因果关联' if has_significant_ivw else '的显著因果关联'}。"
                f"结果需要在更大样本量和不同人群的研究中进一步验证。"
            )
            limitations = (
                "本研究存在以下局限性：(1) MR分析假设工具变量满足相关性、独立性和排他性三大假设，"
                "可能存在残余的水平多效性；(2) 分析主要基于欧洲人群GWAS数据，结果外推至其他人群需谨慎；"
                "(3) 无法排除终生效应与短期效应的差异。"
            )
        else:
            discussion = (
                f"This study evaluated the causal effect of {exp} on {out} using "
                f"Mendelian randomization. The IVW analysis "
                f"{'identified' if has_significant_ivw else 'did not identify'} "
                f"a significant causal association.\n\n"
                f"These findings should be interpreted alongside existing epidemiological evidence. "
                f"MR analysis leverages genetic variants to reduce confounding, but has limitations "
                f"including potential horizontal pleiotropy and insufficient instrument strength."
            )
            conclusion = (
                f"Based on Mendelian randomization analysis, this study "
                f"{'found' if has_significant_ivw else 'did not find'} "
                f"a significant causal relationship between {exp} and {out}. "
                f"Further validation in larger samples and diverse populations is warranted."
            )
            limitations = (
                "This study has several limitations: (1) MR assumes relevance, independence, "
                "and exclusion restriction for instrumental variables, but residual horizontal "
                "pleiotropy may exist; (2) analyses are primarily based on European-ancestry "
                "GWAS data, limiting generalizability; (3) lifelong genetic effects may differ "
                "from short-term interventions."
            )

        # Only override empty sections — preserve any LLM-generated content
        fallback = {
            "title": title,
            "abstract": abstract,
            "introduction": introduction,
            "methods": methods,
            "results": results_text,
            "discussion": discussion,
            "limitations": limitations,
            "conclusion": conclusion,
            "data_availability": (
                "Data and code available from the corresponding author upon reasonable request."
            ),
            "ethics_statement": (
                "This study uses publicly available summary-level data; no individual-level "
                "data was used, thus no ethical approval was required."
            ),
        }
        for key, value in fallback.items():
            if not paper.get(key, "").strip():
                paper[key] = value

        return paper

    def _enforce_numerical_consistency(self, paper: dict[str, str]) -> dict[str, str]:
        """Post-generation pass: force all numbers to match the Results section.

        Uses the raw MR analysis data as the single source of truth.
        For each section in _NUM_CHECK_SECTIONS, sends the text and canonical
        data to the LLM and asks it to fix/delete any inconsistent values.
        """
        from mr_agent.llm.prompts import PAPER_NUMBER_CONSISTENCY
        from mr_agent.paper.sections import _build_results_summary, _format_key_findings
        results = self.state.analysis_results
        if not results:
            return paper
        # Build canonical data: combine detailed summary + key findings
        canonical = (
            _build_results_summary(results)
            + "\n\n---\n\n"
            + _format_key_findings(results)
        )
        for sec_name in self._NUM_CHECK_SECTIONS:
            text = paper.get(sec_name, "")
            if not text or len(text) < 50:
                continue
            try:
                prompt = PAPER_NUMBER_CONSISTENCY.format(
                    canonical_data=canonical,
                    section_name=sec_name,
                    section_text=text,
                )
                revised = self._llm_chat(prompt, max_tokens=4000)
                if revised and len(revised) > len(text) * 0.5:
                    paper[sec_name] = revised
                    logger.info(f"数值一致性修正: {sec_name}")
            except Exception as e:
                logger.warning(f"数值一致性修正失败 ({sec_name}): {e}")
        return paper

    @staticmethod
    def _metadata_value(metadata: dict, key: str) -> str:
        value = metadata.get(key)
        if value is None or str(value).strip().casefold() in {"", "na", "n/a", "unknown", "none"}:
            return "N/A"
        if key in {"sample_size", "nsnp"}:
            try:
                return f"{int(value):,}"
            except (TypeError, ValueError):
                pass
        return str(value)

    @staticmethod
    def _fmt_p(value: float) -> str:
        return f"{value:.3e}"

    def _grounded_methods(self, results: list[MRAnalysisResult]) -> str:
        """Return methods that contain only operations and metadata we observed."""
        zh = self.language == "zh"
        blocks: list[str] = []
        for index, result in enumerate(results, start=1):
            design_section = index * 2 - 1
            analysis_section = index * 2
            exp = result.exposure_name or result.exposure_id
            out = result.outcome_name or result.outcome_id
            em = result.exposure_metadata
            om = result.outcome_metadata
            methods = ", ".join(item.method for item in result.mr_results) or "N/A"
            sensitivity: list[str] = []
            if result.heterogeneity:
                sensitivity.append("Cochran's Q")
            if result.pleiotropy:
                sensitivity.append("MR-Egger intercept")
            if result.steiger_correct is not None:
                sensitivity.append("Steiger directionality")
            if result.presso_global_pval is not None or result.presso_n_outliers is not None:
                sensitivity.append("MR-PRESSO")
            if result.plots.get("loo_plot_png") or result.plots.get("loo_plot_pdf"):
                sensitivity.append("leave-one-out plot")
            sensitivity_text = ", ".join(sensitivity) or "N/A"
            threshold = f"{result.pval_threshold:.1e}"
            if zh:
                overlap = (
                    "源元数据不能确认两个GWAS队列完全不重叠；运行时已标记潜在样本重叠，"
                    "因此不能把两样本设计等同于已证实的独立队列。"
                    if result.sample_overlap_warning else
                    "源元数据未提供足以验证样本是否重叠的队列级信息，故不作无重叠声明。"
                )
                blocks.append(
                    f"### {design_section}. 研究设计与数据源\n"
                    f"本分析采用两样本孟德尔随机化设计，评估{exp}对{out}的效应。"
                    f"暴露GWAS为{result.exposure_id}（性状：{self._metadata_value(em, 'trait')}；"
                    f"样本量：{self._metadata_value(em, 'sample_size')}；人群：{self._metadata_value(em, 'population')}；"
                    f"年份：{self._metadata_value(em, 'year')}；SNP总数：{self._metadata_value(em, 'nsnp')}）。"
                    f"结局GWAS为{result.outcome_id}（性状：{self._metadata_value(om, 'trait')}；"
                    f"样本量：{self._metadata_value(om, 'sample_size')}；人群：{self._metadata_value(om, 'population')}；"
                    f"年份：{self._metadata_value(om, 'year')}；SNP总数：{self._metadata_value(om, 'nsnp')}）。{overlap}\n\n"
                    f"### {analysis_section}. 工具变量与统计分析\n"
                    f"暴露相关SNP的筛选阈值为p < {threshold}，执行LD clumping（r² < 0.001，窗口10,000 kb）"
                    f"及等位基因协调；最终进入分析的工具变量为{result.n_instruments}个。"
                    f"当前结构化结果仅记录平均F统计量"
                    f"{f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'}，"
                    f"不据此虚构逐SNP的F值。实际估计方法为：{methods}。"
                    f"实际产生的敏感性分析为：{sensitivity_text}。"
                    "统计由R/TwoSampleMR流程执行；未由运行时记录的软件版本报告为N/A。"
                    "本次未提供正式功效或最小可检测效应计算，故不作相关推断。"
                )
            else:
                overlap = (
                    "Source metadata did not establish non-overlap and the runtime flagged possible sample overlap; "
                    "a two-sample design is therefore not treated as proof of independent cohorts."
                    if result.sample_overlap_warning else
                    "Source metadata was insufficient to verify cohort overlap, so no non-overlap claim is made."
                )
                blocks.append(
                    f"### {design_section}. Study design and data sources\n"
                    f"This two-sample Mendelian randomization analysis evaluated {exp} against {out}. "
                    f"The exposure GWAS was {result.exposure_id} (trait: {self._metadata_value(em, 'trait')}; "
                    f"sample size: {self._metadata_value(em, 'sample_size')}; population: {self._metadata_value(em, 'population')}; "
                    f"year: {self._metadata_value(em, 'year')}; total SNPs: {self._metadata_value(em, 'nsnp')}). "
                    f"The outcome GWAS was {result.outcome_id} (trait: {self._metadata_value(om, 'trait')}; "
                    f"sample size: {self._metadata_value(om, 'sample_size')}; population: {self._metadata_value(om, 'population')}; "
                    f"year: {self._metadata_value(om, 'year')}; total SNPs: {self._metadata_value(om, 'nsnp')}). {overlap}\n\n"
                    f"### {analysis_section}. Instruments and analysis\n"
                    f"Exposure-associated variants were selected at p < {threshold}, followed by LD clumping "
                    f"(r² < 0.001; 10,000 kb) and allele harmonization. The analysis retained "
                    f"{result.n_instruments} instruments. The runtime recorded a mean F-statistic of "
                    f"{f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'}; "
                    f"this is not represented as a per-SNP value. Estimation methods: {methods}. "
                    f"Observed sensitivity outputs: {sensitivity_text}. Analysis used R/TwoSampleMR; "
                    "unrecorded software versions are reported as N/A. No formal power or minimum-detectable-effect result was supplied."
                )
        return "\n\n".join(blocks)

    def _grounded_results(self, results: list[MRAnalysisResult]) -> str:
        """Render the complete statistical result without model-authored facts."""
        zh = self.language == "zh"
        blocks: list[str] = []
        for index, result in enumerate(results, start=1):
            exp = result.exposure_name or result.exposure_id
            out = result.outcome_name or result.outcome_id
            lines = [
                (f"### {index}. {exp} → {out}"),
                (f"GWAS ID：{result.exposure_id} → {result.outcome_id}；"
                 f"工具变量：{result.n_instruments}个；"
                 f"平均F统计量：{result.f_statistic_mean:.3f}。"
                 if zh and result.f_statistic_mean is not None else
                 f"GWAS IDs: {result.exposure_id} → {result.outcome_id}; instruments: {result.n_instruments}; "
                 f"mean F-statistic: {result.f_statistic_mean:.3f}."
                 if result.f_statistic_mean is not None else
                 f"GWAS ID{'s' if not zh else ''}: {result.exposure_id} → {result.outcome_id}; "
                 f"{'instruments' if not zh else '工具变量'}: {result.n_instruments}; mean F-statistic: N/A."),
            ]
            for method in result.mr_results:
                estimate = (
                    f"beta={method.beta:.4f}，SE={method.se:.4f}，p={self._fmt_p(method.pval)}"
                    if zh else
                    f"beta={method.beta:.4f}, SE={method.se:.4f}, p={self._fmt_p(method.pval)}"
                )
                if method.or_value is not None:
                    ci = (
                        f"OR={method.or_value:.3f}，95% CI {method.ci_lower:.3f}–{method.ci_upper:.3f}"
                        if method.ci_lower is not None and method.ci_upper is not None else
                        f"OR={method.or_value:.3f}"
                    )
                    estimate += (f"，{ci}" if zh else f", {ci}")
                lines.append(f"- {method.method}：{estimate}" if zh else f"- {method.method}: {estimate}")
            ivw = find_ivw(result.mr_results)
            if ivw:
                if zh:
                    direction = "正向" if ivw.beta > 0 else "负向"
                    significance = "统计学显著" if ivw.pval < 0.05 else "未达统计学显著"
                    lines.append(f"主分析效应为{direction}且{significance}；该结果提示关联方向，不单独等同于无条件因果证明。")
                else:
                    direction = "positive" if ivw.beta > 0 else "negative"
                    significance = "statistically significant" if ivw.pval < 0.05 else "not statistically significant"
                    lines.append(f"The primary estimate was {direction} and {significance}; it supports an association direction but is not unconditional proof of causation.")
            if result.heterogeneity:
                lines.append("异质性：" if zh else "Heterogeneity:")
                for item in result.heterogeneity:
                    lines.append(
                        f"- {item.method}：Q={item.q:.2f}，df={item.q_df}，p={self._fmt_p(item.q_pval)}"
                        if zh else
                        f"- {item.method}: Q={item.q:.2f}, df={item.q_df}, p={self._fmt_p(item.q_pval)}"
                    )
            if result.pleiotropy:
                p = result.pleiotropy
                lines.append(
                    f"MR-Egger截距={p.egger_intercept:.4f}，SE={p.se:.4f}，p={self._fmt_p(p.pval)}；"
                    "未检出显著方向性多效性，但不能排除平衡多效性。"
                    if zh else
                    f"MR-Egger intercept={p.egger_intercept:.4f}, SE={p.se:.4f}, p={self._fmt_p(p.pval)}; "
                    "no significant directional pleiotropy was detected, although balanced pleiotropy remains possible."
                )
            if result.steiger_correct is not None:
                lines.append(
                    f"Steiger方向检验：correct_causal_direction={str(result.steiger_correct).lower()}，"
                    f"p={self._fmt_p(result.steiger_pval) if result.steiger_pval is not None else 'N/A'}。"
                    if zh else
                    f"Steiger directionality: correct_causal_direction={str(result.steiger_correct).lower()}, "
                    f"p={self._fmt_p(result.steiger_pval) if result.steiger_pval is not None else 'N/A'}."
                )
            if result.presso_global_pval is not None:
                lines.append(
                    f"MR-PRESSO全局检验p={self._fmt_p(result.presso_global_pval)}，"
                    f"候选离群值={result.presso_n_outliers if result.presso_n_outliers is not None else 'N/A'}。"
                    if zh else
                    f"MR-PRESSO global p={self._fmt_p(result.presso_global_pval)}; candidate outliers="
                    f"{result.presso_n_outliers if result.presso_n_outliers is not None else 'N/A'}."
                )
            elif result.presso_n_outliers is not None:
                lines.append(
                    f"MR-PRESSO运行记录了{result.presso_n_outliers}个候选离群值，但全局检验p值未被解析；不据此声称全局检验阴性。"
                    if zh else
                    f"MR-PRESSO recorded {result.presso_n_outliers} candidate outliers, but the global-test p-value was not parsed; no negative global-test claim is made."
                )
            if result.sample_overlap_warning:
                lines.append(
                    "运行时已标记潜在样本重叠；效应量应结合显著异质性和队列重叠不确定性解读。"
                    if zh else
                    "The runtime flagged possible sample overlap; estimates should be interpreted with the significant heterogeneity and overlap uncertainty."
                )
            blocks.append("\n".join(lines))
        return "\n\n".join(blocks)

    def _grounded_abstract(self, results: list[MRAnalysisResult]) -> str:
        """Render an abstract whose design, estimates, and caveats are runtime facts."""
        zh = self.language == "zh"
        blocks: list[str] = []
        for result in results:
            exp = result.exposure_name or result.exposure_id
            out = result.outcome_name or result.outcome_id
            ivw = find_ivw(result.mr_results)
            estimate = "N/A"
            if ivw:
                if ivw.or_value is not None:
                    ci = (
                        f", 95% CI {ivw.ci_lower:.3f}–{ivw.ci_upper:.3f}"
                        if ivw.ci_lower is not None and ivw.ci_upper is not None else ""
                    )
                    estimate = f"OR={ivw.or_value:.3f}{ci}, p={self._fmt_p(ivw.pval)}"
                else:
                    estimate = f"beta={ivw.beta:.4f}, SE={ivw.se:.4f}, p={self._fmt_p(ivw.pval)}"
            heterogeneity = any(item.q_pval < 0.05 for item in result.heterogeneity)
            pleiotropy = (
                f"MR-Egger intercept p={self._fmt_p(result.pleiotropy.pval)}"
                if result.pleiotropy else "MR-Egger intercept=N/A"
            )
            presso = (
                f"MR-PRESSO global p={self._fmt_p(result.presso_global_pval)}"
                if result.presso_global_pval is not None else
                f"MR-PRESSO global p=N/A; candidate outliers={result.presso_n_outliers}"
                if result.presso_n_outliers is not None else
                "MR-PRESSO=N/A"
            )
            if zh:
                heterogeneity_label = "显著" if heterogeneity else "未达统计学显著"
                direction_label = "正" if ivw and ivw.beta > 0 else "负" if ivw else "无法判定"
                blocks.append(
                    f"背景：本分析评估遗传预测的{exp}与{out}之间的关联。\n"
                    f"方法：采用两样本孟德尔随机化；暴露GWAS {result.exposure_id}"
                    f"（n={self._metadata_value(result.exposure_metadata, 'sample_size')}）与结局GWAS {result.outcome_id}"
                    f"（n={self._metadata_value(result.outcome_metadata, 'sample_size')}）经等位基因协调后保留"
                    f"{result.n_instruments}个工具变量，以IVW为主分析。\n"
                    f"结果：IVW估计为{estimate}；平均F统计量为"
                    f"{f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'}。"
                    f"异质性检验{heterogeneity_label}；{pleiotropy}；{presso}。\n"
                    f"结论：主分析提示{direction_label}向关联。"
                    "该估计必须结合工具变量假设、异质性、多效性与样本重叠不确定性解读，"
                    "不单独等同于无条件因果证明。"
                )
            else:
                blocks.append(
                    f"Background: This analysis evaluated genetically predicted {exp} in relation to {out}.\n"
                    f"Methods: A two-sample Mendelian randomization design used exposure GWAS {result.exposure_id} "
                    f"(n={self._metadata_value(result.exposure_metadata, 'sample_size')}) and outcome GWAS "
                    f"{result.outcome_id} (n={self._metadata_value(result.outcome_metadata, 'sample_size')}); "
                    f"{result.n_instruments} harmonized instruments entered an IVW primary analysis.\n"
                    f"Results: The IVW estimate was {estimate}; the mean F-statistic was "
                    f"{f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'}. "
                    f"Heterogeneity was {'statistically significant' if heterogeneity else 'not statistically significant'}; "
                    f"{pleiotropy}; {presso}.\n"
                    f"Conclusion: The primary estimate indicated a "
                    f"{'positive' if ivw and ivw.beta > 0 else 'negative' if ivw else 'not estimable'} association. "
                    "It must be interpreted with the instrument assumptions, heterogeneity, pleiotropy, and overlap uncertainty; "
                    "it is not unconditional proof of causation."
                )
        return "\n\n".join(blocks)

    def _grounded_discussion(self, results: list[MRAnalysisResult]) -> str:
        """Interpret observed estimates without adding unexecuted methods or external facts."""
        zh = self.language == "zh"
        paragraphs: list[str] = []
        for result in results:
            exp = result.exposure_name or result.exposure_id
            out = result.outcome_name or result.outcome_id
            ivw = find_ivw(result.mr_results)
            if not ivw:
                continue
            sensitivity = [item for item in result.mr_results if item is not ivw]
            same_direction = bool(sensitivity) and all(item.beta * ivw.beta > 0 for item in sensitivity)
            significant_heterogeneity = any(item.q_pval < 0.05 for item in result.heterogeneity)
            directional = result.pleiotropy is not None and result.pleiotropy.pval < 0.05
            if zh:
                association = "正向" if ivw.beta > 0 else "负向"
                significance = "达统计学显著" if ivw.pval < 0.05 else "未达统计学显著"
                concordance = "方向一致" if same_direction else "未全部呈一致方向"
                heterogeneity_label = "显著" if significant_heterogeneity else "未达统计学显著"
                pleiotropy_label = (
                    "提示方向性多效性" if directional else
                    "未检出显著方向性多效性" if result.pleiotropy else
                    "未提供"
                )
                paragraphs.append(
                    f"主分析显示遗传预测的{exp}与{out}呈{association}关联，"
                    f"IVW估计{significance}（beta={ivw.beta:.4f}，p={self._fmt_p(ivw.pval)}）。"
                    f"其他{len(sensitivity)}种已实际运行的估计方法"
                    f"{concordance}。"
                    "这种方法间的一致性可作为稳健性信号，但不能修复共享偏倚或无效工具变量。"
                )
                paragraphs.append(
                    f"异质性检验{heterogeneity_label}。"
                    f"MR-Egger截距{pleiotropy_label}，"
                    "但不显著截距不排除平衡多效性。"
                    + (f"MR-PRESSO记录{result.presso_n_outliers}个候选离群值，但全局p值不可用，因此不作阴性结论。"
                       if result.presso_n_outliers is not None and result.presso_global_pval is None else "")
                )
                paragraphs.append(
                    ("运行时标记了潜在样本重叠，可能将估计推向观察性关联。"
                     if result.sample_overlap_warning else "源元数据不足以证明队列完全无重叠。")
                    + "由于本次未运行反向MR、多变量MR、非线性MR或正式功效分析，不对方向反转、"
                    + "独立性、剂量阈值或最小可检测效应作额外声称。该结果适合作为可追溯的因果推断证据，"
                    + "不直接生成个体诊疗建议或具体干预阈值。"
                )
            else:
                association = "positive" if ivw.beta > 0 else "negative"
                significance = "statistically significant" if ivw.pval < 0.05 else "not statistically significant"
                paragraphs.append(
                    f"The IVW analysis indicated a {association} association between genetically predicted {exp} and {out}; "
                    f"the estimate was {significance} (beta={ivw.beta:.4f}, p={self._fmt_p(ivw.pval)}). "
                    f"The {len(sensitivity)} other executed estimators were "
                    f"{'directionally concordant' if same_direction else 'not uniformly directionally concordant'}. "
                    "Concordance is a robustness signal but cannot repair shared bias or invalid instruments."
                )
                paragraphs.append(
                    f"Heterogeneity was {'statistically significant' if significant_heterogeneity else 'not statistically significant'}. "
                    f"The MR-Egger intercept {'suggested directional pleiotropy' if directional else 'did not detect significant directional pleiotropy' if result.pleiotropy else 'was unavailable'}; "
                    "a non-significant intercept does not exclude balanced pleiotropy. "
                    + (f"MR-PRESSO recorded {result.presso_n_outliers} candidate outliers, but its global p-value was unavailable and no negative conclusion is drawn. "
                       if result.presso_n_outliers is not None and result.presso_global_pval is None else "")
                )
                paragraphs.append(
                    ("The runtime flagged possible sample overlap, which may move estimates toward observational associations. "
                     if result.sample_overlap_warning else
                     "Source metadata was insufficient to establish complete cohort non-overlap. ")
                    + "Reverse, multivariable, and nonlinear MR and formal power analysis were not executed, so no additional claims are made about reversal, independence, thresholds, or minimum detectable effects. "
                    + "The result is causal-inference evidence, not a patient-level treatment recommendation or intervention threshold."
                )
        return "\n\n".join(paragraphs)

    def _grounded_conclusion(self, results: list[MRAnalysisResult]) -> str:
        zh = self.language == "zh"
        conclusions: list[str] = []
        for result in results:
            exp = result.exposure_name or result.exposure_id
            out = result.outcome_name or result.outcome_id
            ivw = find_ivw(result.mr_results)
            if not ivw:
                continue
            direction = (
                "正向" if zh and ivw.beta > 0 else "负向" if zh else
                "positive" if ivw.beta > 0 else "negative"
            )
            if zh:
                conclusions.append(
                    f"在{result.n_instruments}个协调后工具变量的两样本MR中，{exp}与{out}的IVW估计呈{direction}关联"
                    f"（beta={ivw.beta:.4f}，p={self._fmt_p(ivw.pval)}）。"
                    "该结果在遗传工具假设成立时支持相应方向的因果解释，但必须保留对异质性、"
                    "多效性、样本重叠和未运行扩展分析的限制；不由此推导具体剂量、阈值或个体化治疗决策。"
                )
            else:
                conclusions.append(
                    f"In this two-sample MR analysis of {result.n_instruments} harmonized instruments, the IVW estimate for {exp} and {out} was {direction} "
                    f"(beta={ivw.beta:.4f}, p={self._fmt_p(ivw.pval)}). Under the instrumental-variable assumptions, the result supports a causal interpretation in that direction, "
                    "but heterogeneity, pleiotropy, overlap uncertainty, and unexecuted extension analyses remain limitations; no dose, threshold, or patient-level treatment decision is inferred."
                )
        return "\n\n".join(conclusions)

    def _grounded_data_availability(self, results: list[MRAnalysisResult]) -> str:
        zh = self.language == "zh"
        rows = []
        for result in results:
            rows.append(
                f"{result.exposure_id} (https://gwas.mrcieu.ac.uk/datasets/{result.exposure_id}/)"
            )
            rows.append(
                f"{result.outcome_id} (https://gwas.mrcieu.ac.uk/datasets/{result.outcome_id}/)"
            )
        unique = list(dict.fromkeys(rows))
        if zh:
            return (
                "本次运行使用的GWAS标识符与仓库链接为：" + "；".join(unique) + "。"
                "实际可用性、访问条件和版本以源仓库为准。EviMed运行包保留请求、结构化结果、"
                "数据表、图形和报告供本地复核；本运行未声称另有公开GitHub代码仓库。"
            )
        return (
            "GWAS identifiers and repository links used in this run were: " + "; ".join(unique) + ". "
            "Availability, access conditions, and versions remain governed by the source repository. The EviMed run package retains the request, structured results, tables, figures, and report for local review; this run does not claim a separate public GitHub code repository."
        )

    def _grounded_ethics_statement(self) -> str:
        if self.language == "zh":
            return (
                "本次运行仅处理公开的去识别GWAS汇总统计量，未访问个体级数据。"
                "原始研究的伦理审批和知情同意状态应以各数据集原始出版物为准；EviMed运行时没有独立验证这些文件。"
                "在稿件提交或机构使用前，作者仍需按所在机构和期刊要求确认二次分析是否需额外审查。"
            )
        return (
            "This run processed only public, de-identified GWAS summary statistics and did not access individual-level data. "
            "Ethics approval and consent for the source studies must be verified in their original publications; the EviMed runtime did not independently verify those documents. "
            "Before submission or institutional use, authors remain responsible for confirming whether local and journal policies require additional review for this secondary analysis."
        )

    def _grounded_limitations(self, results: list[MRAnalysisResult]) -> str:
        zh = self.language == "zh"
        paragraphs: list[str] = []
        for result in results:
            exp_pop = self._metadata_value(result.exposure_metadata, "population")
            out_pop = self._metadata_value(result.outcome_metadata, "population")
            heterogeneity = any(item.q_pval < 0.05 for item in result.heterogeneity)
            if zh:
                paragraphs.append(
                    f"本分析仅记录平均F统计量"
                    f"{f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'}；"
                    "平均值不能替代每个SNP的强度判断。"
                    f"暴露与结局人群元数据分别为{exp_pop}和{out_pop}；"
                    "对未报告的结局人群不作欧洲血统推断，跨人群外推性需另行验证。"
                )
                paragraphs.append(
                    ("异质性检验显著，" if heterogeneity else "未检出显著异质性，")
                    + ("虽然MR-Egger截距未显著，仍不能排除平衡性或非相关多效性。"
                       if result.pleiotropy and result.pleiotropy.pval >= 0.05 else
                       "方向性多效性仍需审慎评估。")
                    + (f"MR-PRESSO记录了{result.presso_n_outliers}个候选离群值，但全局p值不可用，因此不将其解读为“无多效性”。"
                       if result.presso_n_outliers is not None and result.presso_global_pval is None else "")
                )
                paragraphs.append(
                    ("源数据无法证明暴露与结局队列完全不重叠，且运行时已标记潜在样本重叠；偏倚可能向观察性关联靠近。"
                     if result.sample_overlap_warning else
                     "队列级样本重叠信息不足，不能宣称完全无重叠。")
                    + "其他限制包括GWAS数据库选择偏倚、赢家诅咒、水平多效性、线性平均效应无法刻画阈值/非线性关系，以及未完成反向或多变量MR。"
                )
            else:
                paragraphs.append(
                    f"Only a mean F-statistic of {f'{result.f_statistic_mean:.3f}' if result.f_statistic_mean is not None else 'N/A'} was recorded; "
                    f"it is not a per-SNP strength assessment. Exposure and outcome population metadata were {exp_pop} and {out_pop}; "
                    "unreported ancestry is not inferred, and transportability requires external validation."
                )
                paragraphs.append(
                    ("Heterogeneity was statistically significant. " if heterogeneity else "Significant heterogeneity was not detected. ")
                    + ("A non-significant MR-Egger intercept does not exclude balanced or uncorrelated pleiotropy. "
                       if result.pleiotropy and result.pleiotropy.pval >= 0.05 else
                       "Directional pleiotropy remains uncertain. ")
                    + (f"MR-PRESSO recorded {result.presso_n_outliers} candidate outliers, but its global p-value was unavailable and is not interpreted as a negative test. "
                       if result.presso_n_outliers is not None and result.presso_global_pval is None else "")
                )
                paragraphs.append(
                    ("The source data did not establish complete cohort non-overlap and the runtime flagged possible overlap; bias may move estimates toward observational associations. "
                     if result.sample_overlap_warning else
                     "Cohort-level overlap information was insufficient, so complete non-overlap is not claimed. ")
                    + "Other limitations include GWAS selection bias, winner's curse, horizontal pleiotropy, linear-average effects that cannot describe thresholds, and the absence of reverse or multivariable MR."
                )
        return "\n\n".join(paragraphs)

    def _grounded_table1(self, results: list[MRAnalysisResult]) -> str:
        rows = ["| Characteristic | Exposure | Outcome |", "|---|---:|---:|"]
        for result in results:
            em, om = result.exposure_metadata, result.outcome_metadata
            values = [
                ("Trait", self._metadata_value(em, "trait"), self._metadata_value(om, "trait")),
                ("GWAS ID", result.exposure_id, result.outcome_id),
                ("Sample size", self._metadata_value(em, "sample_size"), self._metadata_value(om, "sample_size")),
                ("Population", self._metadata_value(em, "population"), self._metadata_value(om, "population")),
                ("Year", self._metadata_value(em, "year"), self._metadata_value(om, "year")),
                ("Total SNPs", self._metadata_value(em, "nsnp"), self._metadata_value(om, "nsnp")),
                ("Selected instruments", str(result.n_instruments), str(result.n_instruments)),
                ("Mean F-statistic", f"{result.f_statistic_mean:.3f}" if result.f_statistic_mean is not None else "N/A", "N/A"),
            ]
            rows.extend(f"| {name} | {exp} | {out} |" for name, exp, out in values)
        return "\n".join(rows)

    def _grounded_table2(self, results: list[MRAnalysisResult]) -> str:
        rows = [
            "| Exposure GWAS | Outcome GWAS | Method | nSNP | beta | SE | OR | 95% CI | p-value |",
            "|---|---|---|---:|---:|---:|---:|---|---:|",
        ]
        for result in results:
            for item in result.mr_results:
                ci = (
                    f"{item.ci_lower:.3f}–{item.ci_upper:.3f}"
                    if item.ci_lower is not None and item.ci_upper is not None else "N/A"
                )
                or_value = f"{item.or_value:.3f}" if item.or_value is not None else "N/A"
                rows.append(
                    f"| {result.exposure_id} | {result.outcome_id} | {item.method} | {item.nsnp} | "
                    f"{item.beta:.4f} | {item.se:.4f} | "
                    f"{or_value} | {ci} | {self._fmt_p(item.pval)} |"
                )
        return "\n".join(rows)

    def _enforce_structured_grounding(self, paper: dict[str, str]) -> dict[str, str]:
        results = self.state.analysis_results
        if not results:
            return paper
        paper["abstract"] = self._grounded_abstract(results)
        paper["methods"] = self._grounded_methods(results)
        paper["results"] = self._grounded_results(results)
        paper["discussion"] = self._grounded_discussion(results)
        paper["limitations"] = self._grounded_limitations(results)
        paper["conclusion"] = self._grounded_conclusion(results)
        paper["data_availability"] = self._grounded_data_availability(results)
        paper["ethics_statement"] = self._grounded_ethics_statement()
        paper["table1"] = self._grounded_table1(results)
        paper["table2"] = self._grounded_table2(results)
        return paper

    def _detect_affected_sections(
        self, feedback: str, body: dict[str, str],
    ) -> list[str]:
        """Identify which sections are explicitly mentioned in the feedback."""
        import re
        feedback_lower = feedback.lower()
        affected = []
        for k in body:
            if not body[k]:
                continue
            # Match section name as a whole word to avoid false positives
            # e.g. "results" won't match "the introduction presents results"
            pattern = rf'\b{re.escape(k.lower())}\s+(section|部分)'
            if re.search(pattern, feedback_lower):
                affected.append(k)
                continue
            # Also match capitalized section headers like "Results:" or "RESULTS"
            if k.upper() in feedback or k.capitalize() + ":" in feedback:
                affected.append(k)
        return affected

    def _revise_section(
        self, section_name: str, original: str, feedback: str,
    ) -> str:
        """Revise a single section based on coherence feedback."""
        from mr_agent.llm.prompts import PAPER_SECTION_REVISE
        prompt = PAPER_SECTION_REVISE.format(
            section_name=section_name,
            original_text=original[:2000],
            feedback=feedback,
        )
        return self._llm_chat(prompt, max_tokens=3000)

    def _write_extra_sections(self) -> dict[str, str]:
        """Write data availability and ethics statement."""
        self.on_progress("撰写数据声明...", 0.92)
        results = self.state.analysis_results
        lang = self.language
        data_avail = self._safe_write(
            sections.write_data_availability,
            self.llm, results, self.state.slots, language=lang,
        )
        ethics = self._safe_write(
            sections.write_ethics_statement, self.llm, results, language=lang,
        )
        return {
            "data_availability": data_avail,
            "ethics_statement": ethics,
        }

    def _fix_strobe_gaps(self, body: dict[str, str]) -> dict[str, str]:
        """Check STROBE-MR compliance and fix unaddressed items."""
        try:
            from mr_agent.analysis.strobe_mr import check_strobe_compliance
            report = check_strobe_compliance(body, self.llm)
            self.strobe_report = report
            unaddressed = [i for i in report.items if not i.addressed]
            if not unaddressed:
                return body
            section_gaps = _group_strobe_by_section(unaddressed)
            for section_name, items in section_gaps.items():
                if section_name not in body or not body[section_name]:
                    continue
                body[section_name] = self._patch_section_strobe(
                    section_name, body[section_name], items,
                )
        except Exception as e:
            logger.warning(f"STROBE-MR check failed, skipping: {e}")
        return body

    def _patch_section_strobe(
        self, section_name: str, text: str, items: list,
    ) -> str:
        """Patch a section to address missing STROBE-MR items."""
        from mr_agent.llm.prompts import STROBE_FIX_PROMPT
        missing = "\n".join(
            f"- [{i.item_id}] {i.description}" for i in items
        )
        prompt = STROBE_FIX_PROMPT.format(
            section_name=section_name,
            section_text=text[:2000],
            missing_items=missing,
        )
        return self._llm_chat(prompt, max_tokens=3000)

    def _assemble_dict(
        self, abstract: str, body: dict, tables: dict,
        extras: dict, refs,
    ) -> dict[str, str]:
        """Combine all sections into paper dict."""
        exp, out = self.state.slots.exposure, self.state.slots.outcome
        return {
            "title": self._generate_title(exp, out),
            "abstract": abstract,
            **body,
            **extras,
            **tables,
            "references": references.format_numbered(refs),
        }

    def save_paper(self, output_dir: Path | None = None) -> Path:
        """Save paper to text file."""
        output_dir = output_dir or self.state.output_dir or Path("mr_output")
        output_dir.mkdir(parents=True, exist_ok=True)
        paper = self.state.paper_sections
        if not paper:
            paper = self.generate()
        content = self._assemble_paper(paper)
        filepath = output_dir / "paper.txt"
        filepath.write_text(content, encoding="utf-8")
        logger.info(f"Paper saved to {filepath}")
        return filepath

    def save_paper_docx(self, output_dir: Path | None = None) -> Path:
        """Save paper as Word document."""
        output_dir = output_dir or self.state.output_dir or Path("mr_output")
        output_dir.mkdir(parents=True, exist_ok=True)
        paper = self.state.paper_sections
        if not paper:
            paper = self.generate()
        filepath = output_dir / "paper.docx"
        _write_docx(paper, filepath)
        return filepath

    def _gather_references(self) -> list[PaperReference]:
        """Gather references from PubMed for the paper."""
        exp = self.state.slots.exposure
        out = self.state.slots.outcome
        refs = []
        try:
            refs.extend(pubmed.search_papers(f"{exp}", max_results=10))
            refs.extend(pubmed.search_papers(f"{out}", max_results=10))
            refs.extend(pubmed.search_mr_studies(exp, out, max_results=10))
            # Foundational MR methodology papers
            for query in [
                "Burgess Mendelian randomization inverse variance weighted",
                "Bowden MR-Egger regression Mendelian randomization pleiotropy",
                "Bowden weighted median Mendelian randomization",
                "Verbanck MR-PRESSO pleiotropy residual sum outlier",
                "Hemani MR-Base TwoSampleMR Mendelian randomization",
            ]:
                refs.extend(pubmed.search_papers(query, max_results=2))
        except Exception as e:
            logger.warning(f"Reference gathering failed: {e}")
        return references.deduplicate(refs)

    def _generate_title(self, exposure: str | None, outcome: str | None) -> str:
        """Generate a neutral title that cannot contradict the measured estimate."""
        exposure = exposure or "unknown exposure"
        outcome = outcome or "unknown outcome"
        if self.language == "zh":
            return f"{exposure}与{outcome}：两样本孟德尔随机化分析"
        return f"{exposure} and {outcome}: a two-sample Mendelian randomization analysis"

    def _conclusion_brief(self, results: list[MRAnalysisResult]) -> str:
        """Brief conclusion for title generation."""
        sig_results = [
            r for r in results
            if (ivw := find_ivw(r.mr_results)) and ivw.pval < 0.05
        ]
        if sig_results:
            return "Evidence supports a causal relationship."
        return "No significant causal relationship was found."

    def _assemble_paper(self, paper: dict[str, str]) -> str:
        """Assemble all sections into a single document."""
        parts = []
        for section in SECTION_ORDER:
            text = paper.get(section, "")
            if not text:
                continue
            header = section.upper() if section != "title" else ""
            if header:
                parts.append(f"\n{'='*60}\n{header}\n{'='*60}\n")
            parts.append(text)
        return "\n\n".join(parts)


def _write_docx(paper: dict[str, str], filepath: Path) -> None:
    """Write paper as Word document."""
    doc = _init_docx_doc()
    doc.add_heading(paper.get("title", "MR Study"), level=0)
    _add_docx_sections(doc, paper)
    doc.save(str(filepath))
    logger.info(f"DOCX saved to {filepath}")


# STROBE-MR item → most relevant paper section mapping
_STROBE_SECTION_MAP = {
    "1a": "abstract", "1b": "abstract",
    "2": "introduction", "3": "introduction",
    "4": "methods", "5": "methods", "6a": "methods", "6b": "methods",
    "7": "methods", "8": "methods", "9": "methods", "10": "methods",
    "11": "methods",
    "12a": "results", "12b": "results", "13a": "results",
    "13b": "results", "14": "results",
    "15": "discussion", "16": "discussion", "17": "limitations",
    "18": "discussion", "19": "discussion",
    "20": "conclusion", "21": "results",
}


def _group_strobe_by_section(items: list) -> dict[str, list]:
    """Group unaddressed STROBE items by target section."""
    groups: dict[str, list] = {}
    for item in items:
        section = _STROBE_SECTION_MAP.get(item.item_id, "methods")
        groups.setdefault(section, []).append(item)
    return groups


def _init_docx_doc():
    """Create and configure a new Document."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"
    return doc


def _add_docx_sections(doc, paper: dict[str, str]) -> None:
    """Add all sections to a docx Document."""
    table_keys = {"table1", "table2"}
    section_map = [
        ("abstract", "Abstract"),
        ("introduction", "Introduction"),
        ("methods", "Methods"),
        ("results", "Results"),
        ("discussion", "Discussion"),
        ("limitations", "Limitations"),
        ("conclusion", "Conclusion"),
        ("data_availability", "Data Availability Statement"),
        ("ethics_statement", "Ethics Statement"),
        ("table1", "Table 1: Data Source Characteristics"),
        ("table2", "Table 2: MR Results Summary"),
        ("references", "References"),
    ]
    for key, heading in section_map:
        text = paper.get(key, "")
        if not text:
            continue
        doc.add_heading(heading, level=1)
        if key in table_keys:
            _add_docx_table(doc, text)
        else:
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())


def _parse_markdown_table(text: str) -> list[list[str]] | None:
    """Parse markdown table (| col | col |) into rows of cells."""
    lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
    # Accept lines starting with | (trailing | is optional)
    table_lines = [l for l in lines if l.startswith("|") and "|" in l[1:]]
    if len(table_lines) < 2:
        return None
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator lines (e.g., |---|---|)
        if all(set(c.strip()) <= {"-", ":"} for c in cells if c.strip()):
            continue
        rows.append(cells)
    return rows if len(rows) >= 2 else None


def _add_docx_table(doc, text: str) -> None:
    """Add markdown table as Word table, falling back to paragraphs."""
    # Try to find and parse markdown tables in the text
    sections_parts = text.split("\n\n")
    for part in sections_parts:
        rows = _parse_markdown_table(part)
        if rows:
            from docx.shared import Pt
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = "Table Grid"
            for i, row_data in enumerate(rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j < n_cols:
                        cell = row.cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                            paragraph.paragraph_format.space_after = Pt(0)
                        # Bold header row
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        else:
            # Not a table, add as paragraph
            stripped = part.strip()
            if stripped:
                doc.add_paragraph(stripped)
