"""Report rendering tests: required sections, number fidelity, exports."""

from __future__ import annotations

from datetime import datetime, timezone

import pytest

from safety_agent.analysis.models import (
    AnalysisResult,
    CaseOverview,
    CountBucket,
    Interpretation,
    NormalizedReaction,
    SignalRow,
)
from safety_agent.evidence.label_check import (
    LabelCheckReport,
    LabelCheckResult,
    LabelQuote,
)
from safety_agent.evidence.models import EvidenceLayerResult
from safety_agent.report.docx_export import export_docx, export_pdf
from safety_agent.report.markdown import render_markdown, signal_table_csv

# The signal row mirrors the hand-computed T1 panel (see
# test_signals_known_answers): ROR=10.444444..., PRR=9.5, chi2=51.473862....
T1_ROW = SignalRow(
    reaction="myalgia",
    source="user-specified",
    a=10, b=90, c=20, d=1880, n=2000,
    haldane_anscombe_applied=False,
    ror=10.444444444444445, ror_ci95_lower=4.749574678434984, ror_ci95_upper=22.96761860559505,
    prr=9.5, prr_ci95_lower=4.569056617568817, prr_ci95_upper=19.75243634604419,
    chi2=51.47386232077656,
    ic=2.736965594166206, ic025=1.7345667481436353,
    ebgm=6.353572889693643, eb05=3.4210649828906288,
    is_signal=True,
)


def _result(**overrides) -> AnalysisResult:
    base = dict(
        drug_query="Atorvastatin 20 mg",
        drug_normalized="atorvastatin",
        drug_candidates=["atorvastatin"],
        reactions=[
            NormalizedReaction(query="肌痛", normalized="myalgia", method="zh-map", confidence=1.0)
        ],
        language="zh",
        overview=CaseOverview(
            total_reports=2000,
            yearly=[CountBucket(term="2023", count=100), CountBucket(term="2024", count=120)],
            sex=[CountBucket(term="male", count=55), CountBucket(term="female", count=45)],
            age_buckets=[CountBucket(term="45-64", count=80)],
            outcomes=[CountBucket(term="death (死亡)", count=30)],
            countries=[CountBucket(term="us", count=80)],
            concomitant_drugs=[CountBucket(term="ASPIRIN", count=50)],
            indications=[CountBucket(term="HYPERTENSION", count=25)],
        ),
        signals=[T1_ROW],
        label_check=LabelCheckReport(
            drug="atorvastatin",
            status="ok",
            checks=[
                LabelCheckResult(
                    reaction="myalgia",
                    status="labeled",
                    quotes=[
                        LabelQuote(
                            section="adverse_reactions",
                            sentence="The most common adverse reactions are myalgia.",
                        )
                    ],
                )
            ],
            label_refs=["set-1 (LIPITOR)"],
        ),
        evidence=EvidenceLayerResult(enabled=False, note="证据检索层未启用。"),
        interpretation=Interpretation(
            overview="总览文字",
            demographics="人口学文字",
            outcomes="结局文字",
            signal_commentary="信号解读文字",
            label_commentary="说明书对照文字",
            focus_adrs=[{"reaction": "myalgia", "text": "重点段落文字"}],
        ),
        llm_status="ok",
        degradation_notes=["证据检索层未启用。"],
        query_urls={
            "drug_total": "https://api.fda.gov/drug/event.json?limit=1&search=x",
            "signal_joint[myalgia]": "https://api.fda.gov/drug/event.json?limit=1&search=y",
        },
        generated_at=datetime(2026, 7, 20, tzinfo=timezone.utc),
    )
    base.update(overrides)
    return AnalysisResult(**base)


def test_markdown_contains_all_required_sections():
    md = render_markdown(_result())
    for heading in (
        "## 1. 分析概览",
        "## 2. 输入归一",
        "## 3. 病例概览(FAERS)",
        "### 3.1 年度趋势",
        "### 3.3 结局分布",
        "### 3.5 合并用药",
        "## 4. 失比例信号分析",
        "## 5. 重点 ADR 解读",
        "## 6. 说明书对照(FDA label)",
        "## 7. 循证证据检索(EviMed)",
        "## 8. 局限性声明",
        "## 附录:数据来源与可追溯查询",
    ):
        assert heading in md, heading


def test_markdown_numbers_match_statistics_input():
    md = render_markdown(_result())
    # every figure comes from T1_ROW, formatted to 3 decimals
    assert "10.444 [4.750, 22.968]" in md
    assert "9.500 [4.569, 19.752]" in md
    assert "51.474" in md
    assert "2.737 (1.735)" in md
    assert "6.354 (3.421)" in md
    assert "2,000" in md  # total reports with thousands separator
    # LLM narrative is present but the numbers above are the renderer's own
    assert "总览文字" in md and "重点段落文字" in md
    # label cross-check verdict and verbatim quote
    assert "已标注" in md
    assert "The most common adverse reactions are myalgia." in md


def test_markdown_limitations_and_traceability():
    md = render_markdown(_result())
    assert "信号不等于因果关系" in md
    assert "不能用于推算不良反应发生率" in md
    assert "VigiBase" in md
    assert "https://api.fda.gov/drug/event.json?limit=1&search=y" in md
    assert "证据检索层未启用" in md
    assert "report_contains_suspect_approximation" in md
    assert "不得解释为 PS-only" in md
    assert "unfitted-starting-prior" in md


def test_markdown_frozen_snapshot_provenance_is_visible():
    md = render_markdown(
        _result(
            data_source="frozen_faers",
            suspect_binding="same_drug_object",
            suspect_roles=["PS"],
            ps_only=True,
            drug_field_used="frozen_normalized",
            study_date_from="2018-01-01",
            study_date_to="2020-12-31",
            snapshot_id="faers-2020q4-v1",
            snapshot_source="FDA quarterly files",
            snapshot_sha256="a" * 64,
        )
    )
    assert "冻结 FAERS 快照(faers-2020q4-v1)" in md
    assert "same_drug_object" in md
    assert "2018-01-01 至 2020-12-31" in md
    assert "FDA quarterly files" in md
    assert "`" + "a" * 64 + "`" in md
    assert "报告级近似" not in md


def test_live_report_discloses_report_level_binding_and_unfitted_prior():
    md = render_markdown(_result())
    assert "openFDA live API" in md
    assert "report_contains_suspect_approximation" in md
    assert "不得解释为 PS-only" in md
    assert "unfitted-starting-prior" in md
    assert "不能标作已完成全矩阵经验贝叶斯拟合" in md


def test_frozen_report_discloses_snapshot_scope_and_fitted_prior():
    md = render_markdown(
        _result(
            data_source="frozen_faers",
            suspect_binding="same_drug_object",
            suspect_roles=["PS"],
            study_date_from="2024-01-01",
            study_date_to="2024-12-31",
            snapshot_id="faers-2024q4",
            snapshot_source="FDA quarterly ASCII files",
            snapshot_sha256="a" * 64,
            gps_prior_fitted=True,
            gps_prior_id="prior-2024q4",
        )
    )
    assert "冻结 FAERS 快照(faers-2024q4)" in md
    assert "same_drug_object" in md
    assert "2024-01-01 至 2024-12-31" in md
    assert "GPS prior=fitted (prior-2024q4)" in md
    assert "FDA quarterly ASCII files" in md
    assert "`" + "a" * 64 + "`" in md
    assert "不得解释为 PS-only" not in md
    assert "未拟合的 GPS" not in md


def test_markdown_degraded_run_shows_methodology_note():
    result = _result(interpretation=None, llm_status="degraded",
                     degradation_notes=["LLM 解读失败(backend down),报告降级为仅统计结果。"])
    md = render_markdown(result)
    assert "方法学声明" in md
    assert "LLM 解读缺失" in md
    # statistics survive degradation
    assert "10.444 [4.750, 22.968]" in md


def test_signal_csv_same_data_source():
    import csv as csv_module
    import io

    csv_text = signal_table_csv(_result())
    rows = list(csv_module.reader(io.StringIO(csv_text)))
    assert rows[0][:7] == ["reaction", "source", "a", "b", "c", "d", "N"]
    assert len(rows) == 2
    cells = rows[1]
    assert cells[0] == "myalgia"
    assert cells[2:7] == ["10", "90", "20", "1,880", "2,000"]
    assert cells[7] == "10.444"
    assert cells[-1] == "yes"
    assert rows[0][-4:] == [
        "expected_count",
        "haldane_anscombe_applied",
        "gps_prior_id",
        "is_signal",
    ]


def test_docx_export_roundtrip(tmp_path):
    path = export_docx(_result(), tmp_path / "report.docx")
    assert path.is_file() and path.stat().st_size > 5000
    from docx import Document

    document = Document(path)
    headings = [
        p.text
        for p in document.paragraphs
        if p.style.name.startswith(("Heading", "Title"))
    ]
    assert any("atorvastatin" in h for h in headings)
    assert any("失比例信号分析" in h for h in headings)
    assert any("局限性声明" in h for h in headings)
    tables_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "myalgia" in tables_text
    assert "10.444" in tables_text


def test_pdf_export_via_libreoffice(tmp_path):
    docx = export_docx(_result(), tmp_path / "report.docx")
    pdf = export_pdf(docx, tmp_path)
    if pdf is None:
        pytest.skip("LibreOffice not installed on this machine")
    assert pdf.is_file() and pdf.stat().st_size > 10000
    assert pdf.read_bytes()[:4] == b"%PDF"
