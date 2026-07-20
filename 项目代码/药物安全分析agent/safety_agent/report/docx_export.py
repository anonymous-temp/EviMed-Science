"""docx export (python-docx) and pdf export (LibreOffice headless).

The docx mirrors the markdown sections, built from the same AnalysisResult
— never by re-parsing the markdown. PDF conversion shells out to a local
LibreOffice (``EVIMED_OFFICE_HOME`` override, then the standard macOS
install path, then PATH); when LibreOffice is absent the pdf is skipped
with a warning instead of failing.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from safety_agent.analysis.models import AnalysisResult, CountBucket
from safety_agent.core.logging import get_logger
from safety_agent.report.markdown import _f, _int

logger = get_logger(__name__)

_LIBREOFFICE_MACOS = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")


def export_docx(result: AnalysisResult, path: Path) -> Path:
    """Build a .docx version of the report at ``path``."""
    from docx import Document  # python-docx, imported lazily for testability

    document = Document()
    document.add_heading(f"{result.drug_normalized} — FAERS 药物安全性分析报告", level=0)
    data_source = (
        f"冻结 FAERS 快照({result.snapshot_id or 'ID not provided'})"
        if result.data_source == "frozen_faers"
        else "FAERS via openFDA live API"
    )
    meta = document.add_paragraph()
    meta.add_run(
        f"生成时间:{result.generated_at:%Y-%m-%d %H:%M UTC} | "
        f"数据源:{data_source} | 指标:ROR/PRR/χ²/IC/GPS-EBGM"
    ).italic = True

    _docx_overview(document, result)
    _docx_case_profile(document, result)
    _docx_signals(document, result)
    _docx_focus(document, result)
    _docx_label_check(document, result)
    _docx_limitations(document, result)
    _docx_appendix(document, result)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    logger.info("docx report written to %s", path)
    return path


def export_pdf(docx_path: Path, out_dir: Path, *, timeout: float = 180.0) -> Path | None:
    """Convert the docx to pdf via headless LibreOffice; None when skipped.

    Each invocation gets its own throwaway LibreOffice user profile:
    concurrent ``soffice`` processes share a profile lock otherwise and
    silently fail to produce the pdf (observed when generating samples in
    parallel).
    """
    soffice = _find_soffice()
    if soffice is None:
        logger.warning("LibreOffice not found; pdf export skipped (docx kept)")
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="soffice-profile-") as profile:
        try:
            completed = subprocess.run(
                [
                    str(soffice),
                    f"-env:UserInstallation=file://{profile}",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(out_dir),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            logger.warning("LibreOffice pdf conversion failed to run: %s", exc)
            return None
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if completed.returncode != 0 or not pdf_path.is_file():
        logger.warning(
            "LibreOffice pdf conversion failed (rc=%s): %s",
            completed.returncode,
            (completed.stderr or completed.stdout)[:300],
        )
        return None
    logger.info("pdf report written to %s", pdf_path)
    return pdf_path


def _find_soffice() -> Path | None:
    office_home = os.environ.get("EVIMED_OFFICE_HOME")
    candidates: list[Path] = []
    if office_home:
        candidates.append(Path(office_home) / "Contents" / "MacOS" / "soffice")
        candidates.append(Path(office_home) / "soffice")
    candidates.append(_LIBREOFFICE_MACOS)
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    which = shutil.which("soffice")
    return Path(which) if which else None


# -- sections (mirror of markdown.py) --------------------------------------------


def _docx_overview(document, result: AnalysisResult) -> None:
    document.add_heading("1. 分析概览", level=1)
    requested = [r.normalized for r in result.reactions if r.normalized]
    document.add_paragraph(f"目标药品:{result.drug_normalized}(输入:{result.drug_query})")
    document.add_paragraph("目标 ADR:" + ("、".join(requested) if requested else "未指定"))
    document.add_paragraph(f"FAERS 报告总数:{_int(result.overview.total_reports)} 份")
    document.add_paragraph(f"药品角色口径:{'、'.join(result.suspect_roles)}")
    document.add_paragraph(f"药名/角色绑定:{result.suspect_binding}")
    if result.study_date_from or result.study_date_to:
        document.add_paragraph(
            f"研究时间窗:{result.study_date_from or 'open'} 至 "
            f"{result.study_date_to or 'open'}"
        )
    prior_label = "fitted" if result.gps_prior_fitted else "unfitted-starting-prior"
    prior_id = f" ({result.gps_prior_id})" if result.gps_prior_id else ""
    document.add_paragraph(
        f"统计版本:{result.statistics_version};GPS prior={prior_label}{prior_id}"
    )
    if result.snapshot_id:
        document.add_paragraph(f"冻结快照:{result.snapshot_id}")
    if result.interpretation and result.interpretation.overview:
        document.add_paragraph(result.interpretation.overview)


def _docx_case_profile(document, result: AnalysisResult) -> None:
    overview = result.overview
    document.add_heading("2. 病例概览(FAERS)", level=1)
    _docx_buckets(document, "年度趋势", overview.yearly)
    _docx_buckets(document, "性别分布", overview.sex)
    _docx_buckets(document, "年龄段(近似)", overview.age_buckets)
    _docx_buckets(document, "结局分布", overview.outcomes)
    _docx_buckets(document, "报告国别 top10", overview.countries)
    _docx_buckets(document, "合并用药 top10", overview.concomitant_drugs)
    _docx_buckets(document, "适应症 top10", overview.indications)


def _docx_signals(document, result: AnalysisResult) -> None:
    document.add_heading("3. 失比例信号分析", level=1)
    document.add_paragraph(
        "2×2 定义:a=目标药且目标 ADR,b=目标药其他 ADR,c=其他药目标 ADR,d=N−a−b−c。"
        "信号判定:a≥3 且(ROR 95%CI 下限>1 或(PRR≥2 且 χ²≥4))。"
    )
    headers = ["ADR", "来源", "a", "ROR [95%CI]", "PRR", "χ²", "IC (IC025)", "EBGM (EB05)", "信号"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = text
    for row in result.signals:
        cells = table.add_row().cells
        values = [
            row.reaction,
            "指定" if row.source == "user-specified" else "top",
            _int(row.a),
            f"{_f(row.ror)} [{_f(row.ror_ci95_lower)}, {_f(row.ror_ci95_upper)}]",
            _f(row.prr),
            _f(row.chi2),
            f"{_f(row.ic)} ({_f(row.ic025)})",
            f"{_f(row.ebgm)} ({_f(row.eb05)})",
            "是" if row.is_signal else "否",
        ]
        for cell, text in zip(cells, values, strict=True):
            cell.text = text
    if result.interpretation and result.interpretation.signal_commentary:
        document.add_paragraph(result.interpretation.signal_commentary)


def _docx_focus(document, result: AnalysisResult) -> None:
    document.add_heading("4. 重点 ADR 解读", level=1)
    itp = result.interpretation
    if itp and itp.focus_adrs:
        for focus in itp.focus_adrs:
            document.add_heading(focus.reaction, level=2)
            document.add_paragraph(focus.text)
    else:
        document.add_paragraph(
            f"方法学声明:本次运行 LLM 解读缺失({result.llm_status}),本节仅保留统计结果。"
        )


def _docx_label_check(document, result: AnalysisResult) -> None:
    document.add_heading("5. 说明书对照(FDA label)", level=1)
    check = result.label_check
    if check is None or check.status != "ok":
        note = check.note if check else "未执行"
        document.add_paragraph(f"说明书对照未完成或不可用:{note or '—'}")
        return
    status_zh = {"labeled": "已标注", "partially_labeled": "部分标注", "unlabeled": "未标注"}
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["ADR", "标注状态", "原文引用(章节)"], strict=True):
        cell.text = text
    for item in check.checks:
        cells = table.add_row().cells
        quotes = "\n".join(f"“{q.sentence}”({q.section})" for q in item.quotes) or "—"
        for cell, text in zip(
            cells, [item.reaction, status_zh.get(item.status, item.status), quotes], strict=True
        ):
            cell.text = text


def _docx_limitations(document, result: AnalysisResult) -> None:
    from safety_agent.report.markdown import _BASE_LIMITATIONS

    document.add_heading("6. 局限性声明", level=1)
    for item in _BASE_LIMITATIONS:
        document.add_paragraph(item, style="List Number")
    if result.suspect_binding == "report_contains_suspect_approximation":
        document.add_paragraph(
            "openFDA live 聚合无法把目标药名与 suspect 角色绑定到同一 drug 对象;"
            "当前口径是报告级近似,不是 PS-only。",
            style="List Number",
        )
    elif result.suspect_binding == "target_name_only":
        document.add_paragraph(
            "live 队列只按目标药名筛选,未限定该药在报告中的 PS/SS/C/I 角色。",
            style="List Number",
        )
    if not result.gps_prior_fitted:
        document.add_paragraph(
            "GPS prior 未在完整快照矩阵上拟合,EBGM/EB05 仅供探索。",
            style="List Number",
        )
    for note in result.degradation_notes:
        document.add_paragraph(f"降级项:{note}", style="List Bullet")


def _docx_appendix(document, result: AnalysisResult) -> None:
    document.add_heading("附录:数据来源与可追溯查询", level=1)
    if result.snapshot_id:
        document.add_paragraph(f"snapshot_id: {result.snapshot_id}", style="List Bullet")
        document.add_paragraph(
            f"snapshot_source: {result.snapshot_source or 'not provided'}",
            style="List Bullet",
        )
        document.add_paragraph(
            f"snapshot_sha256: {result.snapshot_sha256 or 'not provided'}",
            style="List Bullet",
        )
        document.add_paragraph(
            f"snapshot_extracted_at: {result.snapshot_extracted_at or 'not provided'}",
            style="List Bullet",
        )
        document.add_paragraph(
            f"snapshot_deduplication: {result.snapshot_deduplication or 'not provided'}",
            style="List Bullet",
        )
    for key, url in result.query_urls.items():
        document.add_paragraph(f"{key}: {url}", style="List Bullet")


def _docx_buckets(document, title: str, buckets: list[CountBucket]) -> None:
    document.add_heading(title, level=2)
    if not buckets:
        document.add_paragraph("(无数据)")
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项"
    table.rows[0].cells[1].text = "报告数"
    for bucket in buckets:
        cells = table.add_row().cells
        cells[0].text = str(bucket.term)
        cells[1].text = _int(bucket.count)
