"""
Document Parser Service - Enhanced with Marker for high-quality PDF parsing

This service extracts text from manuscript formats (.docx, .pdf, .txt, .md)
and produces high-quality Markdown output that preserves:
- Document structure (sections, headings)
- Tables and their content
- Figure captions
- Multi-column layouts

Primary method: Marker (deep learning-based)
Fallback: pypdf for basic extraction
"""
import os
import re
from typing import Optional, Tuple, List, Dict, Any
from pathlib import Path

# Try to import Marker for high-quality PDF parsing
try:
    from marker.converters.pdf import PdfConverter
    from marker.models import create_model_dict
    from marker.output import text_from_rendered
    MARKER_AVAILABLE = True
except ImportError:
    MARKER_AVAILABLE = False
    print("Info: Marker not installed. PDF parsing will use fallback pypdf method.")
    print("For best results, install marker: pip install marker-pdf")

# Fallback parsers
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class DocumentParser:
    """
    Service for parsing manuscript files in various formats (.docx, .pdf, .txt, .md)
    and extracting clean, structured text content.

    Features:
    - High-quality PDF parsing via Marker (preserves tables, structure)
    - DOCX parsing with table extraction
    - Text file support
    - Coverage tracking for downstream Evidence Gate
    """

    SUPPORTED_FORMATS = [".docx", ".pdf", ".txt", ".md"]
    MAX_FILE_SIZE_MB = 50

    def __init__(self, use_marker: bool = True):
        """
        Initialize document parser.

        Args:
            use_marker: Whether to use Marker for PDF parsing (if available)
        """
        self._check_dependencies()
        self.use_marker = use_marker and MARKER_AVAILABLE
        self._marker_models = None

    def _check_dependencies(self):
        """Check if required parsing libraries are available"""
        if not DOCX_AVAILABLE:
            print("Warning: python-docx not installed. .docx parsing will fail.")
        if not PDF_AVAILABLE and not MARKER_AVAILABLE:
            print("Warning: Neither marker nor pypdf installed. .pdf parsing will fail.")

    def _init_marker_models(self):
        """Lazy initialization of Marker models (expensive)"""
        if self._marker_models is None and MARKER_AVAILABLE:
            try:
                self._marker_models = create_model_dict()
            except Exception as e:
                print(f"Warning: Failed to initialize Marker models: {e}")
                self._marker_models = False  # Mark as failed, don't retry

    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a manuscript file and extract text content.

        Args:
            file_path: Path to the manuscript file

        Returns:
            Tuple of (text_content, metadata)
            - text_content: Markdown-formatted text
            - metadata: Dict with parsing info and coverage data

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format not supported or file too large
            RuntimeError: If parsing fails
        """
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File size ({file_size_mb:.2f} MB) exceeds maximum ({self.MAX_FILE_SIZE_MB} MB)"
            )

        # Get file extension
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: {file_ext}. Supported: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        # Parse based on file type
        try:
            if file_ext == ".docx":
                text, metadata = self._parse_docx(file_path)
            elif file_ext == ".pdf":
                text, metadata = self._parse_pdf(file_path)
            elif file_ext in {".txt", ".md"}:
                text, metadata = self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_ext}")

            # Clean and validate text
            text = self._clean_text(text)
            if not text or len(text.strip()) < 100:
                raise RuntimeError("Extracted text is empty or too short (< 100 characters)")

            # Add common metadata
            metadata["file_path"] = file_path
            metadata["file_size_mb"] = file_size_mb
            metadata["format"] = file_ext

            return text, metadata

        except Exception as e:
            raise RuntimeError(f"Failed to parse document: {str(e)}")

    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse PDF file using Marker (primary) or pypdf (fallback).

        Marker produces high-quality Markdown with:
        - Preserved headings and structure
        - Extracted tables
        - Figure captions
        - Multi-column layout handling
        """
        # Try Marker first if available and enabled
        if self.use_marker and MARKER_AVAILABLE:
            try:
                return self._parse_pdf_with_marker(file_path)
            except Exception as e:
                print(f"Marker parsing failed, falling back to pypdf: {e}")

        # Fallback to pypdf
        return self._parse_pdf_with_pypdf(file_path)

    def _parse_pdf_with_marker(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF using Marker for high-quality extraction"""
        self._init_marker_models()

        if not self._marker_models:
            raise RuntimeError("Marker models not available")

        # Create converter and process
        converter = PdfConverter(artifact_dict=self._marker_models)
        rendered = converter(file_path)

        # Extract text from rendered output
        text, _, images = text_from_rendered(rendered)

        # Build metadata with coverage information
        metadata = {
            "parse_method": "marker",
            "parse_quality": "high",
            "page_count": len(rendered.children) if hasattr(rendered, 'children') else 0,
            "tables_found": text.count("| "),  # Rough table count from markdown
            "figures_found": len(images) if images else 0,
            "has_structure": True,
        }

        # Extract table data for coverage tracking
        tables_data = self._extract_tables_from_markdown(text)
        metadata["tables_data"] = tables_data

        return text, metadata

    def _parse_pdf_with_pypdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF using pypdf (fallback method)"""
        if not PDF_AVAILABLE:
            raise ImportError("pypdf not installed. Run: pip install pypdf")

        reader = PdfReader(file_path)

        # Extract text from all pages
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)

        text = "\n\n".join(pages_text)

        # Extract metadata
        pdf_info = reader.metadata
        metadata = {
            "parse_method": "pypdf",
            "parse_quality": "low",  # pypdf doesn't preserve structure well
            "title": pdf_info.get("/Title", "") if pdf_info else "",
            "author": pdf_info.get("/Author", "") if pdf_info else "",
            "created": str(pdf_info.get("/CreationDate", "")) if pdf_info else "",
            "page_count": len(reader.pages),
            "has_structure": False,
            "tables_found": 0,  # pypdf doesn't extract tables
            "figures_found": 0,
        }

        return text, metadata

    def _parse_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse .docx file with table extraction"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(file_path)

        # Extract paragraphs as Markdown
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings based on style
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                paragraphs.append(f"# {text}")
            elif "Heading 2" in style_name:
                paragraphs.append(f"## {text}")
            elif "Heading 3" in style_name:
                paragraphs.append(f"### {text}")
            else:
                # 兜底：识别无 Heading 样式但形式上像标题的段落
                # 条件：短段落（<120字符）+ 全部加粗 + 不含句号
                is_bold_heading = (
                    len(text) < 120
                    and text.rstrip().rstrip('.')  # 末尾无句号
                    and para.runs
                    and all(run.bold for run in para.runs if run.text.strip())
                    and '.' not in text[-2:]  # 末尾2字符无句号
                )
                # 识别数字编号标题：如 "1. Introduction"、"2 Mechanisms"
                is_numbered_heading = bool(
                    len(text) < 120
                    and __import__('re').match(r'^\d+[\.\s]\s*\w', text)
                    and text.count('.') <= 2  # 不是正文句子
                )
                if is_bold_heading or is_numbered_heading:
                    paragraphs.append(f"## {text}")
                else:
                    paragraphs.append(text)

        # Extract tables as Markdown
        tables_markdown = []
        tables_data = []
        for table_idx, table in enumerate(doc.tables, 1):
            table_md, table_info = self._table_to_markdown(table, table_idx)
            tables_markdown.append(table_md)
            tables_data.append(table_info)

        # Combine paragraphs and tables
        text = "\n\n".join(paragraphs)
        if tables_markdown:
            text += "\n\n## Tables\n\n" + "\n\n".join(tables_markdown)

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "parse_method": "docx",
            "parse_quality": "high",
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "tables_found": len(doc.tables),
            "tables_data": tables_data,
            "figures_found": 0,  # DOCX figure extraction needs special handling
            "has_structure": True,
        }

        return text, metadata

    def _parse_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse a plain-text or Markdown manuscript."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        is_markdown = Path(file_path).suffix.lower() == ".md"

        metadata = {
            "parse_method": "markdown" if is_markdown else "txt",
            "parse_quality": "medium",
            "line_count": text.count("\n"),
            "character_count": len(text),
            "has_structure": is_markdown,
            "tables_found": 0,
            "figures_found": 0,
        }

        return text, metadata

    def _table_to_markdown(self, table, table_number: int) -> Tuple[str, Dict[str, Any]]:
        """Convert a DOCX table to Markdown format"""
        rows = []
        max_cols = 0

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            max_cols = max(max_cols, len(cells))
            rows.append(cells)

        if not rows:
            return "", {"table_number": table_number, "row_count": 0, "column_count": 0}

        # Build Markdown table
        md_lines = [f"**Table {table_number}**", ""]

        # Header row
        header = rows[0] if rows else []
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # Data rows
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

        table_info = {
            "table_number": table_number,
            "row_count": len(rows),
            "column_count": max_cols,
            "headers": rows[0] if rows else [],
        }

        return "\n".join(md_lines), table_info

    def _extract_tables_from_markdown(self, text: str) -> List[Dict[str, Any]]:
        """Extract table information from Markdown text"""
        tables = []
        # Simple table detection based on Markdown pipe syntax
        table_pattern = r'\|[^\n]+\|\n\|[-:\s|]+\|(?:\n\|[^\n]+\|)+'
        matches = re.findall(table_pattern, text)

        for idx, match in enumerate(matches, 1):
            lines = match.strip().split("\n")
            row_count = len(lines) - 1  # Exclude separator row
            col_count = lines[0].count("|") - 1 if lines else 0

            tables.append({
                "table_number": idx,
                "row_count": row_count,
                "column_count": col_count,
            })

        return tables

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by:
        - Removing excessive whitespace
        - Normalizing line breaks
        - Removing hidden/invisible characters
        """
        # Remove zero-width characters (potential security issue)
        text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)

        # Normalize whitespace
        text = re.sub(r'\r\n', '\n', text)  # Windows line endings
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double newline

        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate if a file can be parsed without actually parsing it.

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check existence
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"

        # Check format
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            return False, f"Unsupported format: {file_ext}"

        # Check size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.MAX_FILE_SIZE_MB:
            return False, f"File too large: {file_size_mb:.2f} MB (max: {self.MAX_FILE_SIZE_MB} MB)"

        # Check if file is readable
        try:
            with open(file_path, "rb") as f:
                f.read(1)
        except Exception as e:
            return False, f"File not readable: {str(e)}"

        return True, None

    def get_parser_info(self) -> Dict[str, Any]:
        """Get information about available parsers"""
        return {
            "marker_available": MARKER_AVAILABLE,
            "marker_enabled": self.use_marker,
            "pypdf_available": PDF_AVAILABLE,
            "docx_available": DOCX_AVAILABLE,
            "supported_formats": self.SUPPORTED_FORMATS,
        }
